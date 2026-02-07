from __future__ import annotations
import os
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, make_response, send_file
import sqlite3
from werkzeug.security import generate_password_hash, check_password_hash
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.utils import simpleSplit
from io import BytesIO
from datetime import datetime
from checklist import bp_checklist, ensure_checklist_tables

# === Lista de Presença (DOCX/PDF) ===
import zipfile
from docx import Document

from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from biblioteca import bp_biblioteca
from biblioteca import conectar_bd_biblioteca
from reportlab.lib.pagesizes import A4, landscape
from carometro import bp_carometro, init_carometro_db
from conselho import bp_conselho, ensure_conselho_tables
import re
from soe import bp_soe, ensure_soe_table
from termo import bp_termo, ensure_termo_tables, get_termo_ativo, registrar_aceite
from rotina import bp_rotina, ensure_rotina_tables

app = Flask(__name__)
app.secret_key = 'sua_chave_secreta'

# Disciplinas disponíveis para o professor
DISCIPLINAS = [
    'Matemática',
    'Português',
    'Ciências',
    'Geografia',
    'História',
    'Artes',
    'Inglês',
    'Educação Física'
]


# Funções auxiliares de banco

def conectar_bd():
    conn = sqlite3.connect('rfa.db', check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def extrair_nomes_alunos_do_pdf(file_storage):
    file_storage.stream.seek(0)  # 🔧 garante leitura desde o início
    nomes = []
    vistos = set()
    """
    Extrai somente os NOMES do PDF de enturmação (modelo SEEDF/CRE).
    Ignora código, telefone e demais dados.
    Retorna lista na ordem em que aparece no PDF (sem repetidos).
    """
    nomes = []
    vistos = set()

    # Tentativa 1: pdfplumber (recomendado)
    try:
        import pdfplumber
        with pdfplumber.open(file_storage.stream) as pdf:
            texto_total = ""
            for page in pdf.pages:
                texto_total += (page.extract_text() or "") + "\n"
    except Exception:
        # Tentativa 2: PyPDF2 (fallback)
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_storage.stream)
            texto_total = ""
            for p in reader.pages:
                texto_total += (p.extract_text() or "") + "\n"
        except Exception as e:
            raise RuntimeError(f"Não foi possível ler o PDF. Erro: {e}")

    # Normaliza e percorre linha a linha
    for raw in texto_total.splitlines():
        linha = (raw or "").strip()
        if not linha:
            continue

        # Padrão típico do PDF: "123456 NOME COMPLETO (61) 9..."
        m = re.match(r"^\s*\d+\s+(.+?)\s*(?:\(|$)", linha)
        if not m:
            continue

        candidato = m.group(1).strip()

        # Filtra cabeçalhos do relatório caso apareçam como "nome"
        # (no seu PDF vem "Código Nome do Estudante/Nome Social Telefone")
        if "NOME DO ESTUDANTE" in candidato.upper():
            continue

        # Limpa múltiplos espaços
        candidato = re.sub(r"\s{2,}", " ", candidato).strip()

        # Evita duplicados (case-insensitive)
        key = candidato.upper()
        if key not in vistos:
            vistos.add(key)
            nomes.append(candidato)

    return nomes


def inicializar_bd():
    conn = conectar_bd()
    cursor = conn.cursor()

    cursor.execute("PRAGMA foreign_keys = ON")

    # Tabela de professores
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS professores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            login TEXT UNIQUE NOT NULL,
            senha TEXT NOT NULL,
            status TEXT DEFAULT 'pendente',
            turma_id INTEGER DEFAULT NULL,
            FOREIGN KEY (turma_id) REFERENCES turmas(id)
        )
    ''')

    # Tabela de turmas
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS turmas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            turno TEXT NOT NULL,
            UNIQUE (nome, turno)
        )
    ''')

    # Tabela intermediária professor–turma (para vários vínculos)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS professores_turmas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            FOREIGN KEY (professor_id) REFERENCES professores(id) ON DELETE CASCADE,
            FOREIGN KEY (turma_id) REFERENCES turmas(id) ON DELETE CASCADE,
            UNIQUE (professor_id, turma_id)
        )
    ''')

    # Turmas do professor POR disciplina (vínculo correto para status por disciplina)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS professor_turmas_disciplina (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            disciplina TEXT NOT NULL,
            turma_id INTEGER NOT NULL,
            FOREIGN KEY (professor_id) REFERENCES professores(id),
            FOREIGN KEY (turma_id) REFERENCES turmas(id),
            UNIQUE (professor_id, disciplina, turma_id)
        )
    ''')

    # Moderadores
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS moderadores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            login TEXT UNIQUE NOT NULL,
            senha TEXT NOT NULL,
            tipo TEXT DEFAULT 'coordenador'
        )
    ''')

    # Alunos
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS alunos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            turma_id INTEGER NOT NULL,
            FOREIGN KEY (turma_id) REFERENCES turmas(id)
        )
    ''')

    # Transferências de alunos (log/auditoria)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS transferencias_alunos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            aluno_id INTEGER NOT NULL,
            turma_origem_id INTEGER NOT NULL,
            turma_destino_id INTEGER NOT NULL,
            motivo TEXT,
            transferido_por TEXT,
            data_transferencia TEXT DEFAULT (datetime('now','localtime')),
            FOREIGN KEY (aluno_id) REFERENCES alunos(id),
            FOREIGN KEY (turma_origem_id) REFERENCES turmas(id),
            FOREIGN KEY (turma_destino_id) REFERENCES turmas(id)
        )
    ''')

    # Atestados (comprovantes médicos, comparecimento etc.)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS atestados (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            bimestre TEXT NOT NULL,
            turma_id INTEGER NOT NULL,
            aluno_id INTEGER NOT NULL,
            tipo_atestado TEXT NOT NULL,
            outro_tipo TEXT,
            total_dias INTEGER DEFAULT 0,
            data_atestado TEXT NOT NULL,
            criado_em TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (turma_id) REFERENCES turmas(id),
            FOREIGN KEY (aluno_id) REFERENCES alunos(id)
        )
    ''')

    # Responsáveis
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS responsaveis (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            login TEXT UNIQUE NOT NULL,
            senha TEXT NOT NULL,
            telefone TEXT,
            aluno_id INTEGER,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id)
        )
    ''')
    # Logs de acessos (professores e responsáveis)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS logs_acessos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tipo TEXT NOT NULL,
            login TEXT NOT NULL,
            data_hora TEXT NOT NULL DEFAULT (datetime('now','localtime'))
        )
    ''')

    # Recados
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS recados (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            titulo TEXT NOT NULL,
            mensagem TEXT NOT NULL,
            data_envio TEXT DEFAULT (datetime('now','localtime')),
            FOREIGN KEY (professor_id) REFERENCES professores(id),
            FOREIGN KEY (turma_id) REFERENCES turmas(id)
        )
    ''')

    # Recados por aluno (mensagens individuais)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS recados_aluno (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER,
            aluno_id INTEGER,
            turma_id INTEGER,
            conteudo TEXT NOT NULL,
            data_criacao TEXT NOT NULL,
            visualizado INTEGER DEFAULT 0
        )
    ''')

    # Planejamentos
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS planejamentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            disciplina TEXT NOT NULL,
            bimestre TEXT NOT NULL,
            ano INTEGER NOT NULL,
            criado_em TEXT DEFAULT (datetime('now','localtime')),
            observacoes TEXT,
            FOREIGN KEY (professor_id) REFERENCES professores(id)
        )
    ''')

    # Itens do planejamento (conteúdos, habilidades etc.)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS planejamento_itens (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            planejamento_id INTEGER NOT NULL,
            descricao_conteudo TEXT NOT NULL,
            habilidades TEXT,
            forma_avaliacao TEXT,
            pontuacao_total REAL DEFAULT 0,
            FOREIGN KEY (planejamento_id) REFERENCES planejamentos(id) ON DELETE CASCADE
        )
    ''')

    # Relação planejamento x turmas
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS planejamentos_turmas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            planejamento_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            FOREIGN KEY (planejamento_id) REFERENCES planejamentos(id) ON DELETE CASCADE,
            FOREIGN KEY (turma_id) REFERENCES turmas(id),
            UNIQUE (planejamento_id, turma_id)
        )
    ''')

    # Provas/avaliações criadas pelos professores
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS provas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            bimestre TEXT NOT NULL,
            ano INTEGER NOT NULL,
            titulo TEXT NOT NULL,
            descricao TEXT,
            criado_em TEXT DEFAULT (datetime('now','localtime')),
            FOREIGN KEY (professor_id) REFERENCES professores(id),
            FOREIGN KEY (turma_id) REFERENCES turmas(id)
        )
    ''')

    # Lançamentos de notas/avaliações
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS provas_lancamentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            prova_id INTEGER NOT NULL,
            aluno_id INTEGER NOT NULL,
            nota REAL,
            observacao TEXT,
            criado_em TEXT DEFAULT (datetime('now','localtime')),
            FOREIGN KEY (prova_id) REFERENCES provas(id) ON DELETE CASCADE,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id)
        )
    ''')

    # NOVO: tabela só para os modelos de recado por aluno
    cursor.execute("""
                   CREATE TABLE IF NOT EXISTS recados_modelos (
                       id INTEGER PRIMARY KEY AUTOINCREMENT,
                       conteudo TEXT NOT NULL
                   )
               """)

    # Verifica se já existem modelos, senão insere os padrões
    cursor.execute("SELECT COUNT(*) AS total FROM recados_modelos")
    total_recados = cursor.fetchone()["total"]

    if total_recados == 0:
        recados_padrao = [
            # 1) FALTAS
            ("[FALTAS] Informamos que o(a) estudante tem apresentado número elevado de faltas às aulas, "
             "o que pode prejudicar seu acompanhamento dos conteúdos e seu desenvolvimento. "
             "Solicitamos que verifiquem os motivos dessas ausências e incentivem a frequência regular."),

            # 2) NÃO REALIZA ATIVIDADES
            ("[ATIVIDADES] Observamos que o(a) estudante não tem realizado, com regularidade, as atividades "
             "propostas em sala e/ou as tarefas de casa. Pedimos que acompanhem a rotina de estudos e reforcem "
             "com o(a) estudante a importância de concluir as atividades."),

            # 3) DESINTERESSE
            ("[DESINTERESSE] Percebemos, nas últimas aulas, um desinteresse do(a) estudante em relação às atividades, "
             "explicações e momentos de participação. Solicitamos o apoio da família para conversar com o(a) aluno(a) "
             "sobre a importância do compromisso com os estudos."),

            # 4) SONOLÊNCIA
            ("[SONOLÊNCIA] Temos observado que o(a) estudante tem se apresentado muito sonolento(a) em sala de aula, "
             "o que compromete sua atenção e aprendizado. Pedimos que, se possível, organizem a rotina de sono em casa "
             "para que ele(a) consiga aproveitar melhor as aulas."),

            # 5) ELOGIO
            (
                "[ELOGIO] Parabenizamos o(a) estudante pelo bom comportamento, participação nas aulas e dedicação às atividades. "
                "Ele(a) tem se destacado positivamente e demonstrado empenho em sua aprendizagem. Agradecemos o apoio da família."),

            # 6) INDISCIPLINA
            (
                "[INDISCIPLINA] Registramos que o(a) estudante tem apresentado comportamentos de indisciplina em sala de aula, "
                "como desatenção às orientações, conversas em excesso e atitudes que atrapalham o andamento da aula. "
                "Solicitamos o apoio da família para conversar com o(a) aluno(a) e reforçar as regras de convivência."),

            # 7) ATRASOS FREQUENTES
            ("[ATRASOS] Informamos que o(a) estudante tem chegado com frequência atrasado(a) para o início das aulas. "
             "Essa situação prejudica o aproveitamento das atividades e pode gerar dificuldades no acompanhamento dos conteúdos. "
             "Solicitamos o apoio da família para que o(a) estudante chegue no horário estabelecido."),

            # 8) USO DE CELULAR
            (
                "[CELULAR] Informamos que o(a) estudante tem utilizado o telefone celular durante as aulas, mesmo após orientações da escola. "
                "Essa prática prejudica sua atenção, o andamento das atividades e a concentração da turma. "
                "Solicitamos o apoio da família para reforçar com o(a) estudante as regras de uso responsável do aparelho no ambiente escolar."),

            # 9) FALTA DE MATERIAL
            (
                "[MATERIAL] Informamos que o(a) estudante tem vindo com frequência para a escola sem caderno, livros ou materiais necessários para as aulas. "
                "Isso dificulta sua participação nas atividades e o acompanhamento dos conteúdos. "
                "Pedimos o apoio da família para organizar os materiais escolares e verificar diariamente se estão sendo levados para a escola."),

            # 10) CONVERSAS EXCESSIVAS
            (
                "[CONVERSAS] Registramos que o(a) estudante tem mantido conversas em excesso durante as explicações e atividades em sala de aula, "
                "o que prejudica sua própria aprendizagem e a concentração da turma. "
                "Solicitamos o apoio da família para conversar com o(a) aluno(a) sobre a importância do foco e do respeito ao momento de estudo."),

            # 11) DIFICULDADE DE APRENDIZAGEM
            (
                "[DIFICULDADE] Observamos que o(a) estudante tem apresentado dificuldades no acompanhamento dos conteúdos trabalhados em sala de aula. "
                "Sugerimos que a família busque formas de apoio pedagógico complementar e mantenha diálogo constante com a escola para acompanhar a evolução do(a) estudante."),

            # 12) BAIXO RENDIMENTO
            (
                "[RENDIMENTO] Informamos que o(a) estudante tem apresentado rendimento abaixo do esperado nas avaliações e atividades. "
                "Recomendamos que a família acompanhe mais de perto a rotina de estudos, verificando a realização das tarefas e incentivando momentos de revisão dos conteúdos."),

            # 13) UNIFORME
            (
                "[UNIFORME] Informamos que o(a) estudante tem frequentado as aulas sem o uniforme escolar completo ou adequado. "
                "O uso do uniforme é importante para a identificação e segurança dos estudantes. Pedimos que a família se certifique de que o(a) estudante venha uniformizado(a) diariamente."),

            # 14) PROGRESSO ACADÊMICO
            (
                "[PROGRESSO] Temos o prazer de informar que o(a) estudante apresentou excelente evolução em seu desempenho acadêmico neste bimestre. "
                "Ele(a) tem demonstrado maior dedicação, participação e compreensão dos conteúdos. Parabenizamos o(a) estudante e a família pelo apoio."),

            # 15) MUDANÇA DE COMPORTAMENTO
            ("[COMPORTAMENTO] Observamos uma mudança no comportamento habitual do(a) estudante nos últimos dias. "
             "Ele(a) tem se mostrado mais retraído(a)/agitado(a) que o normal. Sugerimos uma conversa com a família para verificar se está tudo bem e se há algo que possamos fazer para ajudar."),

            # 16) SAÍDA ANTECIPADA
            ("[SAÍDA] Registramos que o(a) estudante tem saído antecipadamente da escola com frequência. "
             "Isso prejudica o acompanhamento das aulas finais do turno e pode gerar defasagem no aprendizado. Solicitamos que, sempre que possível, o(a) estudante permaneça até o final do horário regular."),

            # 17) RELACIONAMENTO COM COLEGAS
            (
                "[RELACIONAMENTO] Informamos que o(a) estudante tem apresentado dificuldades no relacionamento com alguns colegas, "
                "gerando conflitos durante as aulas e intervalos. Solicitamos o apoio da família para conversar sobre respeito, empatia e boa convivência no ambiente escolar."),

            # 18) USO INADEQUADO DE LINGUAGEM
            (
                "[LINGUAGEM] Registramos que o(a) estudante tem utilizado linguagem inadequada (palavrões, apelidos ofensivos) no ambiente escolar. "
                "Solicitamos que a família reforce com o(a) estudante a importância do respeito na comunicação e do uso de vocabulário apropriado."),

            # 19) EXCESSO DE BRINCADEIRAS
            (
                "[BRINCADEIRAS] Observamos que o(a) estudante tem extrapolado durante as brincadeiras, não respeitando os limites dos colegas "
                "e gerando situações de desconforto ou risco. Pedimos que conversem sobre a importância de brincar com respeito e segurança."),

            # 20) RECUSA EM PARTICIPAR DE ATIVIDADES
            (
                "[RECUSA] Informamos que o(a) estudante tem se recusado a participar de atividades propostas em sala de aula, "
                "como apresentações, trabalhos em grupo ou atividades físicas. Sugerimos que a família dialogue para entender os motivos dessa recusa e incentive a participação."),

            # 21) ORGANIZAÇÃO DO MATERIAL
            (
                "[ORGANIZAÇÃO] Observamos que o(a) estudante tem apresentado dificuldades em manter seus materiais escolares organizados, "
                "perdendo cadernos, livros ou esquecendo itens importantes. Sugerimos que a família auxilie na criação de uma rotina de organização e verificação diária da mochila."),

            # 22) MELHORA NO COMPORTAMENTO
            (
                "[MELHORA] Parabenizamos o(a) estudante pela significativa melhora apresentada em seu comportamento nas últimas semanas. "
                "Ele(a) tem demonstrado mais respeito às regras, melhor convivência com os colegas e maior foco nas atividades. Agradecemos o apoio da família nesse processo."),

            # 23) FALTA DE LIÇÃO DE CASA
            ("[LIÇÃO] Informamos que o(a) estudante tem deixado de entregar as lições de casa com frequência. "
             "A realização dessas atividades é fundamental para a fixação dos conteúdos e o desenvolvimento da autonomia. "
             "Solicitamos que a família acompanhe e incentive a realização diária das tarefas escolares.")
        ]

        cursor.executemany(
            "INSERT INTO recados_modelos (conteudo) VALUES (?)",
            [(r,) for r in recados_padrao]
        )
        conn.commit()

    # AGORA FORA DO IF: garante que os modelos extras existam mesmo em bancos antigos
    modelos_extras = [
        ("ATRASOS",
         "[ATRASOS] Informamos que o(a) estudante tem chegado com frequência atrasado(a) para o início das aulas. "
         "Essa situação prejudica o aproveitamento das atividades e pode gerar dificuldades no acompanhamento dos conteúdos. "
         "Solicitamos o apoio da família para que o(a) estudante chegue no horário estabelecido."),

        ("CELULAR",
         "[CELULAR] Informamos que o(a) estudante tem utilizado o telefone celular durante as aulas, mesmo após orientações da escola. "
         "Essa prática prejudica sua atenção, o andamento das atividades e a concentração da turma. "
         "Solicitamos o apoio da família para reforçar com o(a) estudante as regras de uso responsável do aparelho no ambiente escolar."),

        ("MATERIAL",
         "[MATERIAL] Informamos que o(a) estudante tem vindo com frequência para a escola sem caderno, livros ou materiais necessários para as aulas. "
         "Isso dificulta sua participação nas atividades e o acompanhamento dos conteúdos. "
         "Pedimos o apoio da família para organizar os materiais escolares e verificar diariamente se estão sendo levados para a escola."),

        ("CONVERSAS",
         "[CONVERSAS] Registramos que o(a) estudante tem mantido conversas em excesso durante as explicações e atividades em sala de aula, "
         "o que prejudica sua própria aprendizagem e a concentração da turma. "
         "Solicitamos o apoio da família para conversar com o(a) aluno(a) sobre a importância do foco e do respeito ao momento de estudo."),

        ("DIFICULDADE",
         "[DIFICULDADE] Observamos que o(a) estudante tem apresentado dificuldades no acompanhamento dos conteúdos trabalhados em sala de aula. "
         "Sugerimos que a família busque formas de apoio pedagógico complementar e mantenha diálogo constante com a escola para acompanhar a evolução do(a) estudante."),

        ("RENDIMENTO",
         "[RENDIMENTO] Informamos que o(a) estudante tem apresentado rendimento abaixo do esperado nas avaliações e atividades. "
         "Recomendamos que a família acompanhe mais de perto a rotina de estudos, verificando a realização das tarefas e incentivando momentos de revisão dos conteúdos."),

        ("UNIFORME",
         "[UNIFORME] Informamos que o(a) estudante tem frequentado as aulas sem o uniforme escolar completo ou adequado. "
         "O uso do uniforme é importante para a identificação e segurança dos estudantes. Pedimos que a família se certifique de que o(a) estudante venha uniformizado(a) diariamente."),

        ("PROGRESSO",
         "[PROGRESSO] Temos o prazer de informar que o(a) estudante apresentou excelente evolução em seu desempenho acadêmico neste bimestre. "
         "Ele(a) tem demonstrado maior dedicação, participação e compreensão dos conteúdos. Parabenizamos o(a) estudante e a família pelo apoio."),

        ("COMPORTAMENTO",
         "[COMPORTAMENTO] Observamos uma mudança no comportamento habitual do(a) estudante nos últimos dias. "
         "Ele(a) tem se mostrado mais retraído(a)/agitado(a) que o normal. Sugerimos uma conversa com a família para verificar se está tudo bem e se há algo que possamos fazer para ajudar."),

        ("SAÍDA",
         "[SAÍDA] Registramos que o(a) estudante tem saído antecipadamente da escola com frequência. "
         "Isso prejudica o acompanhamento das aulas finais do turno e pode gerar defasagem no aprendizado. Solicitamos que, sempre que possível, o(a) estudante permaneça até o final do horário regular."),

        ("RELACIONAMENTO",
         "[RELACIONAMENTO] Informamos que o(a) estudante tem apresentado dificuldades no relacionamento com alguns colegas, "
         "gerando conflitos durante as aulas e intervalos. Solicitamos o apoio da família para conversar sobre respeito, empatia e boa convivência no ambiente escolar."),

        ("LINGUAGEM",
         "[LINGUAGEM] Registramos que o(a) estudante tem utilizado linguagem inadequada (palavrões, apelidos ofensivos) no ambiente escolar. "
         "Solicitamos que a família reforce com o(a) estudante a importância do respeito na comunicação e do uso de vocabulário apropriado."),

        ("BRINCADEIRAS",
         "[BRINCADEIRAS] Observamos que o(a) estudante tem extrapolado durante as brincadeiras, não respeitando os limites dos colegas "
         "e gerando situações de desconforto ou risco. Pedimos que conversem sobre a importância de brincar com respeito e segurança."),

        ("RECUSA",
         "[RECUSA] Informamos que o(a) estudante tem se recusado a participar de atividades propostas em sala de aula, "
         "como apresentações, trabalhos em grupo ou atividades físicas. Sugerimos que a família dialogue para entender os motivos dessa recusa e incentive a participação."),

        ("ORGANIZAÇÃO",
         "[ORGANIZAÇÃO] Observamos que o(a) estudante tem apresentado dificuldades em manter seus materiais escolares organizados, "
         "perdendo cadernos, livros ou esquecendo itens importantes. Sugerimos que a família auxilie na criação de uma rotina de organização e verificação diária da mochila."),

        ("MELHORA",
         "[MELHORA] Parabenizamos o(a) estudante pela significativa melhora apresentada em seu comportamento nas últimas semanas. "
         "Ele(a) tem demonstrado mais respeito às regras, melhor convivência com os colegas e maior foco nas atividades. Agradecemos o apoio da família nesse processo."),

        ("LIÇÃO",
         "[LIÇÃO] Informamos que o(a) estudante tem deixado de entregar as lições de casa com frequência. "
         "A realização dessas atividades é fundamental para a fixação dos conteúdos e o desenvolvimento da autonomia. "
         "Solicitamos que a família acompanhe e incentive a realização diária das tarefas escolares.")
    ]

    for rotulo, texto in modelos_extras:
        cursor.execute(
            "SELECT COUNT(*) AS total FROM recados_modelos WHERE conteudo LIKE ?",
            (f'[{rotulo}] %',)
        )
        tem_modelo = cursor.fetchone()["total"]

        if tem_modelo == 0:
            cursor.execute(
                "INSERT INTO recados_modelos (conteudo) VALUES (?)",
                (texto,)
            )
            conn.commit()

    # Disciplinas do professor
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS professor_disciplinas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            disciplina TEXT NOT NULL,
            FOREIGN KEY (professor_id) REFERENCES professores(id)
        )
    ''')

    # Turmas do professor
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS professor_turmas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            FOREIGN KEY (professor_id) REFERENCES professores(id),
            FOREIGN KEY (turma_id) REFERENCES turmas(id),
            UNIQUE (professor_id, turma_id)
        )
    ''')

    # Planejamentos
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS planejamentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            disciplina TEXT NOT NULL,
            bimestre INTEGER NOT NULL,
            ano INTEGER NOT NULL,
            criado_em TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (professor_id) REFERENCES professores(id)
        )
    ''')

    # Relação planejamento x turmas
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS planejamentos_turmas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            planejamento_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            FOREIGN KEY (planejamento_id) REFERENCES planejamentos(id),
            FOREIGN KEY (turma_id) REFERENCES turmas(id),
            UNIQUE (planejamento_id, turma_id)
        )
    ''')

    # Itens do planejamento
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS planejamento_itens (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            planejamento_id INTEGER NOT NULL,
            conteudo TEXT NOT NULL,
            data_inicio TEXT,
            data_fim TEXT,
            forma_avaliacao TEXT,
            pontuacao_total REAL,
            FOREIGN KEY (planejamento_id) REFERENCES planejamentos(id)
        )
    ''')

    # Conteúdos das avaliações bimestrais
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS avaliacoes_bimestrais (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            disciplina TEXT NOT NULL,
            turma_id INTEGER NOT NULL,
            bimestre INTEGER NOT NULL,
            ano INTEGER NOT NULL,
            tipo_avaliacao TEXT,
            descricao_avaliacao TEXT,
            conteudos TEXT,
            data_avaliacao TEXT,
            pontuacao REAL,
            criado_em TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (professor_id) REFERENCES professores(id),
            FOREIGN KEY (turma_id) REFERENCES turmas(id)
        )
    ''')

    # Recados por aluno (vinculado a professor, aluno e turma)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS recados_aluno (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_id INTEGER NOT NULL,
            aluno_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            conteudo TEXT NOT NULL,
            data_criacao TEXT DEFAULT CURRENT_TIMESTAMP,
            visualizado INTEGER DEFAULT 0,
            FOREIGN KEY (professor_id) REFERENCES professores(id),
            FOREIGN KEY (aluno_id) REFERENCES alunos(id),
            FOREIGN KEY (turma_id) REFERENCES turmas(id)
        )
    ''')
    # --- MIGRAÇÃO: garante colunas esperadas em recados_aluno ---
    try:
        cursor.execute("PRAGMA table_info(recados_aluno)")
        cols = [row[1] for row in cursor.fetchall()]  # (cid, name, type, notnull, dflt_value, pk)

        if 'visualizado' not in cols:
            cursor.execute("ALTER TABLE recados_aluno ADD COLUMN visualizado INTEGER DEFAULT 0")

        if 'conteudo' not in cols:
            cursor.execute("ALTER TABLE recados_aluno ADD COLUMN conteudo TEXT")

        # Se existir recado_id e ainda não tiver conteúdo salvo, tenta preencher pelo modelo
        if 'recado_id' in cols:
            cursor.execute(
                """
                UPDATE recados_aluno
                SET conteudo = (
                    SELECT rm.conteudo
                    FROM recados_modelos rm
                    WHERE rm.id = recados_aluno.recado_id
                )
                WHERE (conteudo IS NULL OR TRIM(conteudo) = '')
                """
            )
    except Exception:
        # Não interrompe a inicialização por falhas de migração
        pass

    # Bibliotecários (login próprio, aprovado pelo moderador)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS bibliotecarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            login TEXT UNIQUE NOT NULL,
            senha TEXT NOT NULL,
            nome TEXT NOT NULL,
            telefone TEXT,
            status TEXT NOT NULL DEFAULT 'pendente'
        )
    ''')

    # Empréstimos da Biblioteca
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS emprestimos_biblioteca (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            aluno_id INTEGER NOT NULL,
            turma_id INTEGER NOT NULL,
            titulo_livro TEXT NOT NULL,
            autor TEXT,
            codigo_interno TEXT,
            data_emprestimo TEXT NOT NULL,
            data_prevista_devolucao TEXT,
            data_devolucao TEXT,
            status TEXT NOT NULL DEFAULT 'Emprestado',
            devolucao_pontual INTEGER DEFAULT NULL,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id),
            FOREIGN KEY (turma_id) REFERENCES turmas(id)
        )
    ''')

    # SEED: moderadores padrão
    from werkzeug.security import generate_password_hash

    senha_hash_reginaldo = generate_password_hash('reginaldodiretor123')
    cursor.execute(
        "INSERT OR IGNORE INTO moderadores (login, senha, tipo) VALUES (?, ?, ?)",
        ('REGINALDO', senha_hash_reginaldo, 'diretor')
    )

    senha_hash_savio = generate_password_hash('reginaldodiretor123')
    cursor.execute(
        "INSERT OR IGNORE INTO moderadores (login, senha, tipo) VALUES (?, ?, ?)",
        ('SAVIO', senha_hash_savio, 'diretor')
    )

    # ----------------- ATENDIMENTOS A RESPONSÁVEIS (MODERADOR) ----------------- #
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS atendimentos_responsaveis (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            protocolo TEXT UNIQUE,
            turno TEXT,
            turma_id INTEGER NOT NULL,
            aluno_id INTEGER NOT NULL,

            -- Informações do responsável (digitado no atendimento)
            responsavel_nome TEXT,
            responsavel_parentesco TEXT,

            -- Mantido por compatibilidade (caso você queira voltar a vincular)
            responsavel_id INTEGER,

            -- Quem registrou o atendimento
            registrador_nome TEXT,
            registrador_cargo TEXT,

            envolve_professor INTEGER DEFAULT 0,
            professor_nome TEXT,

            data_atendimento TEXT NOT NULL,
            hora_atendimento TEXT,
            assunto TEXT,

            relato TEXT NOT NULL,
            combinados TEXT,

            retorno_previsto INTEGER DEFAULT 0,
            retorno_em TEXT,

            reuniao_agendada INTEGER DEFAULT 0,
            reuniao_data TEXT,

            criado_por_login TEXT,
            criado_em TEXT NOT NULL DEFAULT (datetime('now','localtime')),

            FOREIGN KEY (turma_id) REFERENCES turmas(id),
            FOREIGN KEY (aluno_id) REFERENCES alunos(id),
            FOREIGN KEY (responsavel_id) REFERENCES responsaveis(id)
        )
    ''')


    # ----------------- SALA DE RECURSOS (AEE) ----------------- #
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sala_recursos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            aluno_id INTEGER NOT NULL UNIQUE,
            turma_id INTEGER NOT NULL,
            turno TEXT,
            cadastrado_em TEXT NOT NULL DEFAULT (datetime('now','localtime')),
            cadastrado_por TEXT,
            FOREIGN KEY (aluno_id) REFERENCES alunos(id) ON DELETE CASCADE,
            FOREIGN KEY (turma_id) REFERENCES turmas(id) ON DELETE CASCADE
        )
    ''')
    conn.commit()
    cursor.close()
    conn.close()


def atualizar_bd():
    conn = conectar_bd()
    cursor = conn.cursor()

    def colunas_da_tabela(tabela: str):
        cursor.execute(f"PRAGMA table_info({tabela})")
        return [col[1] for col in cursor.fetchall()]

    def adicionar_coluna(tabela, coluna, definicao):
        try:
            cols = colunas_da_tabela(tabela)
            if coluna not in cols:
                cursor.execute(f"ALTER TABLE {tabela} ADD COLUMN {coluna} {definicao}")
                conn.commit()
        except sqlite3.Error:
            pass

    # ---------------- MODERADORES ----------------
    try:
        adicionar_coluna('moderadores', 'tipo', "TEXT DEFAULT 'diretor'")
        adicionar_coluna('moderadores', 'soe_liberado', "INTEGER NOT NULL DEFAULT 0")
        try:
            cursor.execute(
                "UPDATE moderadores SET tipo='diretor' "
                "WHERE tipo IS NULL OR tipo=''"
            )
            conn.commit()
        except sqlite3.Error:
            pass
    except sqlite3.Error:
        pass

    # ---------------- PLANEJAMENTOS ----------------
    # Observações no cabeçalho do planejamento
    try:
        adicionar_coluna('planejamentos', 'observacoes', "TEXT")
    except sqlite3.Error:
        pass

    # Itens do planejamento: garante compatibilidade com banco antigo e novo
    try:
        # colunas essenciais do seu HTML novo (datas_inicio/datas_fim + conteudo)
        adicionar_coluna('planejamento_itens', 'conteudo', "TEXT")
        adicionar_coluna('planejamento_itens', 'data_inicio', "TEXT")
        adicionar_coluna('planejamento_itens', 'data_fim', "TEXT")

        # usadas no seu app (listagem soma pontuacao_total) + compatibilidade
        adicionar_coluna('planejamento_itens', 'forma_avaliacao', "TEXT")
        adicionar_coluna('planejamento_itens', 'pontuacao_total', "REAL")

        # concluído (checkbox)
        adicionar_coluna('planejamento_itens', 'concluido', "INTEGER DEFAULT 0")

        # MIGRAÇÃO: se você tinha uma coluna antiga "descricao_conteudo",
        # copiamos para "conteudo" quando "conteudo" estiver vazio.
        cols = colunas_da_tabela('planejamento_itens')
        if 'descricao_conteudo' in cols:
            try:
                cursor.execute("""
                    UPDATE planejamento_itens
                    SET conteudo = COALESCE(NULLIF(conteudo, ''), descricao_conteudo)
                    WHERE conteudo IS NULL OR TRIM(conteudo) = ''
                """)
                conn.commit()
            except sqlite3.Error:
                pass

    except sqlite3.Error:
        pass

    # ---------------- ATENDIMENTOS ----------------
    try:
        adicionar_coluna('atendimentos_responsaveis', 'responsavel_nome', "TEXT")
        adicionar_coluna('atendimentos_responsaveis', 'responsavel_parentesco', "TEXT")
        adicionar_coluna('atendimentos_responsaveis', 'registrador_nome', "TEXT")
        adicionar_coluna('atendimentos_responsaveis', 'registrador_cargo', "TEXT")
    except sqlite3.Error:
        pass

    # ---------------- RESPONSÁVEIS ----------------
    try:
        adicionar_coluna('responsaveis', 'telefone', "TEXT")
    except sqlite3.Error:
        pass

    try:
        adicionar_coluna('responsaveis', 'principal', "INTEGER DEFAULT 0")
    except sqlite3.Error:
        pass

    
    # ---------------- SALA DE RECURSOS (AEE) ----------------
    try:
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS sala_recursos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                aluno_id INTEGER NOT NULL UNIQUE,
                turma_id INTEGER NOT NULL,
                turno TEXT,
                cadastrado_em TEXT NOT NULL DEFAULT (datetime('now','localtime')),
                cadastrado_por TEXT,
                FOREIGN KEY (aluno_id) REFERENCES alunos(id) ON DELETE CASCADE,
                FOREIGN KEY (turma_id) REFERENCES turmas(id) ON DELETE CASCADE
            )
        ''')
        conn.commit()

        # Garante colunas em bancos antigos
        adicionar_coluna('sala_recursos', 'turno', "TEXT")
        adicionar_coluna('sala_recursos', 'cadastrado_em', "TEXT NOT NULL DEFAULT (datetime('now','localtime'))")
        adicionar_coluna('sala_recursos', 'cadastrado_por', "TEXT")
    except sqlite3.Error:
        pass

# ---------------- RECADOS ----------------
    try:
        adicionar_coluna('recados_aluno', 'visualizado', "INTEGER DEFAULT 0")
    except sqlite3.Error:
        pass

    cursor.close()
    conn.close()



# Funções auxiliares de usuário/professor
def obter_professor_id(login):
    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM professores WHERE login = ?", (login,))
    row = cursor.fetchone()
    cursor.close()
    conn.close()
    return row['id'] if row else None


def obter_turmas_professor(professor_id):
    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute('''
        SELECT t.id, t.nome, t.turno
        FROM turmas t
        JOIN professor_turmas pt ON pt.turma_id = t.id
        WHERE pt.professor_id = ?
        ORDER BY t.turno, t.nome
    ''', (professor_id,))
    turmas = cursor.fetchall()
    cursor.close()
    conn.close()
    return turmas


# Inicializar banco e ajustes
inicializar_bd()
atualizar_bd()
ensure_soe_table()
ensure_termo_tables(conectar_bd)
ensure_checklist_tables()

bp_termo.conectar_bd = conectar_bd
bp_rotina.conectar_bd = conectar_bd
bp_checklist.conectar_bd = conectar_bd
# Rotas da Biblioteca Escolar
app.register_blueprint(bp_biblioteca, url_prefix='/biblioteca')

# Carômetro (depende de professores/turmas/alunos existirem)
init_carometro_db()
app.register_blueprint(bp_carometro)
app.register_blueprint(bp_soe)
app.register_blueprint(bp_conselho)
app.register_blueprint(bp_termo)
app.register_blueprint(bp_checklist)

# Sistema de Rotinas

app.register_blueprint(bp_rotina)

try:
    ensure_conselho_tables()
except Exception as _e:
    # Evita quebrar a inicializacao se o banco ainda nao estiver pronto
    print('[CONSELHO] Falha ao garantir tabelas:', _e)

# Inicializar tabelas de rotinas
try:
    ensure_rotina_tables()
except Exception as e:
    print('[ROTINA] Falha ao garantir tabelas:', e)


# Rotas principais

@app.route('/')
def index():
    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT conteudo FROM recados")
    recados = cursor.fetchall()
    cursor.close()
    conn.close()
    return render_template('index.html', recados=recados)


# Login
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        usuario = request.form['usuario']
        senha = request.form['senha']

        conn = conectar_bd()
        cursor = conn.cursor()

        # Moderador
        cursor.execute("SELECT senha, tipo FROM moderadores WHERE login = ?", (usuario,))
        moderador = cursor.fetchone()

        if moderador and check_password_hash(moderador['senha'], senha):
            session['usuario'] = usuario
            session['tipo'] = 'moderador'
            cursor.close()
            conn.close()
            return redirect(url_for('dashboard_moderador'))

        # Professor
        cursor.execute("SELECT senha, status FROM professores WHERE login = ?", (usuario,))
        professor = cursor.fetchone()

        cursor.close()
        conn.close()

        if professor and check_password_hash(professor['senha'], senha):
            if professor['status'] == 'pendente':
                flash("Cadastro pendente de aprovação.")
                return redirect(url_for('login'))

            session['usuario'] = usuario
            session['tipo'] = 'professor'

            # Registrar log de acesso do professor
            try:
                conn_log = conectar_bd()
                cursor_log = conn_log.cursor()
                cursor_log.execute(
                    "INSERT INTO logs_acessos (tipo, login) VALUES (?, ?)",
                    ('professor', usuario)
                )
                conn_log.commit()
            except sqlite3.Error:
                # Não interrompe o fluxo em caso de erro ao registrar o log
                pass
            finally:
                try:
                    cursor_log.close()
                except Exception:
                    pass
                try:
                    conn_log.close()
                except Exception:
                    pass

            return redirect(url_for('dashboard_professor'))

        flash("Usuário ou senha inválidos.")
        return redirect(url_for('login'))

    return render_template('login.html')


# Moderadores

@app.route('/cadastrar_moderador', methods=['GET', 'POST'])
def cadastrar_moderador():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    if request.method == 'POST':
        login_m = request.form['login']
        senha = request.form['senha']
        tipo = request.form.get('tipo', 'coordenador')
        senha_hash = generate_password_hash(senha)

        conn = conectar_bd()
        cursor = conn.cursor()

        try:
            cursor.execute(
                "INSERT INTO moderadores (login, senha, tipo) VALUES (?, ?, ?)",
                (login_m, senha_hash, tipo)
            )
            conn.commit()
            flash("Moderador cadastrado com sucesso!")
        except sqlite3.IntegrityError:
            flash("Login já está em uso. Escolha outro.")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('dashboard_moderador'))

    return render_template('cadastrar_moderador.html')


@app.route('/visualizar_moderadores')
def visualizar_moderadores():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id, login, tipo FROM moderadores")
    moderadores = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template('visualizar_moderadores.html', moderadores=moderadores)


@app.route('/excluir_moderador/<int:moderador_id>', methods=['POST'])
def excluir_moderador(moderador_id):
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        cursor.execute("DELETE FROM moderadores WHERE id = ?", (moderador_id,))
        conn.commit()
        flash("Moderador excluído com sucesso!")
    except sqlite3.Error as e:
        flash(f"Erro ao excluir moderador: {e}")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('visualizar_moderadores'))


# ------------------------------------------------------
# CONTROLE DE ACESSO AO SOE (somente SAVIO libera/revoga)
# ------------------------------------------------------
@app.route('/soe/controle_acesso', methods=['GET', 'POST'])
def soe_controle_acesso():
    # precisa estar logado como moderador
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    # apenas o SAVIO pode gerenciar a liberação do SOE
    if (session.get('usuario') or '').upper() != 'SAVIO':
        flash("Apenas o moderador SAVIO pode gerenciar o acesso ao SOE.")
        return redirect(url_for('dashboard_moderador'))

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # garante que a coluna soe_liberado existe (caso atualizar_bd ainda não tenha rodado)
    try:
        cursor.execute("PRAGMA table_info(moderadores)")
        cols = [r['name'] for r in cursor.fetchall()]
        if 'soe_liberado' not in cols:
            cursor.execute("ALTER TABLE moderadores ADD COLUMN soe_liberado INTEGER NOT NULL DEFAULT 0")
            conn.commit()
    except Exception:
        pass

    if request.method == 'POST':
        senha_savio = (request.form.get('senha_savio') or '').strip()
        if not _senha_savio_ok(conn, senha_savio):
            cursor.close()
            conn.close()
            flash("Senha do moderador SAVIO incorreta. Ação não permitida.")
            return redirect(url_for('soe_controle_acesso'))

        moderador_id = (request.form.get('moderador_id') or '').strip()
        acao = (request.form.get('acao') or '').strip()  # 'liberar' ou 'revogar'

        if not moderador_id.isdigit():
            flash("Selecione um moderador válido.")
            cursor.close()
            conn.close()
            return redirect(url_for('soe_controle_acesso'))

        cursor.execute("SELECT id, login FROM moderadores WHERE id = ?", (moderador_id,))
        alvo = cursor.fetchone()

        if not alvo:
            flash("Moderador não encontrado.")
        elif (alvo['login'] or '').upper() == 'SAVIO':
            flash("O SAVIO já possui acesso total permanente ao SOE.")
        else:
            novo = 1 if acao == 'liberar' else 0
            cursor.execute("UPDATE moderadores SET soe_liberado = ? WHERE id = ?", (novo, moderador_id))
            conn.commit()
            flash("Permissão do SOE atualizada com sucesso.")

        cursor.close()
        conn.close()
        return redirect(url_for('soe_controle_acesso'))

    cursor.execute("SELECT id, login, tipo, COALESCE(soe_liberado, 0) AS soe_liberado FROM moderadores ORDER BY login")
    moderadores = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template('soe_controle_acesso.html', moderadores=moderadores)


@app.route('/logs_acessos')
def logs_acessos():
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    # Filtros recebidos via query string
    tipo = (request.args.get('tipo') or 'todos').lower()

    # Tipos válidos (pode acrescentar outros se passar a registrar novos)
    tipos_validos = ('todos', 'professor', 'responsavel', 'biblioteca')
    if tipo not in tipos_validos:
        tipo = 'todos'

    login_filtro = (request.args.get('login') or '').strip()
    data_de = (request.args.get('data_de') or '').strip()
    data_ate = (request.args.get('data_ate') or '').strip()

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # Começa sem filtro de tipo (para "todos")
    sql = '''
        SELECT id, tipo, login, data_hora
        FROM logs_acessos
        WHERE 1=1
    '''
    params = []

    # Se escolheu um tipo específico, filtra
    if tipo != 'todos':
        sql += " AND tipo = ?"
        params.append(tipo)

    if login_filtro:
        sql += " AND login LIKE ?"
        params.append(f"%{login_filtro}%")

    if data_de:
        sql += " AND date(data_hora) >= date(?)"
        params.append(data_de)

    if data_ate:
        sql += " AND date(data_hora) <= date(?)"
        params.append(data_ate)

    # Agora sem LIMIT, para realmente listar todos (se quiser, pode colocar LIMIT 1000)
    sql += " ORDER BY datetime(data_hora) DESC"

    cursor.execute(sql, params)
    logs = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'logs_acessos.html',
        logs=logs,
        tipo=tipo,
        login_filtro=login_filtro,
        data_de=data_de,
        data_ate=data_ate
    )


# Professores

@app.route('/cadastro_professor', methods=['GET', 'POST'])
def cadastro_professor():
    if request.method == 'POST':
        login_p = (request.form.get('login') or '').strip()
        senha = request.form.get('senha') or ''

        # ✅ TERMO (obrigatório)
        termo = get_termo_ativo(conectar_bd, "professor")
        aceite = request.form.get("aceite_termo")  # vem "on" quando marcado

        # Validações básicas
        if not login_p or not senha:
            flash("Preencha login e senha.")
            return redirect(url_for('cadastro_professor'))

        # Se não existir termo ativo (evita erro e avisa)
        if not termo:
            flash("Termo de uso não encontrado/ativo. Contate o moderador.")
            return redirect(url_for('cadastro_professor'))

        # Obrigatório aceitar
        if not aceite:
            flash("Para concluir o cadastro, você precisa ler e aceitar o Termo de Uso.")
            return redirect(url_for('cadastro_professor'))

        senha_hash = generate_password_hash(senha)

        # ✅ Apenas 1 disciplina + turmas dela
        disciplina1 = (request.form.get('disciplina1') or '').strip()
        turmas1 = request.form.getlist('turmas1[]')

        if not disciplina1:
            flash("Selecione a disciplina.")
            return redirect(url_for('cadastro_professor'))

        if not turmas1:
            flash("Selecione pelo menos uma turma.")
            return redirect(url_for('cadastro_professor'))

        conn = conectar_bd()
        cursor = conn.cursor()

        try:
            cursor.execute(
                "INSERT INTO professores (login, senha, status) VALUES (?, ?, 'pendente')",
                (login_p, senha_hash)
            )
            professor_id = cursor.lastrowid

            # Salva apenas 1 disciplina
            cursor.execute(
                "INSERT INTO professor_disciplinas (professor_id, disciplina) VALUES (?, ?)",
                (professor_id, disciplina1)
            )

            # Salva vínculo: professor + disciplina + turma
            for turma_id in turmas1:
                turma_id = (turma_id or '').strip()
                if not turma_id:
                    continue

                cursor.execute(
                    """
                    INSERT OR IGNORE INTO professor_turmas_disciplina
                    (professor_id, disciplina, turma_id)
                    VALUES (?, ?, ?)
                    """,
                    (professor_id, disciplina1, int(turma_id))
                )

                # Mantém também professor_turmas por compatibilidade
                cursor.execute(
                    "INSERT OR IGNORE INTO professor_turmas (professor_id, turma_id) VALUES (?, ?)",
                    (professor_id, int(turma_id))
                )

            conn.commit()

            # ✅ Registra o aceite do termo (fora da transação do cadastro, mas seguro)
            # Se der erro aqui, não desfaz o cadastro — apenas avisa.
            try:
                registrar_aceite(conectar_bd, termo, "professor", login_p)
            except Exception as e_aceite:
                flash(f"Atenção: cadastro feito, mas não foi possível registrar o aceite do termo. ({e_aceite})")

            flash("Cadastro realizado com sucesso. Aguarde aprovação pelo moderador.")

        except sqlite3.IntegrityError:
            conn.rollback()
            flash("Usuário já cadastrado. Escolha outro login.")
        except Exception as e:
            conn.rollback()
            flash(f"Erro ao cadastrar professor: {e}")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('index'))

    # GET
    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()
    cursor.close()
    conn.close()

    termo = get_termo_ativo(conectar_bd, "professor")
    return render_template('cadastro_professor.html', turmas=turmas, disciplinas=DISCIPLINAS, termo=termo)


@app.route('/visualizar_professores', methods=['GET'])
def visualizar_professores():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    cursor.execute("""
        SELECT
            p.id,
            p.login,
            GROUP_CONCAT(pd.disciplina, ', ') AS disciplinas
        FROM professores p
        LEFT JOIN professor_disciplinas pd
            ON pd.professor_id = p.id
        WHERE p.status = 'aprovado'
        GROUP BY p.id, p.login
        ORDER BY p.login
    """)

    professores = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template(
        'visualizar_professores.html',
        professores=professores
    )


@app.route('/dashboard_moderador')
def dashboard_moderador():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    cursor.execute("SELECT id, login FROM professores WHERE status = 'pendente'")
    professores_pendentes = cursor.fetchall()

    cursor.execute("SELECT id, login, status FROM professores WHERE status IN ('aprovado', 'rejeitado')")
    professores_aprovados = cursor.fetchall()

    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()

    cursor.execute("""
        SELECT alunos.id, alunos.nome, turmas.nome AS turma_nome
        FROM alunos
        JOIN turmas ON alunos.turma_id = turmas.id
        ORDER BY turmas.nome, alunos.nome
    """)
    alunos = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'dashboard_moderador.html',
        professores_pendentes=professores_pendentes,
        professores_aprovados=professores_aprovados,
        turmas=turmas,
        alunos=alunos
    )


# =====================
# TRANSFERÊNCIA DE ALUNO (MODERADOR)
# =====================

@app.route('/api/alunos_por_turma/<int:turma_id>')
def api_alunos_por_turma(turma_id):
    # Apenas moderador
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        return jsonify({"ok": False, "error": "Acesso não autorizado"}), 403

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome", (turma_id,))
    rows = cursor.fetchall()
    cursor.close()
    conn.close()

    return jsonify([{"id": r["id"], "nome": r["nome"]} for r in rows])


# =====================
# SALA DE RECURSOS (AEE)
# =====================

@app.route('/api/turmas_por_turno/<turno>')
def api_turmas_por_turno(turno):
    # Apenas moderador (para evitar expor a lista completa)
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        return jsonify({"ok": False, "error": "Acesso não autorizado"}), 403

    turno = (turno or '').strip()
    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id, nome FROM turmas WHERE turno = ? ORDER BY nome", (turno,))
    rows = cursor.fetchall()
    cursor.close()
    conn.close()

    return jsonify([{"id": r["id"], "nome": r["nome"]} for r in rows])


@app.route('/sala_recursos', methods=['GET', 'POST'])
def sala_recursos():
    # Moderador (gestão/cadastro)
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # Turnos disponíveis
    cursor.execute("SELECT DISTINCT turno FROM turmas ORDER BY turno")
    turnos = [r["turno"] for r in cursor.fetchall()]

    if request.method == 'POST':
        turno = (request.form.get('turno') or '').strip()
        turma_id = request.form.get('turma_id', type=int)
        alunos_ids = request.form.getlist('alunos_ids')

        if not turno or not turma_id:
            flash("Selecione o turno e a turma.")
            cursor.close()
            conn.close()
            return redirect(url_for('sala_recursos'))

        # Normaliza ids
        alunos_ids_norm = []
        for a in alunos_ids:
            a = (a or '').strip()
            if a.isdigit():
                alunos_ids_norm.append(int(a))

        # Atualiza a turma: remove e insere novamente
        try:
            cursor.execute("DELETE FROM sala_recursos WHERE turma_id = ?", (turma_id,))
            for aluno_id in alunos_ids_norm:
                cursor.execute(
                    "INSERT OR IGNORE INTO sala_recursos (aluno_id, turma_id, turno, cadastrado_por) VALUES (?, ?, ?, ?)",
                    (aluno_id, turma_id, turno, session.get('usuario'))
                )
            conn.commit()
            flash("Sala de Recursos atualizada com sucesso!")
        except sqlite3.Error as e:
            conn.rollback()
            flash(f"Erro ao salvar Sala de Recursos: {e}")

        cursor.close()
        conn.close()
        return redirect(url_for('sala_recursos', turno=turno, turma_id=turma_id))

    # GET (com filtros)
    turno_sel = (request.args.get('turno') or (turnos[0] if turnos else '')).strip()
    turma_sel_id = request.args.get('turma_id', type=int)

    # Turmas do turno selecionado (para pré-carregar no servidor; o JS atualiza também)
    turmas_turno = []
    if turno_sel:
        cursor.execute("SELECT id, nome FROM turmas WHERE turno = ? ORDER BY nome", (turno_sel,))
        turmas_turno = cursor.fetchall()

    # Se não veio turma_id, pega a primeira do turno
    if not turma_sel_id and turmas_turno:
        turma_sel_id = turmas_turno[0]["id"]

    alunos_turma = []
    alunos_marcados = set()
    alunos_sala = []

    if turma_sel_id:
        # Lista completa da turma
        cursor.execute("SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome", (turma_sel_id,))
        alunos_turma = cursor.fetchall()

        # Marcados na sala de recursos
        cursor.execute("SELECT aluno_id FROM sala_recursos WHERE turma_id = ?", (turma_sel_id,))
        alunos_marcados = {r["aluno_id"] for r in cursor.fetchall()}

        # Lista apenas os alunos da sala (para visualização)
        cursor.execute('''
            SELECT a.id, a.nome
            FROM sala_recursos sr
            JOIN alunos a ON a.id = sr.aluno_id
            WHERE sr.turma_id = ?
            ORDER BY a.nome
        ''', (turma_sel_id,))
        alunos_sala = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'sala_recursos.html',
        turnos=turnos,
        turno_sel=turno_sel,
        turmas_turno=turmas_turno,
        turma_sel_id=turma_sel_id,
        alunos_turma=alunos_turma,
        alunos_marcados=alunos_marcados,
        alunos_sala=alunos_sala
    )


@app.route('/sala_recursos_visualizar')
def sala_recursos_visualizar():
    # Professor (visualização)
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_login = session.get('usuario')
    professor_id = obter_professor_id(professor_login)

    turmas_prof = obter_turmas_professor(professor_id) if professor_id else []

    turma_sel_id = request.args.get('turma_id', type=int)
    if not turma_sel_id and turmas_prof:
        turma_sel_id = turmas_prof[0]["id"]

    alunos_sala = []
    turma_nome = None

    if turma_sel_id:
        conn = conectar_bd()
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()

        cursor.execute("SELECT nome FROM turmas WHERE id = ?", (turma_sel_id,))
        trow = cursor.fetchone()
        turma_nome = trow["nome"] if trow else None

        cursor.execute('''
            SELECT a.id, a.nome
            FROM sala_recursos sr
            JOIN alunos a ON a.id = sr.aluno_id
            WHERE sr.turma_id = ?
            ORDER BY a.nome
        ''', (turma_sel_id,))
        alunos_sala = cursor.fetchall()

        cursor.close()
        conn.close()

    return render_template(
        'sala_recursos_professor.html',
        turmas=turmas_prof,
        turma_sel_id=turma_sel_id,
        turma_nome=turma_nome,
        alunos_sala=alunos_sala
    )


@app.route('/transferir_aluno', methods=['GET', 'POST'])
def transferir_aluno():
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    # Turmas (origem/destino)
    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()

    if request.method == 'GET':
        cursor.close()
        conn.close()
        return render_template('transferir_aluno.html', turmas=turmas)

    turma_origem_id = request.form.get('turma_origem_id', type=int)
    aluno_id = request.form.get('aluno_id', type=int)
    turma_destino_id = request.form.get('turma_destino_id', type=int)
    motivo = (request.form.get('motivo') or '').strip()

    # Validações
    if not turma_origem_id or not aluno_id or not turma_destino_id:
        cursor.close()
        conn.close()
        flash("Preencha turma de origem, aluno e turma de destino.")
        return redirect(url_for('transferir_aluno'))

    if turma_origem_id == turma_destino_id:
        cursor.close()
        conn.close()
        flash("A turma de destino não pode ser a mesma turma de origem.")
        return redirect(url_for('transferir_aluno'))

    # Confere aluno
    cursor.execute("SELECT id, nome, turma_id FROM alunos WHERE id = ?", (aluno_id,))
    aluno = cursor.fetchone()
    if not aluno:
        cursor.close()
        conn.close()
        flash("Aluno não encontrado.")
        return redirect(url_for('transferir_aluno'))

    if int(aluno["turma_id"]) != int(turma_origem_id):
        cursor.close()
        conn.close()
        flash("Este aluno não pertence à turma de origem selecionada.")
        return redirect(url_for('transferir_aluno'))

    # Atualiza turma atual do aluno (mantém todo histórico em outras tabelas, pois tudo é por aluno_id)
    cursor.execute("UPDATE alunos SET turma_id = ? WHERE id = ?", (turma_destino_id, aluno_id))

    # Registra a transferência (log)
    cursor.execute(
        "INSERT INTO transferencias_alunos (aluno_id, turma_origem_id, turma_destino_id, motivo, transferido_por) "
        "VALUES (?, ?, ?, ?, ?)",
        (aluno_id, turma_origem_id, turma_destino_id, motivo, session.get('usuario'))
    )

    conn.commit()
    cursor.close()
    conn.close()

    flash(f"Transferência realizada: {aluno['nome']} agora está na nova turma.")
    return redirect(url_for('dashboard_moderador'))


@app.route('/aprovar_professor/<int:professor_id>', methods=['POST'])
def aprovar_professor(professor_id):
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("UPDATE professores SET status = 'aprovado' WHERE id = ?", (professor_id,))
    conn.commit()
    cursor.close()
    conn.close()

    flash("Professor aprovado com sucesso.")
    return redirect(url_for('dashboard_moderador'))


@app.route('/rejeitar_professor/<int:professor_id>', methods=['POST'])
def rejeitar_professor(professor_id):
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("UPDATE professores SET status = 'rejeitado' WHERE id = ?", (professor_id,))
    conn.commit()
    cursor.close()
    conn.close()

    flash("Professor rejeitado com sucesso.")
    return redirect(url_for('dashboard_moderador'))


@app.route('/excluir_professor/<int:professor_id>', methods=['POST'])
def excluir_professor(professor_id):
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM professores WHERE id = ?", (professor_id,))
    conn.commit()
    cursor.close()
    conn.close()

    flash("Professor excluído com sucesso.")
    return redirect(url_for('dashboard_moderador'))


# ------------------------------------------------------
# EXCLUSÃO AVANÇADA + RESET (exige senha do moderador SAVIO)
# ------------------------------------------------------

def _senha_savio_ok(conn, senha_digitada: str) -> bool:
    """
    Valida a senha digitada comparando com o hash do moderador 'SAVIO'
    na tabela moderadores.
    """
    if not senha_digitada:
        return False

    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT senha FROM moderadores WHERE login = 'SAVIO' LIMIT 1")
        savio = cursor.fetchone()
        if not savio:
            return False
        return check_password_hash(savio["senha"], senha_digitada)
    finally:
        cursor.close()


@app.route('/exclusao', methods=['POST', 'GET'])
def exclusao():
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    if request.method == 'POST':
        senha_savio = (request.form.get('senha_savio') or "").strip()
        if not _senha_savio_ok(conn, senha_savio):
            cursor.close()
            conn.close()
            flash("Senha do moderador SAVIO incorreta. Exclusão não permitida.")
            return redirect(url_for('exclusao'))

        categoria = request.form.get('categoria')
        item_id = request.form.get('item_id')

        if not categoria or not item_id:
            cursor.close()
            conn.close()
            flash("Por favor, selecione uma categoria e um item.")
            return redirect(url_for('exclusao'))

        try:
            tabela = {
                'moderador': 'moderadores',
                'professor': 'professores',
                'aluno': 'alunos',
                'responsavel': 'responsaveis',
                'turma': 'turmas',
                'ocorrencia': 'ocorrencias',
            }.get(categoria)

            if tabela:
                cursor.execute(f"DELETE FROM {tabela} WHERE id = ?", (item_id,))
                conn.commit()
                flash(f"Item da categoria '{categoria}' excluído com sucesso!")
            else:
                flash("Categoria inválida.")
        except sqlite3.Error as e:
            conn.rollback()
            flash(f"Erro ao excluir item: {e}")
        finally:
            cursor.close()
            conn.close()

        # ✅ conforme você pediu: ao excluir, volta para dashboard do moderador
        return redirect(url_for('dashboard_moderador'))

    def safe_fetch(query):
        try:
            cursor.execute(query)
            return [dict(row) for row in cursor.fetchall()]
        except Exception:
            return []

    moderadores = safe_fetch("SELECT * FROM moderadores")
    professores = safe_fetch("SELECT * FROM professores")
    turmas = safe_fetch("SELECT * FROM turmas")
    alunos = safe_fetch("SELECT * FROM alunos")
    responsaveis = safe_fetch("SELECT * FROM responsaveis")
    ocorrencias = safe_fetch("SELECT * FROM ocorrencias")

    cursor.close()
    conn.close()

    return render_template(
        'exclusao.html',
        moderadores=moderadores or [],
        professores=professores or [],
        turmas=turmas or [],
        alunos=alunos or [],
        responsaveis=responsaveis or [],
        ocorrencias=ocorrencias or []
    )


@app.route('/resetar_dados', methods=['POST'])
def resetar_dados():
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()

    # 🔒 exige senha do SAVIO para reset geral
    senha_savio = (request.form.get('senha_savio') or "").strip()
    if not _senha_savio_ok(conn, senha_savio):
        conn.close()
        flash("Senha do moderador SAVIO incorreta. Reset não permitido.")
        return redirect(url_for('exclusao'))

    cursor = conn.cursor()

    try:
        cursor.execute("PRAGMA foreign_keys = OFF")

        # ✅ Apaga logs (logins)
        cursor.execute("DELETE FROM logs_acessos")

        # ✅ Apaga dados principais (ajuste conforme suas tabelas)
        cursor.execute("DELETE FROM provas_lancamentos")
        cursor.execute("DELETE FROM provas")
        cursor.execute("DELETE FROM planejamento_itens")
        cursor.execute("DELETE FROM planejamentos_turmas")
        cursor.execute("DELETE FROM planejamentos")
        cursor.execute("DELETE FROM recados_aluno")
        cursor.execute("DELETE FROM recados")
        cursor.execute("DELETE FROM atestados")
        cursor.execute("DELETE FROM emprestimos_biblioteca")
        cursor.execute("DELETE FROM atendimentos_responsaveis")
        cursor.execute("DELETE FROM responsaveis")
        cursor.execute("DELETE FROM alunos")
        cursor.execute("DELETE FROM professor_turmas")
        cursor.execute("DELETE FROM professor_disciplinas")
        cursor.execute("DELETE FROM professores_turmas")
        cursor.execute("DELETE FROM professores")
        cursor.execute("DELETE FROM turmas")

        # Se tiver "ocorrencias" no seu banco, inclua:
        cursor.execute("DELETE FROM ocorrencias")

        cursor.execute("PRAGMA foreign_keys = ON")
        conn.commit()

        flash("Reset concluído! Dados e logs foram apagados (moderadores mantidos).")

    except sqlite3.Error as e:
        conn.rollback()
        flash(f"Erro ao resetar dados: {e}")

    finally:
        cursor.close()
        conn.close()

    # ✅ conforme você pediu: ao resetar, volta para dashboard do moderador
    return redirect(url_for('dashboard_moderador'))


# Dashboard do professor
@app.route('/dashboard_professor')
def dashboard_professor():
    if 'usuario' not in session or session['tipo'] != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_id = obter_professor_id(session['usuario'])

    conn = conectar_bd()
    cursor = conn.cursor()

    # ✅ Disciplinas/Funções escolhidas no cadastro (professor_disciplinas)
    cursor.execute("""
        SELECT disciplina
        FROM professor_disciplinas
        WHERE professor_id = ?
        ORDER BY disciplina
    """, (professor_id,))
    disciplinas_rows = cursor.fetchall()
    disciplinas_professor = [row['disciplina'] for row in disciplinas_rows] if disciplinas_rows else []

    # Turmas vinculadas ao professor
    cursor.execute('''
        SELECT t.id, t.nome, t.turno
        FROM turmas t
        JOIN professor_turmas pt ON pt.turma_id = t.id
        WHERE pt.professor_id = ?
        ORDER BY t.turno, t.nome
    ''', (professor_id,))
    todas_turmas = cursor.fetchall()

    if not todas_turmas:
        cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
        todas_turmas = cursor.fetchall()

    turmas_matutino = [turma for turma in todas_turmas if turma['turno'].lower() == 'matutino']
    turmas_vespertino = [turma for turma in todas_turmas if turma['turno'].lower() == 'vespertino']

    cursor.close()
    conn.close()

    return render_template(
        'dashboard_professor.html',
        turmas_matutino=turmas_matutino,
        turmas_vespertino=turmas_vespertino,
        usuario=session['usuario'],
        disciplinas_professor=disciplinas_professor  # ✅ enviado ao HTML
    )


@app.route('/atestados_professor', methods=['GET'])
def atestados_professor():
    # Apenas professor logado pode acessar
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_login = session['usuario']
    professor_id = obter_professor_id(professor_login)

    if not professor_id:
        flash("Professor não encontrado.")
        return redirect(url_for('dashboard_professor'))

    # Filtros (?turma_id=...&aluno_id=...&bimestre=...&ano=...)
    turma_id = (request.args.get('turma_id') or '').strip()
    aluno_id = (request.args.get('aluno_id') or '').strip()
    bimestre = (request.args.get('bimestre') or '').strip()
    ano = (request.args.get('ano') or '').strip()

    conn = conectar_bd()
    cursor = conn.cursor()

    # 1) Turmas ligadas ao professor
    cursor.execute("""
        SELECT DISTINCT
            t.id,
            t.nome,
            t.turno
        FROM turmas t
        JOIN professor_turmas pt ON pt.turma_id = t.id
        WHERE pt.professor_id = ?
        ORDER BY t.turno, t.nome
    """, (professor_id,))
    turmas_vinculadas = cursor.fetchall()
    tem_vinculo = bool(turmas_vinculadas)

    # Se não houver vínculo em professor_turmas, carrega TODAS as turmas
    if tem_vinculo:
        turmas = turmas_vinculadas
    else:
        cursor.execute("""
            SELECT id, nome, turno
            FROM turmas
            ORDER BY turno, nome
        """)
        turmas = cursor.fetchall()

    # 2) Alunos da turma selecionada (para o combo de estudantes)
    alunos_da_turma = []
    if turma_id:
        cursor.execute("""
            SELECT
                a.id,
                a.nome
            FROM alunos a
            WHERE a.turma_id = ?
            ORDER BY a.nome
        """, (turma_id,))
        alunos_da_turma = cursor.fetchall()

    # 3) Consulta principal dos atestados
    #    - Se o professor tiver vínculo, filtra pelas turmas vinculadas
    #    - Se NÃO tiver vínculo, mostra atestados conforme filtros, sem usar professor_turmas

    base_sql = """
        SELECT
            at.id,
            at.data_atestado,
            at.bimestre,
            at.tipo_atestado,
            at.outro_tipo,
            at.total_dias,
            a.nome AS aluno_nome,
            t.nome AS turma_nome,
            t.turno AS turma_turno
        FROM atestados at
        JOIN alunos a ON a.id  = at.aluno_id
        JOIN turmas t ON t.id  = at.turma_id
    """

    params = []

    if tem_vinculo:
        # Respeita apenas as turmas vinculadas ao professor
        sql = base_sql + """
            JOIN professor_turmas pt ON pt.turma_id = t.id
            WHERE pt.professor_id = ?
        """
        params.append(professor_id)
    else:
        # Sem vínculo: mostra atestados pelas turmas/filtros escolhidos
        sql = base_sql + " WHERE 1=1"

    # Filtro por turma
    if turma_id:
        sql += " AND t.id = ?"
        params.append(turma_id)

    # Filtro por aluno
    if aluno_id:
        sql += " AND a.id = ?"
        params.append(aluno_id)

    # Filtro por bimestre
    if bimestre:
        sql += " AND at.bimestre = ?"
        params.append(bimestre)

    # Filtro por ano, usando data_atestado (formato YYYY-MM-DD)
    if ano:
        sql += " AND strftime('%Y', at.data_atestado) = ?"
        params.append(ano)

    sql += " ORDER BY at.data_atestado DESC, aluno_nome"

    cursor.execute(sql, params)
    atestados = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'atestados_professor.html',
        atestados=atestados,
        turmas=turmas,
        alunos_da_turma=alunos_da_turma,
        filtro_turma=turma_id,
        filtro_aluno=aluno_id,
        filtro_bimestre=bimestre,
        filtro_ano=ano
    )


@app.route('/atestados/<int:atestado_id>')
def visualizar_atestado(atestado_id):
    if 'usuario' not in session:
        flash("Faça login para acessar.")
        return redirect(url_for('login'))

    tipo_usuario = session.get('tipo')

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("""
        SELECT
            at.id,
            at.aluno_id,
            a.nome AS aluno_nome,
            t.id   AS turma_id,
            t.nome AS turma_nome,
            t.turno,
            at.tipo_atestado,
            at.outro_tipo,
            at.data_atestado,
            at.bimestre,
            at.total_dias,
            at.observacoes,
            at.criado_em
        FROM atestados at
        JOIN alunos a ON a.id = at.aluno_id
        JOIN turmas t ON t.id = at.turma_id
        WHERE at.id = ?
    """, (atestado_id,))
    atestado = cursor.fetchone()

    cursor.close()
    conn.close()

    if not atestado:
        flash("Atestado não encontrado.")
        if tipo_usuario == 'professor':
            return redirect(url_for('atestados_professor'))
        elif tipo_usuario == 'moderador':
            return redirect(url_for('visualizar_atestados'))
        elif tipo_usuario == 'responsavel':
            return redirect(url_for('dashboard_responsavel'))
        else:
            return redirect(url_for('index'))

    # Se for professor, confere se a turma do atestado está vinculada a ele
    if tipo_usuario == 'professor':
        professor_id = obter_professor_id(session['usuario'])
        conn = conectar_bd()
        conn.row_factory = sqlite3.Row
        c2 = conn.cursor()
        c2.execute("""
            SELECT 1
            FROM professor_turmas
            WHERE professor_id = ? AND turma_id = ?
        """, (professor_id, atestado['turma_id']))
        vinculo = c2.fetchone()
        c2.close()
        conn.close()

        if not vinculo:
            flash("Você não tem permissão para visualizar este atestado.")
            return redirect(url_for('atestados_professor'))

    return render_template('visualizar_atestado.html', atestado=atestado)


# Turmas

@app.route('/cadastrar_turma', methods=['GET', 'POST'])
def cadastrar_turma():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    if request.method == 'POST':
        nome = request.form['nome']
        turno = request.form['turno']

        conn = conectar_bd()
        cursor = conn.cursor()

        try:
            cursor.execute("INSERT INTO turmas (nome, turno) VALUES (?, ?)", (nome, turno))
            conn.commit()
            flash("Turma cadastrada com sucesso!")
        except sqlite3.IntegrityError:
            flash("Essa turma já existe.")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('dashboard_moderador'))

    return render_template('cadastrar_turma.html')


# Responsáveis

@app.route('/login_responsavel', methods=['GET', 'POST'])
def login_responsavel():
    if request.method == 'POST':
        login_r = request.form['login']
        senha = request.form['senha']

        conn = conectar_bd()
        cursor = conn.cursor()
        cursor.execute("SELECT senha, aluno_id FROM responsaveis WHERE login = ?", (login_r,))
        responsavel = cursor.fetchone()
        cursor.close()
        conn.close()

        if responsavel and check_password_hash(responsavel['senha'], senha):
            session['responsavel'] = login_r
            session['aluno_id'] = responsavel['aluno_id']

            # Registrar log de acesso do responsável
            try:
                conn_log = conectar_bd()
                cursor_log = conn_log.cursor()
                cursor_log.execute(
                    "INSERT INTO logs_acessos (tipo, login) VALUES (?, ?)",
                    ('responsavel', login_r)
                )
                conn_log.commit()
            except sqlite3.Error:
                # Não interrompe o fluxo em caso de erro ao registrar o log
                pass
            finally:
                try:
                    cursor_log.close()
                except Exception:
                    pass
                try:
                    conn_log.close()
                except Exception:
                    pass

            return redirect(url_for('area_responsavel'))

        flash("Login ou senha inválidos.")
        return redirect(url_for('login_responsavel'))

    return render_template('login_responsavel.html')


@app.route('/area_responsavel')
def area_responsavel():
    if 'responsavel' not in session:
        flash("Faça login para acessar.")
        return redirect(url_for('login_responsavel'))

    aluno_id = session.get('aluno_id')

    # ========== CAPTURA DE TODOS OS FILTROS ==========
    # Planejamentos
    bim_plan = (request.args.get('bim_plan') or '').strip()
    disciplina_plan = (request.args.get('disciplina_plan') or '').strip()
    professor_plan = (request.args.get('professor_plan') or '').strip()

    # Avaliações
    bim_av = (request.args.get('bim_av') or '').strip()
    disciplina_av = (request.args.get('disciplina_av') or '').strip()
    professor_av = (request.args.get('professor_av') or '').strip()

    # Ocorrências
    bim_oc = (request.args.get('bim_oc') or '').strip()
    professor_oc = (request.args.get('professor_oc') or '').strip()

    # Recados
    bim_rec = (request.args.get('bim_rec') or '').strip()
    professor_rec = (request.args.get('professor_rec') or '').strip()

    # Atestados
    bim_atest = (request.args.get('bim_atest') or '').strip()

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # BUSCAR INFORMAÇÕES BÁSICAS DO ALUNO E TURMA
    cursor.execute(
        '''
        SELECT
            a.nome AS aluno_nome,
            t.nome AS turma_nome,
            t.id   AS turma_id
        FROM alunos a
        JOIN turmas t ON a.turma_id = t.id
        WHERE a.id = ?
        ''',
        (aluno_id,)
    )
    row = cursor.fetchone()
    if not row:
        cursor.close()
        conn.close()
        flash("Aluno não encontrado para este responsável.")
        return redirect(url_for('login_responsavel'))

    aluno_nome = row['aluno_nome']
    turma_nome = row['turma_nome']
    turma_id = row['turma_id']

    # ========== PLANEJAMENTOS ==========
    sql_plan = '''
        SELECT
            p.id,
            pr.login AS professor_login,
            p.disciplina,
            p.bimestre,
            p.ano,
            p.criado_em,
            p.observacoes,
            GROUP_CONCAT(DISTINCT t.nome || ' (' || t.turno || ')') AS turmas_nomes,
            SUM(COALESCE(pi.pontuacao_total, 0)) AS total_pontos
        FROM planejamentos p
        JOIN professores pr ON p.professor_id = pr.id
        LEFT JOIN planejamentos_turmas pt ON pt.planejamento_id = p.id
        LEFT JOIN turmas t              ON t.id = pt.turma_id
        LEFT JOIN planejamento_itens pi ON pi.planejamento_id = p.id
        WHERE t.id = ?
    '''
    params_plan = [turma_id]

    if bim_plan:
        try:
            b = int(bim_plan)
            sql_plan += " AND p.bimestre = ?"
            params_plan.append(b)
        except ValueError:
            pass

    if disciplina_plan:
        sql_plan += " AND p.disciplina = ?"
        params_plan.append(disciplina_plan)

    if professor_plan:
        sql_plan += " AND pr.login = ?"
        params_plan.append(professor_plan)

    sql_plan += '''
        GROUP BY
            p.id,
            pr.login,
            p.disciplina,
            p.bimestre,
            p.ano,
            p.criado_em,
            p.observacoes
        ORDER BY p.ano DESC, p.bimestre DESC, p.criado_em DESC
    '''

    cursor.execute(sql_plan, params_plan)
    planejamentos = cursor.fetchall()

    # Itens de planejamento
    itens_por_planejamento = {}
    if planejamentos:
        ids = [str(pl['id']) for pl in planejamentos]
        placeholders = ','.join(['?'] * len(ids))
        cursor.execute(f'''
            SELECT
                planejamento_id,
                conteudo,
                data_inicio,
                data_fim,
                forma_avaliacao,
                pontuacao_total,
                COALESCE(concluido, 0) AS concluido
            FROM planejamento_itens
            WHERE planejamento_id IN ({placeholders})
            ORDER BY planejamento_id, id
        ''', ids)
        itens_rows = cursor.fetchall()
        for it in itens_rows:
            pid = it['planejamento_id']
            itens_por_planejamento.setdefault(pid, []).append(it)

    # ========== AVALIAÇÕES BIMESTRAIS ==========
    sql_av = '''
        SELECT
            a.id,
            a.disciplina,
            a.bimestre,
            a.ano,
            a.tipo_avaliacao,
            a.descricao_avaliacao,
            a.conteudos,
            a.data_avaliacao,
            a.pontuacao,
            p.login AS professor_login
        FROM avaliacoes_bimestrais a
        JOIN professores p ON a.professor_id = p.id
        WHERE a.turma_id = ?
    '''
    params_av = [turma_id]

    if bim_av:
        try:
            b = int(bim_av)
            sql_av += " AND a.bimestre = ?"
            params_av.append(b)
        except ValueError:
            pass

    if disciplina_av:
        sql_av += " AND a.disciplina = ?"
        params_av.append(disciplina_av)

    if professor_av:
        sql_av += " AND p.login = ?"
        params_av.append(professor_av)

    sql_av += " ORDER BY a.ano DESC, a.bimestre, a.data_avaliacao"

    cursor.execute(sql_av, params_av)
    avaliacoes_rows = cursor.fetchall()

    avaliacoes = []
    for row in avaliacoes_rows:
        d = dict(row)
        data_str = d.get('data_avaliacao')
        if data_str:
            try:
                dt = datetime.strptime(data_str, '%Y-%m-%d')
                d['data_avaliacao_formatada'] = dt.strftime('%d/%m/%Y')
            except ValueError:
                d['data_avaliacao_formatada'] = data_str
        else:
            d['data_avaliacao_formatada'] = 'Não informada'
        avaliacoes.append(d)

    # ========== OCORRÊNCIAS ==========
    sql_oc = '''
        SELECT
            o.data,
            o.tipo_ocorrencia,
            o.motivo,
            o.total_dias,
            o.professor,
            o.chamar_responsavel,
            COALESCE(o.data_reuniao, 'N/A') AS data_reuniao,
            COALESCE(o.hora_reuniao, 'N/A') AS hora_reuniao
        FROM ocorrencias o
        WHERE o.aluno_id = ?
    '''
    params_oc = [aluno_id]

    if bim_oc:
        try:
            b = int(bim_oc)
            sql_oc += '''
                AND (((CAST(strftime('%m', o.data) AS INTEGER) - 1) / 3) + 1) = ?
            '''
            params_oc.append(b)
        except ValueError:
            pass

    if professor_oc:
        sql_oc += " AND o.professor = ?"
        params_oc.append(professor_oc)

    sql_oc += " ORDER BY o.data DESC"

    cursor.execute(sql_oc, params_oc)
    ocorrencias_rows = cursor.fetchall()

    ocorrencias = []
    for oc in ocorrencias_rows:
        d = dict(oc)
        data_str = d.get('data')
        if data_str:
            try:
                dt = datetime.strptime(data_str, '%Y-%m-%d')
                d['data_br'] = dt.strftime('%d/%m/%Y')
            except ValueError:
                d['data_br'] = data_str
        else:
            d['data_br'] = '-'

        data_reu = d.get('data_reuniao')
        if data_reu and data_reu != 'N/A':
            try:
                dt_reu = datetime.strptime(data_reu, '%Y-%m-%d')
                d['data_reuniao_br'] = dt_reu.strftime('%d/%m/%Y')
            except ValueError:
                d['data_reuniao_br'] = data_reu
        else:
            d['data_reuniao_br'] = ''

        ocorrencias.append(d)

    # ========== ATESTADOS DO ALUNO ==========
    sql_at = '''
        SELECT
            a.id,
            a.data_atestado,
            a.bimestre,
            a.tipo_atestado,
            a.total_dias,
            a.outro_tipo
        FROM atestados a
        WHERE a.aluno_id = ?
    '''
    params_at = [aluno_id]

    if bim_atest:
        try:
            b = int(bim_atest)
            sql_at += " AND a.bimestre = ?"
            params_at.append(b)
        except ValueError:
            pass

    sql_at += " ORDER BY date(a.data_atestado) DESC, a.id DESC"

    cursor.execute(sql_at, params_at)
    atestados_rows = cursor.fetchall()

    atestados_aluno = []
    for at in atestados_rows:
        d = dict(at)
        data_at = d.get('data_atestado')
        if data_at:
            try:
                dt = datetime.strptime(data_at, '%Y-%m-%d')
                d['data_atestado_br'] = dt.strftime('%d/%m/%Y')
            except ValueError:
                d['data_atestado_br'] = data_at
        else:
            d['data_atestado_br'] = '-'
        atestados_aluno.append(d)

    # ========== RECADOS POR ALUNO ==========
    sql_rec = '''
        SELECT
            ra.id,
            ra.data_criacao,
            ra.conteudo,
            ra.visualizado,
            p.login AS professor_login,
            t.nome AS turma_nome,
            GROUP_CONCAT(pd.disciplina, ', ') AS professor_funcao
        FROM recados_aluno ra
        JOIN professores p ON ra.professor_id = p.id
        JOIN turmas t ON ra.turma_id = t.id
        LEFT JOIN professor_disciplinas pd ON pd.professor_id = p.id
        WHERE ra.aluno_id = ?
          AND IFNULL(ra.excluido_para_responsavel, 0) = 0
    '''
    params_rec = [aluno_id]

    if bim_rec:
        try:
            b = int(bim_rec)
            sql_rec += '''
                AND (((CAST(strftime('%m', ra.data_criacao) AS INTEGER) - 1) / 3) + 1) = ?
            '''
            params_rec.append(b)
        except ValueError:
            pass

    if professor_rec:
        sql_rec += " AND p.login = ?"
        params_rec.append(professor_rec)

    sql_rec += '''
        GROUP BY ra.id
        ORDER BY ra.data_criacao DESC
    '''

    cursor.execute(sql_rec, params_rec)
    recados_rows = cursor.fetchall()

    recados_aluno = []
    for r in recados_rows:
        d = dict(r)
        data_cr = d.get('data_criacao')

        if data_cr:
            try:
                if isinstance(data_cr, datetime):
                    d['data_criacao_br'] = data_cr.strftime('%d/%m/%Y')
                elif ' ' in data_cr:
                    dt = datetime.strptime(data_cr, '%Y-%m-%d %H:%M:%S')
                    d['data_criacao_br'] = dt.strftime('%d/%m/%Y')
                else:
                    dt = datetime.strptime(data_cr, '%Y-%m-%d')
                    d['data_criacao_br'] = dt.strftime('%d/%m/%Y')
            except Exception:
                d['data_criacao_br'] = data_cr[:10]
        else:
            d['data_criacao_br'] = '-'

        if not d.get('professor_funcao'):
            d['professor_funcao'] = '-'

        recados_aluno.append(d)

    # ========== EMPRÉSTIMOS DA BIBLIOTECA ==========
    conn_bib = conectar_bd_biblioteca()
    conn_bib.row_factory = sqlite3.Row
    c_bib = conn_bib.cursor()

    c_bib.execute("""
        SELECT
            titulo_livro,
            data_emprestimo,
            data_prevista_devolucao,
            data_devolucao,
            status,
            devolucao_pontual
        FROM emprestimos_biblioteca
        WHERE aluno_id = ?
        ORDER BY data_emprestimo DESC
    """, (aluno_id,))

    emprestimos_raw = c_bib.fetchall()
    c_bib.close()
    conn_bib.close()

    emprestimos_biblioteca = []
    for e in emprestimos_raw:
        d = dict(e)

        def formata_data(val):
            if not val:
                return '-'
            try:
                return datetime.strptime(val, '%Y-%m-%d').strftime('%d/%m/%Y')
            except ValueError:
                return val

        d['data_emprestimo_br'] = formata_data(d.get('data_emprestimo'))
        d['data_prevista_devolucao_br'] = formata_data(d.get('data_prevista_devolucao'))
        d['data_devolucao_br'] = formata_data(d.get('data_devolucao'))
        emprestimos_biblioteca.append(d)

    # ========== BUSCAR LISTAS PARA OS FILTROS ==========

    # Disciplinas de planejamentos da turma
    cursor.execute("""
        SELECT DISTINCT p.disciplina
        FROM planejamentos p
        INNER JOIN planejamentos_turmas pt ON p.id = pt.planejamento_id
        WHERE pt.turma_id = ?
        ORDER BY p.disciplina
    """, (turma_id,))
    disciplinas_planejamento = [row['disciplina'] for row in cursor.fetchall()]

    # Professores de planejamentos da turma
    cursor.execute("""
        SELECT DISTINCT pr.login
        FROM planejamentos p
        INNER JOIN professores pr ON p.professor_id = pr.id
        INNER JOIN planejamentos_turmas pt ON p.id = pt.planejamento_id
        WHERE pt.turma_id = ?
        ORDER BY pr.login
    """, (turma_id,))
    professores_planejamento = [row['login'] for row in cursor.fetchall()]

    # Disciplinas de avaliações da turma
    cursor.execute("""
        SELECT DISTINCT disciplina
        FROM avaliacoes_bimestrais
        WHERE turma_id = ?
        ORDER BY disciplina
    """, (turma_id,))
    disciplinas_avaliacao = [row['disciplina'] for row in cursor.fetchall()]

    # Professores de avaliações da turma
    cursor.execute("""
        SELECT DISTINCT p.login
        FROM avaliacoes_bimestrais a
        INNER JOIN professores p ON a.professor_id = p.id
        WHERE a.turma_id = ?
        ORDER BY p.login
    """, (turma_id,))
    professores_avaliacao = [row['login'] for row in cursor.fetchall()]

    # Professores de ocorrências do aluno
    cursor.execute("""
        SELECT DISTINCT professor
        FROM ocorrencias
        WHERE aluno_id = ?
          AND professor IS NOT NULL
          AND professor != ''
        ORDER BY professor
    """, (aluno_id,))
    professores_ocorrencia = [row['professor'] for row in cursor.fetchall()]

    # Professores de recados do aluno
    cursor.execute("""
        SELECT DISTINCT p.login
        FROM recados_aluno ra
        INNER JOIN professores p ON ra.professor_id = p.id
        WHERE ra.aluno_id = ?
        ORDER BY p.login
    """, (aluno_id,))
    professores_recado = [row['login'] for row in cursor.fetchall()]

    cursor.close()
    conn.close()

    # ========== RETORNAR TEMPLATE COM TODAS AS VARIÁVEIS ==========
    return render_template(
        'area_responsavel.html',
        aluno_nome=aluno_nome,
        turma_nome=turma_nome,

        # Dados
        planejamentos=planejamentos,
        itens_por_planejamento=itens_por_planejamento,
        avaliacoes=avaliacoes,
        ocorrencias=ocorrencias,
        atestados_aluno=atestados_aluno,
        recados_aluno=recados_aluno,
        emprestimos_biblioteca=emprestimos_biblioteca,

        # Filtros atuais
        bim_plan=bim_plan,
        disciplina_plan=disciplina_plan,
        professor_plan=professor_plan,

        bim_av=bim_av,
        disciplina_av=disciplina_av,
        professor_av=professor_av,

        bim_oc=bim_oc,
        professor_oc=professor_oc,

        bim_rec=bim_rec,
        professor_rec=professor_rec,

        bim_atest=bim_atest,

        # Listas para popular os dropdowns (✅ ESTAS ESTAVAM FALTANDO!)
        disciplinas_planejamento=disciplinas_planejamento,
        professores_planejamento=professores_planejamento,

        disciplinas_avaliacao=disciplinas_avaliacao,
        professores_avaliacao=professores_avaliacao,

        professores_ocorrencia=professores_ocorrencia,
        professores_recado=professores_recado
    )


@app.route('/cadastrar_responsavel', methods=['GET', 'POST'])
def cadastrar_responsavel():
    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # ✅ Busca termo ativo (para exibir no GET e validar no POST)
    termo = get_termo_ativo(conectar_bd, "responsavel")

    if request.method == 'POST':
        login_r = request.form['login']
        senha = request.form['senha']
        aluno_id = request.form['aluno_id']
        telefone = (request.form.get('telefone') or '').strip()
        senha_hash = generate_password_hash(senha)

        # ✅ Exige termo ativo
        if not termo:
            flash("Termo não encontrado/ativo. Contate a equipe da escola.")
            cursor.close()
            conn.close()
            return redirect(url_for('cadastrar_responsavel'))

        # ✅ Exige aceite marcado
        if not request.form.get("aceite_termo"):
            flash("Para concluir o cadastro, você precisa aceitar o Termo de Uso.")
            cursor.close()
            conn.close()
            return redirect(url_for('cadastrar_responsavel'))

        if not telefone:
            flash("Por favor, informe um número de telefone/WhatsApp atualizado.")
            cursor.close()
            conn.close()
            return redirect(url_for('cadastrar_responsavel'))

        # Verifica limite de até 2 responsáveis por aluno
        cursor.execute("SELECT COUNT(*) AS n FROM responsaveis WHERE aluno_id = ?", (aluno_id,))
        num_responsaveis = cursor.fetchone()["n"]

        if num_responsaveis >= 2:
            flash("Este aluno já possui dois responsáveis cadastrados.")
            cursor.close()
            conn.close()
            return redirect(url_for('cadastrar_responsavel'))

        try:
            cursor.execute(
                "INSERT INTO responsaveis (login, senha, telefone, aluno_id) VALUES (?, ?, ?, ?)",
                (login_r, senha_hash, telefone, aluno_id)
            )
            conn.commit()

            # ✅ Registra aceite do termo (tipo_cadastro = responsavel)
            registrar_aceite(conectar_bd, termo, "responsavel", login_r)

            flash("Responsável cadastrado com sucesso!")
        except sqlite3.IntegrityError:
            flash("Login já está em uso. Escolha outro.")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('login_responsavel'))

    # GET – carrega turmas + termo
    cursor.execute("SELECT id, nome FROM turmas ORDER BY nome")
    turmas = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template('cadastrar_responsavel.html', turmas=turmas, termo=termo)


@app.route('/redefinir_senha', methods=['GET', 'POST'])
def redefinir_senha():
    if request.method == 'POST':
        login_r = request.form.get('login')
        nova_senha = request.form.get('nova_senha')

        if not login_r or not nova_senha:
            flash("Por favor, preencha todos os campos.")
            return redirect(url_for('redefinir_senha'))

        conn = conectar_bd()
        cursor = conn.cursor()

        try:
            cursor.execute("SELECT id FROM responsaveis WHERE login = ?", (login_r,))
            responsavel = cursor.fetchone()

            if not responsavel:
                flash("Login não encontrado.")
                cursor.close()
                conn.close()
                return redirect(url_for('redefinir_senha'))

            nova_senha_hash = generate_password_hash(nova_senha)
            cursor.execute("UPDATE responsaveis SET senha = ? WHERE login = ?", (nova_senha_hash, login_r))
            conn.commit()
            flash("Senha redefinida com sucesso!")
        except sqlite3.Error as e:
            flash(f"Erro ao redefinir senha: {e}")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('login_responsavel'))

    return render_template('redefinir_senha.html')


@app.route('/visualizar_responsaveis', methods=['GET'])
def visualizar_responsaveis():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    cursor.execute("SELECT id, nome FROM turmas")
    turmas = cursor.fetchall()

    turma_id = request.args.get('turma_id')
    responsaveis = []
    if turma_id:
        cursor.execute('''
            SELECT
                responsaveis.id,
                responsaveis.login,
                responsaveis.telefone,
                alunos.nome AS aluno_nome
            FROM responsaveis
            JOIN alunos ON responsaveis.aluno_id = alunos.id
            WHERE alunos.turma_id = ?
        ''', (turma_id,))
        responsaveis = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template('visualizar_responsaveis.html', turmas=turmas, responsaveis=responsaveis)


@app.route('/excluir_responsavel', methods=['POST'])
def excluir_responsavel():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    responsavel_id = request.form['responsavel_id']

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM responsaveis WHERE id = ?", (responsavel_id,))
    conn.commit()
    cursor.close()
    conn.close()

    return redirect(url_for('visualizar_responsaveis'))


# Alunos

@app.route('/cadastrar_aluno', methods=['GET', 'POST'])
def cadastrar_aluno():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id, nome, turno FROM turmas")
    turmas = cursor.fetchall()

    if request.method == 'POST':
        turma_id = (request.form.get('turma_id') or "").strip()

        # nomes digitados manualmente (já existia no seu sistema)
        nomes = request.form.getlist('nomes[]')

        # NOVO: PDF opcional
        pdf_file = request.files.get('pdf_alunos')

        if not turma_id:
            flash("Por favor, selecione uma turma.")
            cursor.close()
            conn.close()
            return redirect(url_for('cadastrar_aluno'))

        # Se veio PDF, extrai e junta com os nomes digitados
        if pdf_file and pdf_file.filename:
            try:
                nomes_pdf = extrair_nomes_alunos_do_pdf(pdf_file)
                nomes.extend(nomes_pdf)
            except Exception as e:
                flash(f"Não foi possível importar do PDF: {e}")
                cursor.close()
                conn.close()
                return redirect(url_for('cadastrar_aluno'))

        # Limpa lista final
        nomes = [(n or "").strip() for n in nomes]
        nomes = [n for n in nomes if n]

        if not nomes:
            flash("Insira ao menos um nome OU envie um PDF para importar.")
            cursor.close()
            conn.close()
            return redirect(url_for('cadastrar_aluno'))

        # Evitar duplicar nomes já existentes na turma
        cursor.execute("SELECT nome FROM alunos WHERE turma_id = ?", (turma_id,))
        existentes = {(row["nome"] or "").strip().upper() for row in cursor.fetchall()}

        novos = []
        vistos_form = set()

        for n in nomes:
            key = n.upper()
            if key in existentes:
                continue
            if key in vistos_form:
                continue
            vistos_form.add(key)
            novos.append((n, turma_id))

        try:
            if novos:
                antes = conn.total_changes
                cursor.executemany("INSERT OR IGNORE INTO alunos (nome, turma_id) VALUES (?, ?)", novos)
                conn.commit()
                inseridos = conn.total_changes - antes
                ignorados = len(novos) - inseridos
                flash(
                    f"Importação concluída: {inseridos} aluno(s) cadastrado(s). {ignorados} ignorado(s) por já existirem.")
            else:
                flash("Nenhum aluno novo para cadastrar (todos já existiam na turma).")
        except sqlite3.Error as e:
            conn.rollback()
            flash(f"Erro ao cadastrar alunos: {e}")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('cadastrar_aluno'))

    cursor.close()
    conn.close()
    return render_template('cadastrar_aluno.html', turmas=turmas)


@app.route('/selecionar_turma', methods=['GET'])
def selecionar_turma():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    turma_id = request.args.get('turma_id')
    if not turma_id:
        flash("Turma não selecionada.")
        return redirect(url_for('dashboard_moderador'))

    conn = conectar_bd()
    cursor = conn.cursor()

    cursor.execute('''
        SELECT alunos.id, alunos.nome
        FROM alunos
        WHERE alunos.turma_id = ?
    ''', (turma_id,))
    alunos_da_turma = cursor.fetchall()

    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()

    cursor.execute("SELECT id, login FROM professores WHERE status = 'pendente'")
    professores_pendentes = cursor.fetchall()

    cursor.execute("SELECT id, login, status FROM professores WHERE status IN ('aprovado', 'rejeitado')")
    professores_aprovados = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'dashboard_moderador.html',
        turmas=turmas,
        alunos_da_turma=alunos_da_turma,
        professores_pendentes=professores_pendentes,
        professores_aprovados=professores_aprovados
    )


@app.route('/excluir_aluno/<int:aluno_id>', methods=['POST'])
def excluir_aluno(aluno_id):
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        cursor.execute("DELETE FROM alunos WHERE id = ?", (aluno_id,))
        conn.commit()
        flash("Aluno excluído com sucesso.")
    except sqlite3.Error as e:
        flash(f"Erro ao excluir aluno: {e}")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('dashboard_moderador'))


@app.route('/excluir_turma/<int:turma_id>', methods=['POST'])
def excluir_turma(turma_id):
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT COUNT(*) AS total FROM alunos WHERE turma_id = ?", (turma_id,))
        total_alunos = cursor.fetchone()['total']

        if total_alunos > 0:
            flash("Não é possível excluir a turma. Existem alunos cadastrados nela.")
            cursor.close()
            conn.close()
            return redirect(url_for('dashboard_moderador'))

        cursor.execute("DELETE FROM turmas WHERE id = ?", (turma_id,))
        conn.commit()
        flash("Turma excluída com sucesso.")
    except sqlite3.Error as e:
        flash(f"Erro ao excluir turma: {e}")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('dashboard_moderador'))


# Ocorrências

@app.route('/registrar_ocorrencia', methods=['GET', 'POST'])
def registrar_ocorrencia():
    if 'usuario' not in session or session.get('tipo') not in ('professor', 'moderador'):
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    if request.method == 'POST':
        turma_id = request.form['turma_id']
        aluno_id = request.form['aluno_id']
        data_ocorrencia = request.form['data_ocorrencia']
        tipo_ocorrencia = request.form['tipo_ocorrencia']
        motivo = request.form['motivo']
        professor = request.form['professor']
        chamar_responsavel = request.form.get('chamar_responsavel', 'nao')
        data_reuniao = request.form.get('data_reuniao')
        hora_reuniao = request.form.get('hora_reuniao')
        total_dias = request.form.get('total_dias', 0) if tipo_ocorrencia == 'SUSPENSAO' else 0

        conn = conectar_bd()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO ocorrencias (aluno_id, turma_id, data, tipo_ocorrencia, motivo, professor,
                                     chamar_responsavel, data_reuniao, hora_reuniao, total_dias)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (aluno_id, turma_id, data_ocorrencia, tipo_ocorrencia, motivo, professor,
              chamar_responsavel, data_reuniao, hora_reuniao, total_dias))
        conn.commit()
        cursor.close()
        conn.close()

        if session['tipo'] == 'professor':
            return redirect(url_for('dashboard_professor'))
        return redirect(url_for('dashboard_moderador'))

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id, nome FROM turmas")
    turmas = cursor.fetchall()
    cursor.execute("SELECT id, nome FROM alunos")
    alunos = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template('registro_ocorrencia.html', turmas=turmas, alunos=alunos)


@app.route('/visualizar_ocorrencias', methods=['GET'])
def visualizar_ocorrencias():
    if 'usuario' not in session or session.get('tipo') not in ('professor', 'moderador'):
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    # Carregar todas as turmas
    cursor.execute("SELECT id, nome FROM turmas ORDER BY nome")
    turmas = cursor.fetchall()

    # Filtros recebidos pela URL
    turma_id = request.args.get('turma_id', '').strip()
    aluno_id = request.args.get('aluno_id', '').strip()

    alunos_da_turma = []
    ocorrencias = []

    # Se uma turma foi selecionada, carregar alunos da turma
    if turma_id:
        cursor.execute(
            "SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome",
            (turma_id,)
        )
        alunos_da_turma = cursor.fetchall()

        # Montar SQL das ocorrências com professor
        sql = '''
            SELECT
                ocorrencias.id,                -- 0
                alunos.nome AS aluno_nome,     -- 1
                turmas.nome AS turma_nome,     -- 2
                ocorrencias.professor,         -- 3
                ocorrencias.data,              -- 4
                ocorrencias.tipo_ocorrencia,   -- 5
                ocorrencias.motivo,            -- 6
                ocorrencias.total_dias,        -- 7
                ocorrencias.chamar_responsavel,-- 8
                COALESCE(ocorrencias.data_reuniao, 'N/A') AS data_reuniao, -- 9
                COALESCE(ocorrencias.hora_reuniao, 'N/A') AS hora_reuniao  -- 10
            FROM ocorrencias
            JOIN alunos ON ocorrencias.aluno_id = alunos.id
            JOIN turmas ON ocorrencias.turma_id = turmas.id
            WHERE turmas.id = ?
        '''
        params = [turma_id]

        # Se um aluno específico foi selecionado, filtra também
        if aluno_id:
            sql += " AND alunos.id = ?"
            params.append(aluno_id)

        sql += " ORDER BY ocorrencias.data DESC"

        cursor.execute(sql, params)
        ocorrencias = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'visualizar_ocorrencias.html',
        ocorrencias=ocorrencias,
        turmas=turmas,
        turma_selecionada=turma_id,
        alunos_da_turma=alunos_da_turma,
        aluno_selecionado=aluno_id
    )


@app.route('/ocorrencias/download_pdf/<int:turma_id>')
def download_ocorrencias_pdf(turma_id):
    if 'usuario' not in session or session.get('tipo') not in ('professor', 'moderador'):
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    cursor.execute('''
        SELECT
            alunos.nome AS aluno_nome,
            turmas.nome AS turma_nome,
            ocorrencias.professor,
            ocorrencias.data,
            ocorrencias.tipo_ocorrencia,
            ocorrencias.motivo,
            ocorrencias.total_dias,
            ocorrencias.chamar_responsavel,
            COALESCE(ocorrencias.data_reuniao, 'N/A') AS data_reuniao,
            COALESCE(ocorrencias.hora_reuniao, 'N/A') AS hora_reuniao
        FROM ocorrencias
        JOIN alunos ON ocorrencias.aluno_id = alunos.id
        JOIN turmas ON ocorrencias.turma_id = turmas.id
        WHERE turmas.id = ?
        ORDER BY ocorrencias.data DESC
    ''', (turma_id,))
    ocorrencias = cursor.fetchall()
    cursor.close()
    conn.close()

    pdf_buffer = BytesIO()
    pdf = canvas.Canvas(pdf_buffer, pagesize=letter)
    pdf.setTitle("Relatório de Ocorrências")

    turma_nome = ocorrencias[0]["turma_nome"] if ocorrencias else "N/A"
    pdf.drawString(80, 750, f"Relatório de Ocorrências - Turma {turma_nome}")

    y_position = 720
    for ocorrencia in ocorrencias:
        pdf.drawString(50, y_position,
                       f"Aluno: {ocorrencia['aluno_nome']}  |  Professor: {ocorrencia['professor'] or '-'}")
        pdf.drawString(50, y_position - 15,
                       f"Data: {ocorrencia['data']}  |  Tipo: {ocorrencia['tipo_ocorrencia']}")
        pdf.drawString(50, y_position - 30,
                       f"Motivo: {ocorrencia['motivo']}")
        pdf.drawString(50, y_position - 45,
                       f"Chamar responsável: {'Sim' if ocorrencia['chamar_responsavel'] == 'sim' else 'Não'}  "
                       f"| Dias de suspensão: {ocorrencia['total_dias'] or 0}")
        pdf.drawString(50, y_position - 60,
                       f"Reunião: {ocorrencia['data_reuniao']} às {ocorrencia['hora_reuniao']}")
        y_position -= 80

        if y_position < 100:
            pdf.showPage()
            y_position = 750

    pdf.save()
    pdf_buffer.seek(0)
    return send_file(
        pdf_buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f"ocorrencias_turma_{turma_id}.pdf"
    )


@app.route('/excluir_ocorrencia', methods=['GET', 'POST'])
def excluir_ocorrencia():
    if 'usuario' not in session or session.get('tipo') not in ('professor', 'moderador'):
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    if request.method == 'POST':
        ocorrencia_id = request.form['ocorrencia_id']
        cursor.execute("DELETE FROM ocorrencias WHERE id = ?", (ocorrencia_id,))
        conn.commit()
        cursor.close()
        conn.close()
        return redirect(url_for('excluir_ocorrencia'))

    cursor.execute('''
        SELECT ocorrencias.id, alunos.nome, turmas.nome, ocorrencias.data, ocorrencias.tipo_ocorrencia, ocorrencias.motivo
        FROM ocorrencias
        JOIN alunos ON ocorrencias.aluno_id = alunos.id
        JOIN turmas ON ocorrencias.turma_id = turmas.id
        ORDER BY ocorrencias.data DESC
    ''')
    ocorrencias = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template('excluir_ocorrencia.html', ocorrencias=ocorrencias)


# Atestados (gestão de comprovantes de alunos)

@app.route('/registrar_atestado', methods=['GET', 'POST'])
def registrar_atestado():
    # Apenas moderador (direção/coordenação entram como moderador no sistema)
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    if request.method == 'POST':
        bimestre = request.form.get('bimestre')
        turma_id = request.form.get('turma_id')
        aluno_id = request.form.get('aluno_id')
        tipo_atestado = (request.form.get('tipo_atestado') or '').strip().upper()
        data_atestado = request.form.get('data_atestado')
        total_dias_str = (request.form.get('total_dias') or '').strip()
        outro_tipo = (request.form.get('outro_tipo') or '').strip()

        # Campos obrigatórios
        if not bimestre or not turma_id or not aluno_id or not tipo_atestado or not data_atestado:
            flash("Preencha todos os campos obrigatórios.")
            cursor.close()
            conn.close()
            return redirect(url_for('registrar_atestado'))

        # Tratamento de dias de afastamento
        total_dias = None
        if tipo_atestado == 'AFASTAMENTO' and total_dias_str:
            try:
                total_dias = int(total_dias_str)
            except ValueError:
                total_dias = None

        # Tratamento do campo "OUTROS"
        if tipo_atestado == 'OUTROS':
            if not outro_tipo:
                flash("Descreva o tipo de atestado em 'Outros'.")
                cursor.close()
                conn.close()
                return redirect(url_for('registrar_atestado'))
        else:
            outro_tipo = None

        try:
            cursor.execute("""
                INSERT INTO atestados
                    (bimestre, turma_id, aluno_id, tipo_atestado, data_atestado, total_dias, outro_tipo)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (bimestre, turma_id, aluno_id, tipo_atestado, data_atestado, total_dias, outro_tipo))
            conn.commit()
            flash("Atestado registrado com sucesso!")
        except sqlite3.Error as e:
            conn.rollback()
            flash(f"Erro ao registrar atestado: {e}")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('visualizar_atestados'))

    # GET – carregar turmas e alunos
    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()

    cursor.execute("SELECT id, nome, turma_id FROM alunos ORDER BY nome")
    alunos_rows = cursor.fetchall()

    # CONVERTE Row -> dict PARA USAR |tojson NO TEMPLATE
    alunos = [
        {
            "id": row["id"],
            "nome": row["nome"],
            "turma_id": row["turma_id"]
        }
        for row in alunos_rows
    ]

    data_hoje = datetime.now().strftime('%Y-%m-%d')

    cursor.close()
    conn.close()

    return render_template(
        'registrar_atestado.html',
        turmas=turmas,
        alunos=alunos,
        data_hoje=data_hoje
    )


@app.route('/visualizar_atestados', methods=['GET'])
def visualizar_atestados():
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    bimestre = request.args.get('bimestre', '').strip()
    turma_id = request.args.get('turma_id', '').strip()
    aluno_id = request.args.get('aluno_id', '').strip()

    conn = conectar_bd()
    cursor = conn.cursor()

    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()

    alunos_da_turma = []
    if turma_id:
        cursor.execute("SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome", (turma_id,))
        alunos_da_turma = cursor.fetchall()

    query = '''
        SELECT
            atestados.id,
            atestados.bimestre,
            atestados.data_atestado,
            atestados.tipo_atestado,
            atestados.outro_tipo,
            atestados.total_dias,
            alunos.nome AS aluno_nome,
            turmas.nome AS turma_nome,
            turmas.turno AS turma_turno
        FROM atestados
        JOIN alunos ON atestados.aluno_id = alunos.id
        JOIN turmas ON atestados.turma_id = turmas.id
        WHERE 1=1
    '''
    params = []

    if bimestre:
        query += " AND atestados.bimestre = ?"
        params.append(bimestre)

    if turma_id:
        query += " AND atestados.turma_id = ?"
        params.append(turma_id)

    if aluno_id:
        query += " AND atestados.aluno_id = ?"
        params.append(aluno_id)

    query += " ORDER BY atestados.data_atestado DESC, alunos.nome"

    cursor.execute(query, params)
    atestados = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template(
        'visualizar_atestados.html',
        atestados=atestados,
        turmas=turmas,
        alunos_da_turma=alunos_da_turma,
        filtro_bimestre=bimestre,
        filtro_turma=turma_id,
        filtro_aluno=aluno_id
    )


@app.route('/editar_atestado/<int:atestado_id>', methods=['GET', 'POST'])
def editar_atestado(atestado_id):
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    if request.method == 'POST':
        senha_moderador = request.form.get('senha_moderador', '').strip()
        cursor.execute("SELECT senha FROM moderadores WHERE login = ?", (session.get('usuario'),))
        moderador = cursor.fetchone()

        if not moderador or not check_password_hash(moderador['senha'], senha_moderador):
            cursor.close()
            conn.close()
            flash('Senha de moderador incorreta.')
            return redirect(url_for('editar_atestado', atestado_id=atestado_id))

        bimestre = request.form.get('bimestre', '').strip()
        turma_id = request.form.get('turma_id')
        aluno_id = request.form.get('aluno_id')
        tipo_atestado = request.form.get('tipo_atestado', '').strip()
        outro_tipo = request.form.get('outro_tipo', '').strip()
        data_atestado = request.form.get('data_atestado', '').strip()
        total_dias_raw = request.form.get('total_dias', '').strip()

        total_dias = 0
        if tipo_atestado.upper() == 'AFASTAMENTO':
            try:
                total_dias = int(total_dias_raw) if total_dias_raw else 0
            except ValueError:
                total_dias = 0

        if tipo_atestado.upper() != 'OUTROS':
            outro_tipo = None

        cursor.execute(
            '''
            UPDATE atestados
            SET bimestre = ?, turma_id = ?, aluno_id = ?, tipo_atestado = ?, outro_tipo = ?, total_dias = ?, data_atestado = ?
            WHERE id = ?
            ''',
            (bimestre, turma_id, aluno_id, tipo_atestado, outro_tipo, total_dias, data_atestado, atestado_id)
        )
        conn.commit()
        cursor.close()
        conn.close()

        flash('Atestado atualizado com sucesso!')
        return redirect(url_for('visualizar_atestados'))

    # GET - carregar dados do atestado
    cursor.execute(
        '''
        SELECT
            atestados.*,
            alunos.nome AS aluno_nome,
            turmas.nome AS turma_nome,
            turmas.turno AS turma_turno
        FROM atestados
        JOIN alunos ON atestados.aluno_id = alunos.id
        JOIN turmas ON atestados.turma_id = turmas.id
        WHERE atestados.id = ?
        ''',
        (atestado_id,)
    )
    atestado = cursor.fetchone()

    if not atestado:
        cursor.close()
        conn.close()
        flash('Atestado não encontrado.')
        return redirect(url_for('visualizar_atestados'))

    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()
    cursor.execute("SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome", (atestado['turma_id'],))
    alunos_da_turma = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'editar_atestado.html',
        atestado=atestado,
        turmas=turmas,
        alunos_da_turma=alunos_da_turma
    )


@app.route('/excluir_atestado/<int:atestado_id>', methods=['POST'])
def excluir_atestado(atestado_id):
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    senha_moderador = request.form.get('senha_moderador', '').strip()

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT senha FROM moderadores WHERE login = ?", (session.get('usuario'),))
    moderador = cursor.fetchone()

    if not moderador or not check_password_hash(moderador['senha'], senha_moderador):
        cursor.close()
        conn.close()
        flash('Senha de moderador incorreta. Atestado não excluído.')
        return redirect(url_for('visualizar_atestados'))

    cursor.execute("DELETE FROM atestados WHERE id = ?", (atestado_id,))
    conn.commit()
    cursor.close()
    conn.close()

    flash('Atestado excluído com sucesso!')
    return redirect(url_for('visualizar_atestados'))


# Recados

@app.route('/registrar_recado', methods=['GET', 'POST'])
def registrar_recado():
    if 'usuario' not in session or session['tipo'] != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    if request.method == 'POST':
        conteudo = request.form.get('conteudo')

        if not conteudo:
            flash("Por favor, preencha o recado.")
            return redirect(url_for('registrar_recado'))

        conn = conectar_bd()
        cursor = conn.cursor()

        try:
            cursor.execute("INSERT INTO recados (conteudo) VALUES (?)", (conteudo,))
            conn.commit()
            flash("Recado registrado com sucesso!")
        except sqlite3.Error as e:
            flash(f"Erro ao registrar recado: {e}")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('dashboard_professor'))

    return render_template('registrar_recado.html')


@app.route('/visualizar_recados', methods=['GET', 'POST'])
def visualizar_recados():
    if 'usuario' not in session or session['tipo'] not in ('professor', 'moderador'):
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id, conteudo, data_criacao FROM recados ORDER BY data_criacao DESC")
    recados = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template('visualizar_recados.html', recados=recados)


@app.route('/excluir_recado/<int:recado_id>', methods=['POST'])
def excluir_recado(recado_id):
    if 'usuario' not in session or session['tipo'] not in ('professor', 'moderador'):
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        cursor.execute("DELETE FROM recados WHERE id = ?", (recado_id,))
        conn.commit()
        flash("Recado excluído com sucesso!")
    except sqlite3.Error as e:
        flash(f"Erro ao excluir recado: {e}")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('visualizar_recados'))


# Recados POR ALUNO (professor -> aluno específico)

@app.route('/recados_aluno/registrar', methods=['GET', 'POST'])
def registrar_recado_aluno():
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_login = session['usuario']
    professor_id = obter_professor_id(professor_login)

    conn = conectar_bd()
    cursor = conn.cursor()

    # Turmas em que o professor leciona
    turmas = obter_turmas_professor(professor_id)
    if not turmas:
        cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
        turmas = cursor.fetchall()

    if request.method == 'POST':
        turma_id = request.form.get('turma_id')
        alunos_ids = request.form.getlist('alunos[]')
        recado_id = (request.form.get('recado_id') or "").strip()
        data_registro = request.form.get('data_registro') or datetime.now().strftime('%Y-%m-%d')

        if not turma_id or not alunos_ids or not recado_id:
            flash("Selecione a turma, pelo menos um aluno e um modelo de recado.")
            cursor.close()
            conn.close()
            return redirect(url_for('registrar_recado_aluno'))

        # Busca o texto do recado-modelo escolhido (AGORA NA TABELA CERTA)
        cursor.execute("SELECT conteudo FROM recados_modelos WHERE id = ?", (recado_id,))
        recado_row = cursor.fetchone()
        if not recado_row:
            flash("Modelo de recado não encontrado.")
            cursor.close()
            conn.close()
            return redirect(url_for('registrar_recado_aluno'))

        conteudo = recado_row['conteudo']

        try:
            for aluno_id in alunos_ids:
                cursor.execute('''
                    INSERT INTO recados_aluno
                    (professor_id, aluno_id, turma_id, conteudo, data_criacao)
                    VALUES (?, ?, ?, ?, ?)
                ''', (professor_id, aluno_id, turma_id, conteudo, data_registro))
            conn.commit()
            flash("Recado registrado com sucesso para o(s) aluno(s) selecionado(s).")
        except sqlite3.Error as e:
            conn.rollback()
            flash(f"Erro ao registrar recado: {e}")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('listar_recados_aluno'))

    # GET – carrega turmas e modelos de recado A PARTIR DE recados_modelos
    cursor.execute("SELECT id, conteudo FROM recados_modelos ORDER BY id")
    recados_modelos = cursor.fetchall()

    data_hoje = datetime.now().strftime('%Y-%m-%d')
    cursor.close()
    conn.close()
    return render_template(
        'registrar_recado_aluno.html',
        turmas=turmas,
        data_hoje=data_hoje,
        recados_modelos=recados_modelos
    )


@app.route('/recados_aluno/listar')
def listar_recados_aluno():
    # Apenas professor logado
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_login = session['usuario']
    professor_id = obter_professor_id(professor_login)

    turma_id = (request.args.get('turma_id') or '').strip()
    aluno_id = (request.args.get('aluno_id') or '').strip()

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # Turmas do professor para o filtro
    turmas = obter_turmas_professor(professor_id)
    if not turmas:
        cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
        turmas = cursor.fetchall()

    # Alunos da turma selecionada (para o combo de alunos)
    alunos = []
    if turma_id:
        cursor.execute(
            "SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome",
            (turma_id,)
        )
        alunos = cursor.fetchall()

    # Recados lançados pelo professor (traz ra.visualizado e o conteúdo gravado no recado)
    sql = """
        SELECT
            ra.id,
            ra.data_criacao,
            ra.visualizado,
            IFNULL(ra.excluido_para_responsavel, 0) AS excluido_para_responsavel,
            ra.excluido_em,
            ra.excluido_por_login,
            t.nome  AS turma_nome,
            a.nome  AS aluno_nome,
            ra.conteudo
        FROM recados_aluno ra
        JOIN alunos a  ON a.id = ra.aluno_id
        JOIN turmas t  ON t.id = ra.turma_id
        WHERE ra.professor_id = ?
    """
    params = [professor_id]
    params = [professor_id]

    if turma_id:
        sql += " AND ra.turma_id = ?"
        params.append(turma_id)

    if aluno_id:
        sql += " AND ra.aluno_id = ?"
        params.append(aluno_id)

    sql += " ORDER BY ra.data_criacao DESC, turma_nome, aluno_nome"

    cursor.execute(sql, params)
    recados = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'listar_recados_aluno.html',
        recados=recados,
        turmas=turmas,
        alunos=alunos,
        turma_selecionada=turma_id,
        aluno_selecionado=aluno_id,
    )


@app.route('/recados_aluno/excluir/<int:recado_aluno_id>', methods=['POST'])
def excluir_recado_aluno_para_responsavel(recado_aluno_id):
    # Apenas professor logado
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_login = session['usuario']
    professor_id = obter_professor_id(professor_login)
    if not professor_id:
        flash("Professor não encontrado.")
        return redirect(url_for('dashboard_professor'))

    conn = conectar_bd()
    cursor = conn.cursor()
    try:
        # Confere se o recado pertence ao professor logado
        cursor.execute(
            "SELECT id FROM recados_aluno WHERE id = ? AND professor_id = ?",
            (recado_aluno_id, professor_id)
        )
        row = cursor.fetchone()
        if not row:
            flash("Você não tem permissão para excluir este recado.")
            return redirect(url_for('listar_recados_aluno'))

        cursor.execute(
            """UPDATE recados_aluno
               SET excluido_para_responsavel = 1,
                   excluido_em = datetime('now','localtime'),
                   excluido_por_login = ?
               WHERE id = ? AND professor_id = ?""",
            (professor_login, recado_aluno_id, professor_id)
        )
        conn.commit()
        flash("Recado removido da Área do Responsável (mantido no histórico interno).")
    except sqlite3.Error as e:
        conn.rollback()
        flash(f"Erro ao excluir recado: {e}")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('listar_recados_aluno'))


@app.route('/recados_aluno/marcar_lido/<int:recado_id>', methods=['POST'])
def marcar_recado_aluno_lido(recado_id):
    """
    Marca um recado como visualizado quando o responsável abre o detalhe.
    Essa rota é chamada via fetch() na área do responsável.
    """
    # Se quiser, você pode restringir para quem estiver na área do responsável,
    # mas como o acesso já é controlado pela sessão, geralmente basta assim.

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        cursor.execute(
            "UPDATE recados_aluno SET visualizado = 1 WHERE id = ?",
            (recado_id,)
        )
        conn.commit()
        sucesso = cursor.rowcount > 0
    except sqlite3.Error:
        conn.rollback()
        sucesso = False
    finally:
        cursor.close()
        conn.close()

    return jsonify({'success': sucesso})


@app.route('/recados_aluno/gestor', methods=['GET'])
def listar_recados_aluno_gestor():
    # Somente moderador/coordenação
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    turma_id = (request.args.get('turma_id') or '').strip()
    aluno_id = (request.args.get('aluno_id') or '').strip()
    professor_id = (request.args.get('professor_id') or '').strip()

    conn = conectar_bd()
    cursor = conn.cursor()

    # Turmas para o filtro
    cursor.execute("SELECT id, nome FROM turmas ORDER BY nome")
    turmas = cursor.fetchall()

    # Professores aprovados para o filtro
    cursor.execute("SELECT id, login FROM professores WHERE status = 'aprovado' ORDER BY login")
    professores = cursor.fetchall()

    # Alunos (se tiver turma escolhida, filtra por ela)
    if turma_id:
        cursor.execute("SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome", (turma_id,))
    else:
        cursor.execute("SELECT id, nome FROM alunos ORDER BY nome")
    alunos = cursor.fetchall()

    # Consulta principal dos recados (AGORA COM excluido_para_responsavel)
    sql = '''
        SELECT
            ra.id,
            ra.data_criacao,
            ra.visualizado,
            IFNULL(ra.excluido_para_responsavel, 0) AS excluido_para_responsavel,
            ra.excluido_em,
            ra.excluido_por_login,

            a.id  AS aluno_id,
            a.nome AS aluno_nome,

            t.id  AS turma_id,
            t.nome AS turma_nome,
            t.turno AS turno,

            p.id  AS professor_id,
            p.login AS professor_login,

            ra.conteudo
        FROM recados_aluno ra
        JOIN alunos a       ON ra.aluno_id     = a.id
        JOIN turmas t       ON ra.turma_id     = t.id
        JOIN professores p  ON ra.professor_id = p.id
        WHERE 1 = 1
    '''
    params = []

    if turma_id:
        sql += ' AND ra.turma_id = ?'
        params.append(int(turma_id))

    if aluno_id:
        sql += ' AND ra.aluno_id = ?'
        params.append(int(aluno_id))

    if professor_id:
        sql += ' AND ra.professor_id = ?'
        params.append(int(professor_id))

    sql += '''
        ORDER BY
            ra.data_criacao DESC,
            t.turno,
            t.nome,
            a.nome,
            p.login
    '''

    cursor.execute(sql, params)
    recados = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        'recados_aluno_gestor.html',
        recados=recados,
        turmas=turmas,
        alunos=alunos,
        professores=professores,
        turma_selecionada=turma_id,
        aluno_selecionado=aluno_id,
        professor_selecionado=professor_id
    )


# Verificar senha para ações especiais

@app.route('/verificar_senha', methods=['POST'])
def verificar_senha():
    senha_digitada = request.form.get('senha')
    destino = request.form.get('destino')

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT senha FROM moderadores WHERE login = 'SAVIO'")
        moderador = cursor.fetchone()

        if moderador and check_password_hash(moderador['senha'], senha_digitada):
            if destino == 'registrar':
                return redirect(url_for('registrar_recado'))
            elif destino == 'visualizar':
                return redirect(url_for('visualizar_recados'))
            elif destino == 'gerar_pdf':
                return redirect(url_for('gerar_pdf'))
            else:
                flash("Destino inválido.")
                return redirect(url_for('dashboard_professor'))
        else:
            flash("Senha do moderador incorreta.")
            return redirect(url_for('dashboard_professor'))
    except Exception as e:
        flash(f"Erro ao verificar senha: {e}")
        return redirect(url_for('dashboard_professor'))
    finally:
        cursor.close()
        conn.close()


# Gerar PDF geral

@app.route('/gerar_pdf', methods=['GET', 'POST'])
def gerar_pdf():
    # Apenas moderador pode gerar PDF
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    # GET: só mostra a tela com os filtros
    if request.method == 'GET':
        conn = conectar_bd()
        cursor = conn.cursor()

        # Turmas para o filtro de atestados
        cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
        turmas = cursor.fetchall()

        cursor.close()
        conn.close()
        return render_template('gerar_pdf.html', turmas=turmas)

    # POST: gera o PDF de fato
    tabelas_selecionadas = request.form.getlist('tabelas')

    if not tabelas_selecionadas:
        flash("Selecione pelo menos uma opção para gerar o PDF.")
        return redirect(url_for('gerar_pdf'))

    # Filtros específicos
    bimestre_planejamento = (request.form.get('bimestre_planejamento') or "").strip()
    bimestre_atestado = (request.form.get('bimestre_atestado') or "").strip()
    turma_atestado = (request.form.get('turma_atestado') or "").strip()

    conn = conectar_bd()
    cursor = conn.cursor()

    # ------------------------------------
    # Função auxiliar para formatar datas
    # ------------------------------------
    def fmt_data(data_str):
        if not data_str:
            return ""
        try:
            for padrao in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y"):
                try:
                    d = datetime.strptime(str(data_str)[:10], padrao)
                    return d.strftime("%d/%m/%Y")
                except ValueError:
                    continue
            return str(data_str)
        except Exception:
            return str(data_str)

    # ------------------------------------
    # Montagem do PDF
    # ------------------------------------
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),  # A4 HORIZONTAL
        leftMargin=24,
        rightMargin=24,
        topMargin=30,
        bottomMargin=24
    )

    styles = getSampleStyleSheet()
    titulo_style = styles["Title"]
    h2 = styles["Heading2"]
    h3 = styles["Heading3"]
    normal = styles["Normal"]

    wrap_style = ParagraphStyle(
        "wrap_style",
        parent=normal,
        fontSize=7,
        leading=9,
        spaceAfter=0,
        spaceBefore=0
    )

    def _esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _pwrap(s: str):
        s = "" if s is None else str(s)
        s = s.strip()
        if not s:
            return Paragraph("—", wrap_style)
        s = _esc(s)
        s = s.replace("\r\n", "<br/>").replace("\n", "<br/>").replace("\r", "<br/>")
        return Paragraph(s, wrap_style)

    story = []

    static_dir = os.path.join(app.root_path, 'static')
    logo_left_path = os.path.join(static_dir, 'logo.jpg')
    logo_right_path = os.path.join(static_dir, 'logo1.PNG')

    def _mk_logo(p, w=52, h=52):
        try:
            if p and os.path.exists(p):
                return Image(p, width=w, height=h)
        except Exception:
            pass
        return ''

    header_tbl = Table(
        [[_mk_logo(logo_left_path), Paragraph('<b>RELATÓRIO – DE OLHO NA ESCOLA (ESCOLA CLASSE 16)</b>', normal),
          _mk_logo(logo_right_path)]],
        colWidths=[60, None, 60],
    )
    header_tbl.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),
        ('ALIGN', (2, 0), (2, 0), 'RIGHT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(header_tbl)
    story.append(Spacer(1, 10))

    story.append(Paragraph("Relatório Geral – De Olho na Escola", titulo_style))
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", normal))
    story.append(Spacer(1, 16))

    # =========================
    # 1) PROFESSORES
    # =========================
    if "professores" in tabelas_selecionadas:
        cursor.execute("""
            SELECT id, login, status
            FROM professores
            ORDER BY login
        """)
        rows = cursor.fetchall()

        story.append(Paragraph("Professores", h2))
        story.append(Spacer(1, 6))

        if rows:
            dados_tabela = [["ID", "Login", "Status"]]
            for r in rows:
                dados_tabela.append([r["id"], r["login"], r["status"]])

            tabela = Table(dados_tabela, repeatRows=1)
            tabela.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]))
            story.append(tabela)
        else:
            story.append(Paragraph("Nenhum professor cadastrado.", normal))

        story.append(Spacer(1, 16))

    # =========================
    # 2) TURMAS
    # =========================
    if "turmas" in tabelas_selecionadas:
        cursor.execute("""
            SELECT id, nome, turno
            FROM turmas
            ORDER BY turno, nome
        """)
        rows = cursor.fetchall()

        story.append(Paragraph("Turmas", h2))
        story.append(Spacer(1, 6))

        if rows:
            dados_tabela = [["ID", "Nome", "Turno"]]
            for r in rows:
                dados_tabela.append([r["id"], r["nome"], r["turno"]])

            tabela = Table(dados_tabela, repeatRows=1)
            tabela.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]))
            story.append(tabela)
        else:
            story.append(Paragraph("Nenhuma turma cadastrada.", normal))

        story.append(Spacer(1, 16))

    # =========================
    # 3) ALUNOS
    # =========================
    if "alunos" in tabelas_selecionadas:
        cursor.execute("""
            SELECT a.id,
                   a.nome,
                   t.nome  AS turma,
                   t.turno AS turno
            FROM alunos a
            JOIN turmas t ON a.turma_id = t.id
            ORDER BY t.turno, t.nome, a.nome
        """)
        rows = cursor.fetchall()

        story.append(Paragraph("Alunos", h2))
        story.append(Spacer(1, 6))

        if rows:
            dados_tabela = [["ID", "Nome", "Turma", "Turno"]]
            for r in rows:
                dados_tabela.append([r["id"], r["nome"], r["turma"], r["turno"]])

            tabela = Table(dados_tabela, repeatRows=1)
            tabela.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]))
            story.append(tabela)
        else:
            story.append(Paragraph("Nenhum aluno cadastrado.", normal))

        story.append(Spacer(1, 16))

    # =========================
    # 4) RESPONSÁVEIS
    # =========================
    if "responsaveis" in tabelas_selecionadas:
        cursor.execute("""
            SELECT r.id,
                   r.login,
                   r.telefone,
                   a.nome AS aluno
            FROM responsaveis r
            LEFT JOIN alunos a ON r.aluno_id = a.id
            ORDER BY r.login
        """)
        rows = cursor.fetchall()

        story.append(Paragraph("Responsáveis", h2))
        story.append(Spacer(1, 6))

        if rows:
            dados_tabela = [["ID", "Login", "Telefone", "Aluno"]]
            for r in rows:
                dados_tabela.append([
                    r["id"],
                    r["login"],
                    r["telefone"] or "",
                    r["aluno"] or "",
                ])

            tabela = Table(dados_tabela, repeatRows=1)
            tabela.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]))
            story.append(tabela)
        else:
            story.append(Paragraph("Nenhum responsável cadastrado.", normal))

        story.append(Spacer(1, 16))

    # =========================
    # 5) OCORRÊNCIAS
    # =========================
    if "ocorrencias" in tabelas_selecionadas:
        try:
            cursor.execute("""
                SELECT
                    o.id,
                    a.nome AS aluno,
                    t.nome AS turma,
                    t.turno AS turno,
                    o.data,
                    o.tipo_ocorrencia AS tipo,
                    o.motivo,
                    o.professor,
                    COALESCE(o.data_reuniao, '') AS data_reuniao,
                    COALESCE(o.hora_reuniao, '') AS hora_reuniao,
                    COALESCE(o.descricao, '') AS descricao
                FROM ocorrencias o
                JOIN alunos a ON o.aluno_id = a.id
                JOIN turmas t ON a.turma_id = t.id
                ORDER BY t.turno, t.nome, o.data
            """)
            rows = cursor.fetchall()

            story.append(Paragraph("Ocorrências", h2))
            story.append(Spacer(1, 6))

            if rows:
                dados_tabela = [["ID", "Aluno", "Turma", "Turno",
                                 "Data", "Tipo", "Motivo",
                                 "Professor", "Descrição", "Data Reunião", "Hora Reunião"]]
                for r in rows:
                    dados_tabela.append([
                        r["id"],
                        r["aluno"],
                        r["turma"],
                        r["turno"],
                        fmt_data(r["data"]),
                        r["tipo"],
                        r["motivo"],
                        r["professor"],
                        _pwrap(r["descricao"]),
                        fmt_data(r["data_reuniao"]) if r["data_reuniao"] else "",
                        r["hora_reuniao"] or "",
                    ])

                tabela = Table(dados_tabela, repeatRows=1)
                tabela.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                ]))
                story.append(tabela)
            else:
                story.append(Paragraph("Nenhuma ocorrência cadastrada.", normal))

            story.append(Spacer(1, 16))
        except sqlite3.Error:
            story.append(Paragraph("Não foi possível carregar as ocorrências.", normal))
            story.append(Spacer(1, 16))

    # =========================
    # 6) ATESTADOS (por turma e bimestre)
    # =========================
    if "atestados" in tabelas_selecionadas:
        sql_at = """
            SELECT
                at.data_atestado,
                at.bimestre,
                t.nome  AS turma,
                t.turno AS turno,
                a.nome  AS aluno,
                at.tipo_atestado,
                COALESCE(at.outro_tipo, '') AS outro_tipo,
                at.total_dias
            FROM atestados at
            JOIN alunos a ON a.id = at.aluno_id
            JOIN turmas t ON t.id = at.turma_id
            WHERE 1=1
        """
        params_at = []

        if bimestre_atestado:
            sql_at += " AND at.bimestre = ?"
            params_at.append(bimestre_atestado)

        if turma_atestado:
            sql_at += " AND t.id = ?"
            params_at.append(turma_atestado)

        sql_at += """
            ORDER BY t.turno, t.nome, at.bimestre, at.data_atestado, a.nome
        """

        cursor.execute(sql_at, params_at)
        rows = cursor.fetchall()

        story.append(Paragraph("Atestados por Turma e Bimestre", h2))
        story.append(Spacer(1, 6))

        filtros_txt = []
        if bimestre_atestado:
            filtros_txt.append(f"Bimestre: {bimestre_atestado}")
        if turma_atestado:
            cursor.execute("SELECT nome, turno FROM turmas WHERE id = ?", (turma_atestado,))
            turma_sel = cursor.fetchone()
            if turma_sel:
                filtros_txt.append(f"Turma: {turma_sel['nome']} ({turma_sel['turno']})")
        if filtros_txt:
            story.append(Paragraph("Filtros aplicados: " + " | ".join(filtros_txt), normal))
            story.append(Spacer(1, 6))

        if rows:
            dados_tabela = [["Data", "Bimestre", "Turma", "Turno",
                             "Aluno", "Tipo", "Dias"]]
            for r in rows:
                tipo = r["tipo_atestado"]
                if tipo == "outro" and r["outro_tipo"]:
                    tipo = r["outro_tipo"]

                dados_tabela.append([
                    fmt_data(r["data_atestado"]),
                    r["bimestre"],
                    r["turma"],
                    r["turno"],
                    r["aluno"],
                    tipo,
                    r["total_dias"],
                ])

            tabela = Table(dados_tabela, repeatRows=1)
            tabela.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]))
            story.append(tabela)
        else:
            story.append(Paragraph("Nenhum atestado encontrado com os filtros informados.", normal))

        story.append(Spacer(1, 16))

    # =========================
    # 7) ATENDIMENTOS – EQUIPE GESTORA (MODERADORES)
    # =========================
    if "atendimentos_gestao" in tabelas_selecionadas:
        try:
            cursor.execute("""
                SELECT
                    ar.id,
                    ar.data_atendimento,
                    ar.hora_atendimento,
                    t.nome  AS turma,
                    t.turno AS turno,
                    a.nome  AS aluno,
                    COALESCE(NULLIF(ar.responsavel_nome,''), r.login) AS responsavel,
                    COALESCE(ar.registrador_nome,'')  AS registrador_nome,
                    COALESCE(ar.registrador_cargo,'') AS registrador_cargo,
                    COALESCE(ar.assunto,'')           AS assunto,
                    ar.envolve_professor,
                    COALESCE(ar.professor_nome,'')    AS professor_nome,
                    COALESCE(ar.relato,'')            AS relato,
                    COALESCE(ar.combinados,'')        AS combinados,
                    ar.retorno_previsto,
                    COALESCE(ar.retorno_em,'')        AS retorno_em,
                    ar.reuniao_agendada,
                    COALESCE(ar.reuniao_data,'')      AS reuniao_data
                FROM atendimentos_responsaveis ar
                JOIN turmas t ON t.id = ar.turma_id
                JOIN alunos a ON a.id = ar.aluno_id
                LEFT JOIN responsaveis r ON r.id = ar.responsavel_id
                ORDER BY ar.data_atendimento DESC, ar.hora_atendimento DESC
            """)
            rows = cursor.fetchall()

            story.append(Paragraph("Atendimentos – Equipe Gestora (Moderadores)", h2))
            story.append(Spacer(1, 6))

            if rows:
                dados_tabela = [[
                    "Data", "Hora", "Turma", "Aluno(a)", "Responsável",
                    "Assunto", "Envolve professor", "Relato", "Combinados",
                    "Retorno", "Reunião", "Registrado por"
                ]]

                for r in rows:
                    retorno_txt = "Não"
                    if r["retorno_previsto"]:
                        retorno_txt = "Sim"
                        if r["retorno_em"]:
                            retorno_txt += f" ({fmt_data(r['retorno_em'])})"

                    reuniao_txt = "Não"
                    if r["reuniao_agendada"]:
                        reuniao_txt = "Sim"
                        if r["reuniao_data"]:
                            reuniao_txt += f" ({fmt_data(r['reuniao_data'])})"

                    envolve_prof_txt = "Não"
                    if r["envolve_professor"]:
                        envolve_prof_txt = "Sim"
                        if r["professor_nome"]:
                            envolve_prof_txt += f" – {r['professor_nome']}"

                    registrador_txt = (r["registrador_nome"] or "").strip()
                    cargo_txt = (r["registrador_cargo"] or "").strip()
                    if cargo_txt:
                        registrador_txt = (registrador_txt + " – " + cargo_txt).strip(" –")
                    if not registrador_txt:
                        registrador_txt = "—"

                    dados_tabela.append([
                        fmt_data(r["data_atendimento"]),
                        (r["hora_atendimento"] or ""),
                        f"{r['turma']} ({r['turno']})",
                        r["aluno"],
                        (r["responsavel"] or "—"),
                        _pwrap(r["assunto"]),
                        envolve_prof_txt,
                        _pwrap(r["relato"]),
                        _pwrap(r["combinados"]),
                        retorno_txt,
                        reuniao_txt,
                        _pwrap(registrador_txt),
                    ])

                tabela = Table(dados_tabela, repeatRows=1)
                tabela.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                ]))
                story.append(tabela)
            else:
                story.append(Paragraph("Nenhum atendimento de gestão cadastrado.", normal))

            story.append(Spacer(1, 16))
        except sqlite3.Error:
            story.append(Paragraph("Não foi possível carregar os atendimentos da gestão.", normal))
            story.append(Spacer(1, 16))

    # =========================
    # 8) ATENDIMENTOS – SOE (sem encaminhamentos)
    # =========================
    if "soe_atendimentos" in tabelas_selecionadas:
        try:
            try:
                ensure_soe_table()
            except Exception:
                pass

            cursor.execute("""
                SELECT
                    s.id,
                    s.data_atendimento,
                    s.hora_atendimento,
                    t.nome  AS turma,
                    t.turno AS turno,
                    a.nome  AS aluno,
                    COALESCE(s.responsavel_nome,'')  AS responsavel_nome,
                    COALESCE(s.orientadora_nome,'')  AS orientadora_nome,
                    COALESCE(s.assunto,'')           AS assunto,
                    COALESCE(s.relato,'')            AS relato,
                    COALESCE(s.combinados,'')        AS combinados,
                    s.retorno_previsto,
                    COALESCE(s.retorno_em,'')        AS retorno_em,
                    s.reuniao_agendada,
                    COALESCE(s.reuniao_data,'')      AS reuniao_data
                FROM soe_atendimentos s
                JOIN turmas t ON t.id = s.turma_id
                JOIN alunos a ON a.id = s.aluno_id
                ORDER BY s.data_atendimento DESC, s.hora_atendimento DESC
            """)
            rows = cursor.fetchall()

            story.append(Paragraph("Atendimentos – SOE", h2))
            story.append(Spacer(1, 6))

            if rows:
                dados_tabela = [[
                    "Data", "Hora", "Turma", "Aluno(a)", "Responsável",
                    "Orientadora", "Assunto", "Relato", "Combinados",
                    "Retorno", "Reunião"
                ]]

                for r in rows:
                    retorno_txt = "Não"
                    if r["retorno_previsto"]:
                        retorno_txt = "Sim"
                        if r["retorno_em"]:
                            retorno_txt += f" ({fmt_data(r['retorno_em'])})"

                    reuniao_txt = "Não"
                    if r["reuniao_agendada"]:
                        reuniao_txt = "Sim"
                        if r["reuniao_data"]:
                            reuniao_txt += f" ({fmt_data(r['reuniao_data'])})"

                    dados_tabela.append([
                        fmt_data(r["data_atendimento"]),
                        (r["hora_atendimento"] or ""),
                        f"{r['turma']} ({r['turno']})",
                        r["aluno"],
                        _pwrap(r["responsavel_nome"]),
                        _pwrap(r["orientadora_nome"]),
                        _pwrap(r["assunto"]),
                        _pwrap(r["relato"]),
                        _pwrap(r["combinados"]),
                        retorno_txt,
                        reuniao_txt,
                    ])

                tabela = Table(dados_tabela, repeatRows=1)
                tabela.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                ]))
                story.append(tabela)
            else:
                story.append(Paragraph("Nenhum atendimento do SOE cadastrado.", normal))

            story.append(Spacer(1, 16))
        except sqlite3.Error:
            story.append(Paragraph("Não foi possível carregar os atendimentos do SOE.", normal))
            story.append(Spacer(1, 16))

    # =========================
    # 9) BIBLIOTECA – Empréstimos (datas e status)
    # =========================
    if "emprestimos_biblioteca" in tabelas_selecionadas:
        try:
            cursor.execute("""
                SELECT
                    e.id,
                    e.data_emprestimo,
                    e.data_prevista_devolucao,
                    e.data_devolucao,
                    e.status,
                    e.titulo_livro,
                    COALESCE(e.autor,'') AS autor,
                    COALESCE(e.codigo_interno,'') AS codigo_interno,
                    a.nome AS aluno,
                    t.nome AS turma,
                    t.turno AS turno
                FROM emprestimos_biblioteca e
                JOIN alunos a ON a.id = e.aluno_id
                JOIN turmas t ON t.id = e.turma_id
                ORDER BY e.data_emprestimo DESC
            """)
            rows = cursor.fetchall()

            story.append(Paragraph("Biblioteca – Empréstimos de Livros", h2))
            story.append(Spacer(1, 6))

            if rows:
                dados_tabela = [[
                    "Data empréstimo", "Prevista devolução", "Data devolução", "Status",
                    "Turma", "Aluno(a)", "Título", "Autor", "Código"
                ]]
                for r in rows:
                    dados_tabela.append([
                        fmt_data(r["data_emprestimo"]),
                        fmt_data(r["data_prevista_devolucao"]),
                        fmt_data(r["data_devolucao"]),
                        r["status"],
                        f"{r['turma']} ({r['turno']})",
                        r["aluno"],
                        _pwrap(r["titulo_livro"]),
                        _pwrap(r["autor"]),
                        _pwrap(r["codigo_interno"]),
                    ])

                tabela = Table(dados_tabela, repeatRows=1)
                tabela.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                ]))
                story.append(tabela)
            else:
                story.append(Paragraph("Nenhum empréstimo de biblioteca cadastrado.", normal))

            story.append(Spacer(1, 16))
        except sqlite3.Error:
            story.append(Paragraph("Não foi possível carregar os empréstimos da biblioteca.", normal))
            story.append(Spacer(1, 16))

    # =========================
    # 10) PLANEJAMENTOS
    # =========================
    if "planejamentos" in tabelas_selecionadas:
        sql_pl = """
            SELECT
                p.id,
                pr.login AS professor,
                p.disciplina,
                p.bimestre,
                p.ano,
                p.criado_em,
                COALESCE(p.observacoes, '') AS observacoes
            FROM planejamentos p
            JOIN professores pr ON pr.id = p.professor_id
            WHERE 1=1
        """
        params_pl = []

        if bimestre_planejamento:
            sql_pl += " AND p.bimestre = ?"
            params_pl.append(bimestre_planejamento)

        sql_pl += " ORDER BY p.ano DESC, p.bimestre, pr.login, p.disciplina, p.id"

        cursor.execute(sql_pl, params_pl)
        planejamentos = cursor.fetchall()

        story.append(Paragraph("Planejamentos – Resumo", h2))
        story.append(Spacer(1, 6))

        if bimestre_planejamento:
            story.append(Paragraph(f"Filtrado pelo bimestre: {bimestre_planejamento}", normal))
            story.append(Spacer(1, 6))

        if not planejamentos:
            story.append(Paragraph("Nenhum planejamento encontrado com os filtros informados.", normal))
            story.append(Spacer(1, 16))
        else:
            for p in planejamentos:
                cursor.execute("""
                    SELECT t.nome, t.turno
                    FROM planejamentos_turmas pt
                    JOIN turmas t ON t.id = pt.turma_id
                    WHERE pt.planejamento_id = ?
                    ORDER BY t.turno, t.nome
                """, (p["id"],))
                turmas_pl = cursor.fetchall()
                turmas_txt = ", ".join(
                    f"{t['nome']} ({t['turno']})" for t in turmas_pl
                ) or "Sem turma vinculada"

                story.append(Paragraph(
                    f"Planejamento #{p['id']} – Prof.: {p['professor']} – "
                    f"{p['disciplina']} – {p['bimestre']}º bimestre/{p['ano']}",
                    h3
                ))
                story.append(Paragraph(f"Turmas: {turmas_txt}", normal))
                if p["observacoes"]:
                    story.append(Paragraph(f"Observações: {p['observacoes']}", normal))

                story.append(Spacer(1, 4))

                cursor.execute("""
                    SELECT
                        conteudo,
                        data_inicio,
                        data_fim
                    FROM planejamento_itens
                    WHERE planejamento_id = ?
                    ORDER BY data_inicio, data_fim, conteudo
                """, (p["id"],))
                itens = cursor.fetchall()

                if itens:
                    dados_tabela = [["Conteúdo", "Período"]]
                    for item in itens:
                        periodo = ""
                        if item["data_inicio"] or item["data_fim"]:
                            inicio = fmt_data(item["data_inicio"]) if item["data_inicio"] else ""
                            fim = fmt_data(item["data_fim"]) if item["data_fim"] else ""
                            if inicio and fim:
                                periodo = f"{inicio} a {fim}"
                            else:
                                periodo = inicio or fim

                        dados_tabela.append([
                            _pwrap(item["conteudo"]),
                            periodo
                        ])

                    tabela = Table(dados_tabela, colWidths=[350, 150], repeatRows=1)
                    tabela.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ]))
                    story.append(tabela)
                else:
                    story.append(Paragraph("Nenhum item de conteúdo cadastrado neste planejamento.", normal))

                story.append(Spacer(1, 16))

    # =========================
    # 11) AVALIAÇÕES BIMESTRAIS
    # =========================
    if "avaliacoes" in tabelas_selecionadas:
        sql_av = """
            SELECT
                a.id                             AS id,
                t.nome || ' (' || t.turno || ')' AS turma,
                a.disciplina                     AS disciplina,
                a.bimestre                       AS bimestre,
                a.ano                            AS ano,
                a.tipo_avaliacao                 AS tipo,
                a.descricao_avaliacao            AS descricao,
                a.conteudos                      AS conteudos,
                a.data_avaliacao                 AS data_avaliacao,
                a.pontuacao                      AS pontuacao,
                p.login                          AS professor
            FROM avaliacoes_bimestrais a
            JOIN professores p ON a.professor_id = p.id
            JOIN turmas t      ON a.turma_id = t.id
            WHERE 1=1
        """
        params_av = []

        if bimestre_planejamento:
            try:
                b = int(bimestre_planejamento)
                sql_av += " AND a.bimestre = ?"
                params_av.append(b)
            except ValueError:
                pass

        sql_av += " ORDER BY a.ano DESC, a.bimestre, a.data_avaliacao"

        cursor.execute(sql_av, params_av)
        rows = cursor.fetchall()

        story.append(Paragraph("Avaliações Bimestrais", h2))
        story.append(Spacer(1, 6))

        if rows:
            dados_tabela = [[
                "ID", "Turma", "Disciplina", "Bimestre", "Ano",
                "Tipo", "Descrição", "Conteúdos", "Data", "Pontuação", "Professor"
            ]]

            for r in rows:
                dados_tabela.append([
                    r["id"],
                    r["turma"],
                    r["disciplina"],
                    f"{r['bimestre']}º",
                    r["ano"],
                    r["tipo"],
                    _pwrap(r["descricao"]),
                    _pwrap(r["conteudos"]),
                    fmt_data(r["data_avaliacao"]),
                    r["pontuacao"],
                    r["professor"],
                ])

            tabela = Table(dados_tabela, repeatRows=1)
            tabela.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
            ]))
            story.append(tabela)
        else:
            story.append(Paragraph("Nenhuma avaliação bimestral cadastrada.", normal))

        story.append(Spacer(1, 16))

    cursor.close()
    conn.close()

    doc.build(story)
    pdf = buffer.getvalue()
    buffer.close()

    response = make_response(pdf)
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = "attachment; filename=relatorio_geral.pdf"
    return response


@app.route('/visualizar_turmas/<int:turma_id>')
def visualizar_alunos_turma(turma_id):
    conn = conectar_bd()
    cursor = conn.cursor()

    cursor.execute("SELECT nome FROM turmas WHERE id = ?", (turma_id,))
    turma = cursor.fetchone()

    cursor.execute("SELECT id, nome FROM alunos WHERE turma_id = ?", (turma_id,))
    alunos = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template('visualizar_alunos_turma.html', turma=turma, alunos=alunos)


@app.route('/get_alunos/<int:turma_id>')
def get_alunos(turma_id):
    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id, nome FROM alunos WHERE turma_id = ?", (turma_id,))
    alunos = cursor.fetchall()
    cursor.close()
    conn.close()

    return jsonify([{'id': aluno['id'], 'nome': aluno['nome']} for aluno in alunos])


@app.route('/get_alunos_turma/<int:turma_id>')
def get_alunos_turma(turma_id):
    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT id, nome FROM alunos WHERE turma_id = ?", (turma_id,))
    alunos = cursor.fetchall()
    cursor.close()
    conn.close()

    return jsonify([{'id': aluno['id'], 'nome': aluno['nome']} for aluno in alunos])


# PLANEJAMENTO UNIFICADO (bimestral + mensal)

@app.route('/planejamento/registrar', methods=['GET', 'POST'])
def registrar_planejamento():
    if 'usuario' not in session or session['tipo'] != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_login = session['usuario']
    professor_id = obter_professor_id(professor_login)

    hoje = datetime.now()
    ano_atual = hoje.year
    bimestre_atual = ((hoje.month - 1) // 3) + 1

    conn = conectar_bd()
    cursor = conn.cursor()

    # Disciplinas do professor
    cursor.execute("""
        SELECT disciplina
        FROM professor_disciplinas
        WHERE professor_id = ?
        ORDER BY disciplina
    """, (professor_id,))
    disciplinas_rows = cursor.fetchall()
    disciplinas_professor = [row['disciplina'] for row in disciplinas_rows] if disciplinas_rows else []

    # Turmas
    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY nome")
    turmas = cursor.fetchall()

    if request.method == 'POST':
        # ----------- Dados gerais -----------
        try:
            ano = int(request.form.get('ano_letivo', ano_atual))
        except:
            ano = ano_atual

        try:
            bimestre = int(request.form.get('bimestre', bimestre_atual))
        except:
            bimestre = bimestre_atual

        disciplina = (request.form.get('disciplina') or "").strip()
        observacoes = (request.form.get('observacoes_gerais') or "").strip()

        turmas_ids = request.form.getlist('turmas[]') or []
        conteudos = request.form.getlist('conteudos[]') or []

        # ✅ NOVO (seu formulário novo)
        datas_inicio = request.form.getlist('datas_inicio[]') or []
        datas_fim = request.form.getlist('datas_fim[]') or []

        # ✅ ANTIGO (mantido por compatibilidade, se existir algum template velho ainda)
        periodos = request.form.getlist('periodos[]') or []
        instrumentos = request.form.getlist('instrumentos[]') or []
        pontuacoes = request.form.getlist('pontuacoes[]') or []

        # valida disciplina do professor
        if disciplinas_professor and disciplina not in disciplinas_professor:
            flash("Disciplina inválida. Selecione uma disciplina cadastrada no seu perfil.")
            cursor.close()
            conn.close()
            return redirect(url_for('registrar_planejamento'))

        if not disciplina or not turmas_ids or not any((c or "").strip() for c in conteudos):
            flash("Preencha disciplina, pelo menos uma turma e um conteúdo.")
            cursor.close()
            conn.close()
            return redirect(url_for('registrar_planejamento'))

        try:
            # cria cabeçalho
            cursor.execute("""
                INSERT INTO planejamentos
                (professor_id, disciplina, bimestre, ano, observacoes)
                VALUES (?, ?, ?, ?, ?)
            """, (professor_id, disciplina, bimestre, ano, observacoes))
            planejamento_id = cursor.lastrowid

            # vincula turmas
            for turma_id in turmas_ids:
                if turma_id and str(turma_id).strip():
                    cursor.execute("""
                        INSERT INTO planejamentos_turmas (planejamento_id, turma_id)
                        VALUES (?, ?)
                    """, (planejamento_id, int(turma_id)))

            # cria itens
            for i, conteudo in enumerate(conteudos):
                texto = (conteudo or "").strip()
                if not texto:
                    continue

                # 1) prioridade: datas do formulário novo
                di = (datas_inicio[i] if i < len(datas_inicio) else "") or ""
                df = (datas_fim[i] if i < len(datas_fim) else "") or ""

                data_inicio = di.strip() or None
                data_fim = df.strip() or None

                # 2) fallback: período antigo ("YYYY-MM-DD a YYYY-MM-DD")
                if (not data_inicio and not data_fim) and i < len(periodos):
                    periodo = (periodos[i] or "").strip()
                    partes = periodo.split(" a ")
                    data_inicio = partes[0].strip() if len(partes) > 0 and partes[0].strip() else None
                    data_fim = partes[1].strip() if len(partes) > 1 and partes[1].strip() else None

                forma_avaliacao = (instrumentos[i] if i < len(instrumentos) else None)
                forma_avaliacao = forma_avaliacao.strip() if isinstance(forma_avaliacao, str) else None

                pont_valor = None
                if i < len(pontuacoes):
                    p = (pontuacoes[i] or "").strip()
                    if p:
                        try:
                            pont_valor = float(p.replace(',', '.'))
                        except:
                            pont_valor = None

                cursor.execute("""
                    INSERT INTO planejamento_itens
                    (planejamento_id, conteudo, data_inicio, data_fim, forma_avaliacao, pontuacao_total)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (planejamento_id, texto, data_inicio, data_fim, forma_avaliacao, pont_valor))

            conn.commit()
            flash("Planejamento registrado com sucesso!")
            cursor.close()
            conn.close()
            return redirect(url_for('listar_planejamentos_professor'))

        except sqlite3.Error as e:
            conn.rollback()
            flash(f"Erro ao registrar planejamento: {e}")
            cursor.close()
            conn.close()
            return redirect(url_for('registrar_planejamento'))

    # GET
    cursor.close()
    conn.close()
    return render_template(
        'registrar_planejamento.html',
        ano_atual=ano_atual,
        bimestre_atual=bimestre_atual,
        turmas=turmas,
        disciplinas_professor=disciplinas_professor
    )


@app.route('/planejamento/professor')
def listar_planejamentos_professor():
    if 'usuario' not in session or session['tipo'] != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_id = obter_professor_id(session['usuario'])

    ano_letivo = (request.args.get('ano_letivo') or '').strip()
    bimestre = (request.args.get('bimestre') or '').strip()
    disciplina_filtro = (request.args.get('disciplina') or '').strip()
    turma_id = (request.args.get('turma_id') or '').strip()

    filtros = {
        'ano_letivo': ano_letivo,
        'bimestre': bimestre,
        'disciplina': disciplina_filtro,
        'turma_id': int(turma_id) if turma_id else None
    }

    conn = conectar_bd()
    cursor = conn.cursor()

    sql = '''
        SELECT
            p.id,
            p.disciplina,
            p.bimestre,
            p.ano            AS ano_letivo,
            p.criado_em,
            p.observacoes,
            GROUP_CONCAT(
                DISTINCT t.nome || ' (' || t.turno || ')'
            ) AS turmas_nomes,
            -- novos: conteúdos cadastrados
            GROUP_CONCAT(
                DISTINCT TRIM(pi.conteudo)
            ) AS conteudos_cadastrados,
            -- mantemos total_pontos só para quem ainda usa em outro lugar
            SUM(
                COALESCE(pi.pontuacao_total, 0)
            ) AS total_pontos
        FROM planejamentos p
        LEFT JOIN planejamentos_turmas pt
            ON pt.planejamento_id = p.id
        LEFT JOIN turmas t
            ON t.id = pt.turma_id
        LEFT JOIN planejamento_itens pi
            ON pi.planejamento_id = p.id
        WHERE p.professor_id = ?
    '''

    params = [professor_id]

    if ano_letivo:
        try:
            sql += ' AND p.ano = ?'
            params.append(int(ano_letivo))
        except ValueError:
            pass

    if bimestre:
        try:
            sql += ' AND p.bimestre = ?'
            params.append(int(bimestre))
        except ValueError:
            pass

    if disciplina_filtro:
        sql += ' AND p.disciplina = ?'
        params.append(disciplina_filtro)

    if turma_id:
        try:
            sql += ' AND t.id = ?'
            params.append(int(turma_id))
        except ValueError:
            pass

    sql += '''
        GROUP BY p.id
        ORDER BY p.ano DESC, p.bimestre, p.disciplina
    '''

    cursor.execute(sql, params)
    planejamentos = cursor.fetchall()

    cursor.execute(
        "SELECT DISTINCT disciplina FROM professor_disciplinas WHERE professor_id = ?",
        (professor_id,)
    )
    disciplinas_professor_rows = cursor.fetchall()
    disciplinas_professor = [row['disciplina'] for row in disciplinas_professor_rows]

    turmas = obter_turmas_professor(professor_id)
    turmas_sem_planejamento = []

    cursor.close()
    conn.close()

    return render_template(
        'planejamentos_professor.html',
        planejamentos=planejamentos,
        filtros=filtros,
        turmas=turmas,
        disciplinas_professor=disciplinas_professor,
        turmas_sem_planejamento=turmas_sem_planejamento
    )


@app.route('/planejamento/gestor')
def listar_planejamentos_gestor():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    # ---- Filtros vindos da URL (?professor_login=...&turma_id=...) ----
    professor_login = (request.args.get('professor_login') or '').strip()
    turma_id = (request.args.get('turma_id') or '').strip()
    disciplina_filtro = (request.args.get('disciplina') or '').strip()
    bimestre = (request.args.get('bimestre') or '').strip()

    filtros = {
        'professor_login': professor_login,
        'turma_id': turma_id,
        'disciplina': disciplina_filtro,
        'bimestre': bimestre
    }

    conn = conectar_bd()
    cursor = conn.cursor()

    # ---- Consulta principal com filtros dinâmicos ----
    sql = '''
        SELECT
            p.id,
            p.disciplina,
            p.bimestre,
            p.ano AS ano,
            p.criado_em,
            pr.login AS professor_login,
            GROUP_CONCAT(DISTINCT t.nome || ' (' || t.turno || ')') AS turmas_nomes,
            SUM(COALESCE(pi.pontuacao_total, 0)) AS total_pontos
        FROM planejamentos p
        JOIN professores pr ON p.professor_id = pr.id
        LEFT JOIN planejamentos_turmas pt
            ON pt.planejamento_id = p.id
        LEFT JOIN turmas t
            ON t.id = pt.turma_id
        LEFT JOIN planejamento_itens pi
            ON pi.planejamento_id = p.id
        WHERE 1=1
    '''
    params = []

    # Filtro por professor
    if professor_login:
        sql += ' AND pr.login = ?'
        params.append(professor_login)

    # Filtro por disciplina
    if disciplina_filtro:
        sql += ' AND p.disciplina = ?'
        params.append(disciplina_filtro)

    # Filtro por turma
    if turma_id:
        try:
            sql += ' AND t.id = ?'
            params.append(int(turma_id))
        except ValueError:
            pass

    # Filtro por bimestre
    if bimestre:
        try:
            sql += ' AND p.bimestre = ?'
            params.append(int(bimestre))
        except ValueError:
            pass

    sql += '''
        GROUP BY p.id
        ORDER BY p.ano DESC, p.bimestre, pr.login, p.disciplina
    '''

    cursor.execute(sql, params)
    planejamentos = cursor.fetchall()

    # ---- Listas para popular os filtros (combos) ----
    # Professores aprovados
    cursor.execute("""
        SELECT DISTINCT login
        FROM professores
        WHERE status = 'aprovado'
        ORDER BY login
    """)
    professores_rows = cursor.fetchall()
    professores = [row['login'] for row in professores_rows]

    # Turmas
    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()

    # Disciplinas (já usadas em planejamentos)
    cursor.execute("SELECT DISTINCT disciplina FROM planejamentos ORDER BY disciplina")
    disciplinas_rows = cursor.fetchall()
    disciplinas = [row['disciplina'] for row in disciplinas_rows]

    # Bimestres existentes
    cursor.execute("SELECT DISTINCT bimestre FROM planejamentos ORDER BY bimestre")
    bimestres_rows = cursor.fetchall()
    bimestres = [row['bimestre'] for row in bimestres_rows]

    cursor.close()
    conn.close()

    return render_template(
        'planejamentos_gestor.html',
        planejamentos=planejamentos,
        filtros=filtros,
        professores=professores,
        turmas=turmas,
        disciplinas=disciplinas,
        bimestres=bimestres
    )


@app.route('/planejamento/detalhar/<int:planejamento_id>')
def detalhar_planejamento(planejamento_id):
    # Garante que só usuário logado pode acessar
    if 'usuario' not in session:
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    # Busca os dados gerais do planejamento
    cursor.execute('''
        SELECT
            p.id,
            p.disciplina,
            p.bimestre,
            p.ano,
            p.criado_em,
            p.observacoes,
            pr.login AS professor_login
        FROM planejamentos p
        JOIN professores pr ON p.professor_id = pr.id
        WHERE p.id = ?
    ''', (planejamento_id,))
    planejamento = cursor.fetchone()

    if not planejamento:
        cursor.close()
        conn.close()
        flash("Planejamento não encontrado.")
        return redirect(url_for('listar_planejamentos_professor'))

    # Turmas vinculadas ao planejamento
    cursor.execute('''
        SELECT t.nome, t.turno
        FROM planejamentos_turmas pt
        JOIN turmas t ON pt.turma_id = t.id
        WHERE pt.planejamento_id = ?
        ORDER BY t.nome
    ''', (planejamento_id,))
    turmas = cursor.fetchall()

    # Itens do planejamento com formatação de datas
    cursor.execute('''
        SELECT
            id,
            conteudo,
            data_inicio,
            data_fim,
            COALESCE(concluido, 0) AS concluido
        FROM planejamento_itens
        WHERE planejamento_id = ?
        ORDER BY id
    ''', (planejamento_id,))
    itens_raw = cursor.fetchall()

    cursor.close()
    conn.close()

    # Formatar datas para exibição (DD/MM/YYYY)
    itens = []
    for item in itens_raw:
        item_dict = dict(item)

        # Formatar data_inicio
        if item_dict['data_inicio']:
            try:
                dt = datetime.strptime(item_dict['data_inicio'], '%Y-%m-%d')
                item_dict['data_inicio_formatada'] = dt.strftime('%d/%m/%Y')
            except:
                item_dict['data_inicio_formatada'] = item_dict['data_inicio']
        else:
            item_dict['data_inicio_formatada'] = None

        # Formatar data_fim
        if item_dict['data_fim']:
            try:
                dt = datetime.strptime(item_dict['data_fim'], '%Y-%m-%d')
                item_dict['data_fim_formatada'] = dt.strftime('%d/%m/%Y')
            except:
                item_dict['data_fim_formatada'] = item_dict['data_fim']
        else:
            item_dict['data_fim_formatada'] = None

        itens.append(item_dict)

    return render_template(
        'detalhar_planejamento.html',
        planejamento=planejamento,
        turmas=turmas,
        itens=itens
    )



@app.route('/planejamento/item/<int:item_id>/marcar', methods=['POST'])
def marcar_item_planejamento(item_id):
    # Apenas professor pode marcar/desmarcar item como concluído
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        # Descobre a qual planejamento esse item pertence e o status atual
        cursor.execute('''
            SELECT
                planejamento_id,
                COALESCE(concluido, 0) AS concluido
            FROM planejamento_itens
            WHERE id = ?
        ''', (item_id,))
        item = cursor.fetchone()

        if not item:
            cursor.close()
            conn.close()
            flash("Item do planejamento não encontrado.")
            return redirect(url_for('listar_planejamentos_professor'))

        planejamento_id = item['planejamento_id']
        status_atual = item['concluido']

        # Alterna status: 0 -> 1, 1 -> 0
        novo_status = 0 if status_atual else 1

        cursor.execute('''
            UPDATE planejamento_itens
            SET concluido = ?
            WHERE id = ?
        ''', (novo_status, item_id))
        conn.commit()

        flash("Status do item atualizado com sucesso.")

    except sqlite3.Error as e:
        conn.rollback()
        flash(f"Erro ao atualizar status do item: {e}")
    finally:
        cursor.close()
        conn.close()

    # Volta para a tela de detalhes do planejamento
    return redirect(url_for('detalhar_planejamento', planejamento_id=planejamento_id))


# EDITAR PLANEJAMENTO
@app.route('/planejamento/editar/<int:planejamento_id>', methods=['POST'])
def atualizar_planejamento(planejamento_id):
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_id = obter_professor_id(session['usuario'])
    if not professor_id:
        flash("Professor não encontrado.")
        return redirect(url_for('dashboard_professor'))

    disciplina = request.form.get('disciplina')
    ano_letivo = request.form.get('ano_letivo')
    bimestre = request.form.get('bimestre')
    observacoes = (request.form.get('observacoes_gerais') or '').strip()
    turmas_ids_form = request.form.getlist('turmas[]') or []
    item_ids = request.form.getlist('item_ids[]') or []
    conteudos = request.form.getlist('conteudos[]') or []
    datas_inicio = request.form.getlist('datas_inicio[]') or []
    datas_fim = request.form.getlist('datas_fim[]') or []

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            SELECT id FROM planejamentos
            WHERE id = ? AND professor_id = ?
        """, (planejamento_id, professor_id))
        existe = cursor.fetchone()
        if not existe:
            flash("Planejamento não encontrado ou não pertence a você.")
            cursor.close()
            conn.close()
            return redirect(url_for('listar_planejamentos_professor'))

        cursor.execute("""
            UPDATE planejamentos
            SET disciplina = ?, ano = ?, bimestre = ?, observacoes = ?
            WHERE id = ? AND professor_id = ?
        """, (disciplina, int(ano_letivo), int(bimestre), observacoes,
              planejamento_id, professor_id))

        cursor.execute(
            "DELETE FROM planejamentos_turmas WHERE planejamento_id = ?",
            (planejamento_id,)
        )

        for turma_id in turmas_ids_form:
            if turma_id and turma_id.strip():
                cursor.execute("""
                    INSERT OR IGNORE INTO planejamentos_turmas
                    (planejamento_id, turma_id)
                    VALUES (?, ?)
                """, (planejamento_id, int(turma_id)))

        # Primeiro, remove itens que não estão mais no formulário
        ids_existentes = [i for i in item_ids if i and i.strip()]
        if ids_existentes:
            placeholders = ','.join(['?'] * len(ids_existentes))
            cursor.execute(f"""
                DELETE FROM planejamento_itens
                WHERE planejamento_id = ? AND id NOT IN ({placeholders})
            """, [planejamento_id] + ids_existentes)

        # Atualiza/insere itens
        for idx, conteudo in enumerate(conteudos):
            if not conteudo or not conteudo.strip():
                continue

            item_id = item_ids[idx].strip() if idx < len(item_ids) and item_ids[idx].strip() else None
            data_inicio = (datas_inicio[idx] if idx < len(datas_inicio) else '').strip() or None
            data_fim = (datas_fim[idx] if idx < len(datas_fim) else '').strip() or None

            if item_id:
                # Atualiza item existente
                cursor.execute("""
                    UPDATE planejamento_itens
                    SET conteudo = ?, data_inicio = ?, data_fim = ?
                    WHERE id = ? AND planejamento_id = ?
                """, (conteudo.strip(), data_inicio, data_fim, int(item_id), planejamento_id))
            else:
                # Insere novo item
                cursor.execute("""
                    INSERT INTO planejamento_itens
                    (planejamento_id, conteudo, data_inicio, data_fim)
                    VALUES (?, ?, ?, ?)
                """, (planejamento_id, conteudo.strip(), data_inicio, data_fim))

        conn.commit()
        flash("Planejamento atualizado com sucesso!")

    except Exception as e:
        conn.rollback()
        flash(f"Erro ao atualizar planejamento: {e}")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('detalhar_planejamento', planejamento_id=planejamento_id))


# EXCLUIR PLANEJAMENTO
@app.route('/planejamento/excluir/<int:planejamento_id>', methods=['POST'])
def excluir_planejamento(planejamento_id):
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_id = obter_professor_id(session['usuario'])
    if not professor_id:
        flash("Professor não encontrado.")
        return redirect(url_for('dashboard_professor'))

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            SELECT id FROM planejamentos
            WHERE id = ? AND professor_id = ?
        """, (planejamento_id, professor_id))
        planejamento = cursor.fetchone()

        if not planejamento:
            cursor.close()
            conn.close()
            flash("Planejamento não encontrado ou não pertence a você.", "warning")
            return redirect(url_for('listar_planejamentos_professor'))

        cursor.execute(
            "DELETE FROM planejamento_itens WHERE planejamento_id = ?",
            (planejamento_id,)
        )

        cursor.execute(
            "DELETE FROM planejamentos_turmas WHERE planejamento_id = ?",
            (planejamento_id,)
        )

        cursor.execute(
            "DELETE FROM planejamentos WHERE id = ? AND professor_id = ?",
            (planejamento_id, professor_id)
        )

        conn.commit()
        flash("Planejamento removido com sucesso!", "success")
    except sqlite3.Error as e:
        conn.rollback()
        flash(f"Erro ao excluir planejamento: {e}", "danger")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('listar_planejamentos_professor'))

@app.route('/planejamento/editar/<int:planejamento_id>/form', methods=['GET'])
def editar_planejamento_form(planejamento_id):
    """Exibe formulário para editar planejamento existente"""
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_id = obter_professor_id(session['usuario'])
    if not professor_id:
        flash("Professor não encontrado.")
        return redirect(url_for('dashboard_professor'))

    conn = conectar_bd()
    cursor = conn.cursor()

    # Busca o planejamento
    cursor.execute("""
        SELECT id, disciplina, bimestre, ano, observacoes
        FROM planejamentos
        WHERE id = ? AND professor_id = ?
    """, (planejamento_id, professor_id))
    planejamento = cursor.fetchone()

    if not planejamento:
        cursor.close()
        conn.close()
        flash("Planejamento não encontrado ou você não tem permissão.")
        return redirect(url_for('listar_planejamentos_professor'))

    # Busca turmas do planejamento
    cursor.execute("""
        SELECT turma_id
        FROM planejamentos_turmas
        WHERE planejamento_id = ?
    """, (planejamento_id,))
    turmas_selecionadas = [row['turma_id'] for row in cursor.fetchall()]

    # Busca todas as turmas disponíveis
    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    todas_turmas = cursor.fetchall()

    # Busca itens do planejamento
    cursor.execute("""
        SELECT id, conteudo, data_inicio, data_fim
        FROM planejamento_itens
        WHERE planejamento_id = ?
        ORDER BY id
    """, (planejamento_id,))
    itens = cursor.fetchall()

    # Disciplinas do professor
    cursor.execute("""
        SELECT disciplina
        FROM professor_disciplinas
        WHERE professor_id = ?
        ORDER BY disciplina
    """, (professor_id,))
    disciplinas_rows = cursor.fetchall()
    disciplinas_professor = [row['disciplina'] for row in disciplinas_rows] if disciplinas_rows else []

    cursor.close()
    conn.close()

    return render_template(
        'editar_planejamento.html',
        planejamento=planejamento,
        turmas=todas_turmas,
        turmas_selecionadas=turmas_selecionadas,
        itens=itens,
        disciplinas_professor=disciplinas_professor
    )

# CONTEÚDOS DAS AVALIAÇÕES BIMESTRAIS

@app.route('/avaliacoes/registrar', methods=['GET', 'POST'])
def registrar_avaliacao_bimestral():
    if 'usuario' not in session or session['tipo'] != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_id = obter_professor_id(session['usuario'])
    if not professor_id:
        flash("Professor não encontrado.")
        return redirect(url_for('dashboard_professor'))

    hoje = datetime.now()
    ano_atual = hoje.year
    bimestre_atual = ((hoje.month - 1) // 3) + 1

    # ------------------------ GET (base) + lista de disciplinas do professor ------------------------
    # (vamos buscar antes, porque precisamos usar no GET e também validar no POST)
    conn = conectar_bd()
    cursor = conn.cursor()

    # ✅ Disciplinas cadastradas pelo professor (para preencher o select)
    cursor.execute("""
        SELECT disciplina
        FROM professor_disciplinas
        WHERE professor_id = ?
        ORDER BY disciplina
    """, (professor_id,))
    disciplinas_rows = cursor.fetchall()
    disciplinas_professor = [r["disciplina"] for r in disciplinas_rows] if disciplinas_rows else []

    # Fallback: se por algum motivo o professor ainda não tiver disciplina cadastrada,
    # você pode manter DISCIPLINAS como plano B (evita travar a tela).
    if not disciplinas_professor:
        disciplinas_professor = DISCIPLINAS

    cursor.close()
    conn.close()

    # ------------------------ POST: salvar avaliações ------------------------
    if request.method == 'POST':
        disciplina = request.form.get('disciplina')
        turmas_ids = request.form.getlist('turmas_ids')
        ano = request.form.get('ano') or str(ano_atual)
        bimestre = request.form.get('bimestre') or str(bimestre_atual)

        tipos = request.form.getlist('tipo_avaliacao')
        descricoes = request.form.getlist('descricao_avaliacao')
        conteudos_list = request.form.getlist('conteudos')
        datas = request.form.getlist('data_avaliacao')
        pontuacoes = request.form.getlist('pontuacao')

        if not disciplina or not turmas_ids:
            flash("Selecione a disciplina e pelo menos uma turma.")
            return redirect(url_for('registrar_avaliacao_bimestral'))

        # ✅ Validação: disciplina precisa pertencer ao professor
        if disciplina not in disciplinas_professor:
            flash("Disciplina inválida para o seu cadastro. Selecione uma disciplina da sua lista.")
            return redirect(url_for('registrar_avaliacao_bimestral'))

        tem_conteudo = any((c or "").strip() for c in conteudos_list)
        if not tem_conteudo:
            flash("Preencha ao menos um bloco de avaliação com conteúdos.")
            return redirect(url_for('registrar_avaliacao_bimestral'))

        conn = conectar_bd()
        cursor = conn.cursor()

        try:
            for i, conteudo in enumerate(conteudos_list):
                texto = (conteudo or "").strip()
                if not texto:
                    continue

                tipo_i = (tipos[i] if i < len(tipos) else None) or None
                desc_i = (descricoes[i] if i < len(descricoes) else None) or None
                data_i = (datas[i] if i < len(datas) and datas[i] else None)

                pont_str = pontuacoes[i] if i < len(pontuacoes) else ""
                if pont_str:
                    try:
                        pont_valor = float(pont_str.replace(',', '.'))
                    except ValueError:
                        pont_valor = None
                else:
                    pont_valor = None

                for turma_id in turmas_ids:
                    cursor.execute(
                        """
                        INSERT INTO avaliacoes_bimestrais
                        (professor_id, disciplina, turma_id, bimestre, ano,
                         tipo_avaliacao, descricao_avaliacao, conteudos,
                         data_avaliacao, pontuacao)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            professor_id,
                            disciplina,
                            int(turma_id),
                            int(bimestre),
                            int(ano),
                            tipo_i,
                            desc_i,
                            texto,
                            data_i,
                            pont_valor,
                        ),
                    )

            conn.commit()
            flash("Avaliações registradas com sucesso!")
        except sqlite3.Error as e:
            conn.rollback()
            flash(f"Erro ao salvar avaliações: {e}")
        finally:
            cursor.close()
            conn.close()

        return redirect(url_for('listar_avaliacoes_professor'))

    # ------------------------ GET: montar formulário ------------------------
    conn = conectar_bd()
    cursor = conn.cursor()

    # Turmas para o professor (ou todas, como você estava usando)
    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()

    # SQL corrigido (mantido como você trouxe)
    cursor.execute(
        """
        SELECT
            pi.conteudo,
            p.disciplina,
            p.bimestre,
            p.ano,
            GROUP_CONCAT(DISTINCT t.nome || ' (' || t.turno || ')') AS turmas
        FROM planejamento_itens pi
        JOIN planejamentos p ON pi.planejamento_id = p.id
        LEFT JOIN planejamentos_turmas pt ON pt.planejamento_id = p.id
        LEFT JOIN turmas t ON t.id = pt.turma_id
        WHERE p.professor_id = ?
          AND p.ano = ?
        GROUP BY pi.id, pi.conteudo, p.disciplina, p.bimestre, p.ano
        ORDER BY p.ano DESC, p.bimestre, p.disciplina, pi.id
        """,
        (professor_id, ano_atual),
    )
    conteudos_planejados = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        "registrar_avaliacao.html",
        turmas=turmas,
        disciplinas=disciplinas_professor,  # ✅ AGORA SÓ AS DISCIPLINAS DO PROFESSOR
        ano_atual=ano_atual,
        bimestre_atual=bimestre_atual,
        conteudos_planejados=conteudos_planejados,
    )


@app.route('/avaliacoes/excluir/<int:avaliacao_id>', methods=['POST'])
def excluir_avaliacao_bimestral(avaliacao_id):
    # Garante que é professor logado
    if 'usuario' not in session or session.get('tipo') != 'professor':
        flash("Acesso não autorizado.", "danger")
        return redirect(url_for('login'))

    professor_id = obter_professor_id(session['usuario'])
    if not professor_id:
        flash("Professor não encontrado.", "danger")
        return redirect(url_for('dashboard_professor'))

    conn = conectar_bd()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            SELECT id FROM avaliacoes_bimestrais
            WHERE id = ? AND professor_id = ?
        """, (avaliacao_id, professor_id))
        avaliacao = cursor.fetchone()

        if not avaliacao:
            flash("Avaliação não encontrada ou você não tem permissão para excluí-la.", "warning")
            cursor.close()
            conn.close()
            return redirect(url_for('listar_avaliacoes_professor'))

        cursor.execute(
            "DELETE FROM avaliacoes_bimestrais WHERE id = ? AND professor_id = ?",
            (avaliacao_id, professor_id)
        )

        conn.commit()
        flash("Avaliação excluída com sucesso!", "success")

    except Exception as e:
        conn.rollback()
        flash(f"Erro ao excluir avaliação: {str(e)}", "danger")
    finally:
        cursor.close()
        conn.close()

    return redirect(url_for('listar_avaliacoes_professor'))


@app.route('/avaliacoes/professor')
def listar_avaliacoes_professor():
    if 'usuario' not in session or session['tipo'] != 'professor':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    professor_id = obter_professor_id(session['usuario'])
    if not professor_id:
        flash("Professor não encontrado.")
        return redirect(url_for('dashboard_professor'))

    # --------- FILTROS (GET) ----------
    ano = (request.args.get('ano') or '').strip()
    bimestre = (request.args.get('bimestre') or '').strip()
    disciplina = (request.args.get('disciplina') or '').strip()
    turma_id = (request.args.get('turma_id') or '').strip()

    filtros = {
        'ano': ano,
        'bimestre': bimestre,
        'disciplina': disciplina,
        'turma_id': turma_id
    }

    conn = conectar_bd()
    cursor = conn.cursor()

    # --------- CONSULTA PRINCIPAL ----------
    sql = '''
        SELECT a.id,
               a.disciplina,
               a.bimestre,
               a.ano,
               a.tipo_avaliacao,
               a.descricao_avaliacao,
               a.conteudos,
               a.data_avaliacao,
               a.pontuacao,
               t.id   AS turma_id,
               t.nome AS turma_nome
        FROM avaliacoes_bimestrais a
        JOIN turmas t ON a.turma_id = t.id
        WHERE a.professor_id = ?
    '''
    params = [professor_id]

    if ano:
        try:
            sql += ' AND a.ano = ?'
            params.append(int(ano))
        except ValueError:
            pass

    if bimestre:
        try:
            sql += ' AND a.bimestre = ?'
            params.append(int(bimestre))
        except ValueError:
            pass

    if disciplina:
        sql += ' AND a.disciplina = ?'
        params.append(disciplina)

    if turma_id:
        try:
            sql += ' AND t.id = ?'
            params.append(int(turma_id))
        except ValueError:
            pass

    sql += '''
        ORDER BY a.ano DESC, a.bimestre, t.nome, a.disciplina
    '''

    cursor.execute(sql, params)
    rows = cursor.fetchall()

    # --------- LISTAS PARA OS SELECTS ----------
    # Anos que esse professor já usou em avaliações
    cursor.execute("""
        SELECT DISTINCT ano
        FROM avaliacoes_bimestrais
        WHERE professor_id = ?
        ORDER BY ano DESC
    """, (professor_id,))
    anos_rows = cursor.fetchall()
    anos = [r['ano'] for r in anos_rows]

    # Disciplinas que esse professor já usou em avaliações
    cursor.execute("""
        SELECT DISTINCT disciplina
        FROM avaliacoes_bimestrais
        WHERE professor_id = ?
        ORDER BY disciplina
    """, (professor_id,))
    disc_rows = cursor.fetchall()
    disciplinas = [r['disciplina'] for r in disc_rows]

    cursor.close()
    conn.close()

    # Turmas vinculadas ao professor (função auxiliar já existente)
    turmas = obter_turmas_professor(professor_id)

    # --------- FORMATAÇÃO DA LISTA PARA O TEMPLATE ----------
    avaliacoes = []
    for row in rows:
        d = dict(row)
        data_str = d.get('data_avaliacao')

        if data_str:
            try:
                dt = datetime.strptime(data_str, '%Y-%m-%d')
                d['data_avaliacao'] = dt.strftime('%d/%m/%Y')
            except ValueError:
                # Mantém formato original se não estiver no padrão esperado
                pass

        avaliacoes.append(d)

    return render_template(
        'avaliacoes_professor.html',
        avaliacoes=avaliacoes,
        filtros=filtros,
        anos=anos,
        disciplinas=disciplinas,
        turmas=turmas
    )


@app.route('/avaliacoes/visualizar/<int:id>')
def visualizar_avaliacao(id):
    """
    Rota usada pelo PROFESSOR (e também aceita moderador, se acessar por aqui).
    Professor só vê avaliações dele; moderador pode ver qualquer uma.
    Renderiza o template visualizar_avaliacao.html (versão do professor).
    """
    if 'usuario' not in session or session.get('tipo') not in ('professor', 'moderador'):
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    tipo_usuario = session.get('tipo')
    conn = conectar_bd()
    cursor = conn.cursor()

    sql = '''
        SELECT
            a.id,
            a.disciplina,
            a.bimestre,
            a.ano,
            a.tipo_avaliacao,
            a.descricao_avaliacao,
            a.conteudos,
            a.data_avaliacao,
            a.pontuacao,
            t.nome AS turma_nome,
            t.turno AS turma_turno,
            p.login AS professor_login
        FROM avaliacoes_bimestrais a
        JOIN turmas t ON a.turma_id = t.id
        JOIN professores p ON a.professor_id = p.id
        WHERE a.id = ?
    '''
    params = [id]

    # Se for professor, restringe à autoria
    if tipo_usuario == 'professor':
        professor_id = obter_professor_id(session['usuario'])
        if not professor_id:
            cursor.close()
            conn.close()
            flash("Professor não encontrado.")
            return redirect(url_for('dashboard_professor'))

        sql += ' AND a.professor_id = ?'
        params.append(professor_id)

    cursor.execute(sql, params)
    row = cursor.fetchone()
    cursor.close()
    conn.close()

    if not row:
        flash("Avaliação não encontrada ou não autorizada.")
        if tipo_usuario == 'professor':
            return redirect(url_for('listar_avaliacoes_professor'))
        else:
            return redirect(url_for('listar_avaliacoes_gestor'))

    avaliacao = dict(row)

    # Formata data
    data_str = avaliacao.get('data_avaliacao')
    if data_str:
        try:
            dt = datetime.strptime(data_str, '%Y-%m-%d')
            avaliacao['data_avaliacao_formatada'] = dt.strftime('%d/%m/%Y')
        except ValueError:
            avaliacao['data_avaliacao_formatada'] = data_str
    else:
        avaliacao['data_avaliacao_formatada'] = 'Não informada'

    # Formata pontuação
    pont = avaliacao.get('pontuacao')
    if pont is not None:
        try:
            avaliacao['pontuacao_formatada'] = f"{float(pont):.1f}".replace('.', ',') + " pontos"
        except Exception:
            avaliacao['pontuacao_formatada'] = str(pont)
    else:
        avaliacao['pontuacao_formatada'] = 'Não informada'

    # Texto do bimestre
    avaliacao['bimestre_formatado'] = f"{avaliacao.get('bimestre', '')}º bimestre / {avaliacao.get('ano', '')}"

    return render_template('visualizar_avaliacao.html', avaliacao=avaliacao)


@app.route('/avaliacoes/gestor')
def listar_avaliacoes_gestor():
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    # -----------------------------
    # LER FILTROS DA URL (?professor_id=... etc.)
    # -----------------------------
    professor_id_raw = (request.args.get('professor_id') or '').strip()
    turma_id_raw = (request.args.get('turma_id') or '').strip()
    disciplina_filtro = (request.args.get('disciplina') or '').strip()
    bimestre_raw = (request.args.get('bimestre') or '').strip()

    professor_id = None
    turma_id = None
    bimestre = None

    if professor_id_raw:
        try:
            professor_id = int(professor_id_raw)
        except ValueError:
            professor_id = None

    if turma_id_raw:
        try:
            turma_id = int(turma_id_raw)
        except ValueError:
            turma_id = None

    if bimestre_raw:
        try:
            bimestre = int(bimestre_raw)
        except ValueError:
            bimestre = None

    filtros = {
        'professor_id': professor_id,
        'turma_id': turma_id,
        'disciplina': disciplina_filtro,
        'bimestre': bimestre
    }

    conn = conectar_bd()
    cursor = conn.cursor()

    # -----------------------------
    # BUSCAR LISTAS PARA OS SELECTS
    # -----------------------------
    cursor.execute("""
        SELECT id, login
        FROM professores
        WHERE status = 'aprovado'
        ORDER BY login
    """)
    professores = cursor.fetchall()

    cursor.execute("""
        SELECT id, nome
        FROM turmas
        ORDER BY nome
    """)
    turmas = cursor.fetchall()

    cursor.execute("""
        SELECT DISTINCT disciplina
        FROM avaliacoes_bimestrais
        ORDER BY disciplina
    """)
    disciplinas_rows = cursor.fetchall()
    disciplinas = [row['disciplina'] for row in disciplinas_rows]

    # -----------------------------
    # MONTAR CONSULTA PRINCIPAL COM FILTROS
    # -----------------------------
    sql = '''
        SELECT
            a.id,
            a.disciplina,
            a.bimestre,
            a.ano,
            a.data_avaliacao,
            t.nome AS turma_nome,
            p.login AS professor_login
        FROM avaliacoes_bimestrais a
        JOIN turmas t      ON a.turma_id    = t.id
        JOIN professores p ON a.professor_id = p.id
        WHERE 1 = 1
    '''
    params = []

    if professor_id is not None:
        sql += " AND p.id = ?"
        params.append(professor_id)

    if turma_id is not None:
        sql += " AND t.id = ?"
        params.append(turma_id)

    if disciplina_filtro:
        sql += " AND a.disciplina = ?"
        params.append(disciplina_filtro)

    if bimestre is not None:
        sql += " AND a.bimestre = ?"
        params.append(bimestre)

    sql += " ORDER BY a.ano DESC, a.bimestre, t.nome, p.login, a.disciplina"

    cursor.execute(sql, params)
    rows = cursor.fetchall()

    cursor.close()
    conn.close()

    # -----------------------------
    # FORMATAR DATAS PARA DD/MM/AAAA
    # -----------------------------
    avaliacoes = []
    for row in rows:
        d = dict(row)
        data_str = d.get('data_avaliacao')

        if data_str:
            try:
                dt = datetime.strptime(data_str, '%Y-%m-%d')
                d['data_avaliacao_formatada'] = dt.strftime('%d/%m/%Y')
            except ValueError:
                d['data_avaliacao_formatada'] = data_str
        else:
            d['data_avaliacao_formatada'] = 'Não informada'

        avaliacoes.append(d)

    return render_template(
        'avaliacoes_gestor.html',
        avaliacoes=avaliacoes,
        professores=professores,
        turmas=turmas,
        disciplinas=disciplinas,
        filtros=filtros
    )


@app.route('/avaliacoes/gestor/visualizar/<int:id>')
def visualizar_avaliacao_gestor(id):
    """
    Detalhe da avaliação para o GESTOR.
    Usa o template visualizar_avaliacao_gestor.html
    (linkado a partir de avaliacoes_gestor.html)
    """
    if 'usuario' not in session or session['tipo'] != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    cursor = conn.cursor()

    cursor.execute('''
        SELECT
            a.id,
            a.disciplina,
            a.bimestre,
            a.ano,
            a.tipo_avaliacao,
            a.descricao_avaliacao,
            a.conteudos,
            a.data_avaliacao,
            a.pontuacao,
            t.nome AS turma_nome,
            t.turno AS turma_turno,
            p.login AS professor_login
        FROM avaliacoes_bimestrais a
        JOIN turmas t ON a.turma_id = t.id
        JOIN professores p ON a.professor_id = p.id
        WHERE a.id = ?
    ''', (id,))
    row = cursor.fetchone()
    cursor.close()
    conn.close()

    if not row:
        flash("Avaliação não encontrada.")
        return redirect(url_for('listar_avaliacoes_gestor'))

    avaliacao = dict(row)

    # Formata data da avaliação
    data_str = avaliacao.get('data_avaliacao')
    if data_str:
        try:
            dt = datetime.strptime(data_str, '%Y-%m-%d')
            avaliacao['data_avaliacao_formatada'] = dt.strftime('%d/%m/%Y')
        except ValueError:
            avaliacao['data_avaliacao_formatada'] = data_str
    else:
        avaliacao['data_avaliacao_formatada'] = 'Não informada'

    # Formata pontuação
    pont = avaliacao.get('pontuacao')
    if pont is not None:
        try:
            avaliacao['pontuacao_formatada'] = f"{float(pont):.1f}".replace('.', ',') + " pontos"
        except Exception:
            avaliacao['pontuacao_formatada'] = str(pont)
    else:
        avaliacao['pontuacao_formatada'] = 'Não informada'

    # Texto do bimestre
    avaliacao['bimestre_formatado'] = f"{avaliacao.get('bimestre', '')}º bimestre / {avaliacao.get('ano', '')}"

    return render_template('visualizar_avaliacao_gestor.html', avaliacao=avaliacao)


# ----------------- ATENDIMENTOS A RESPONSÁVEIS (MODERADOR) ----------------- #

@app.route('/atendimentos/novo', methods=['GET', 'POST'])
def atendimentos_novo():
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    # Turnos disponíveis a partir das turmas cadastradas
    conn = conectar_bd()
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT turno FROM turmas ORDER BY turno")
    turnos = [r['turno'] for r in cursor.fetchall() if r['turno']]
    cursor.close()
    conn.close()

    if request.method == 'POST':
        protocolo = (request.form.get('protocolo') or '').strip() or None
        turma_id = (request.form.get('turma_id') or '').strip()
        aluno_id = (request.form.get('aluno_id') or '').strip()

        # Agora o responsável é digitado (não é mais seletor)
        responsavel_nome = (request.form.get('responsavel_nome') or '').strip()
        responsavel_parentesco = (request.form.get('responsavel_parentesco') or '').strip()

        # Quem registrou
        registrador_nome = (request.form.get('registrador_nome') or '').strip()
        registrador_cargo = (request.form.get('registrador_cargo') or '').strip()

        envolve_professor = 1 if request.form.get('envolve_professor') else 0
        professor_nome = (request.form.get('professor_nome') or '').strip() or None

        data_atendimento = (request.form.get('data_atendimento') or '').strip()
        hora_atendimento = (request.form.get('hora_atendimento') or '').strip() or None
        assunto = (request.form.get('assunto') or '').strip() or None

        relato = (request.form.get('relato') or '').strip()
        combinados = (request.form.get('combinados') or '').strip() or None

        retorno_previsto = 1 if request.form.get('retorno_previsto') else 0
        retorno_em = (request.form.get('retorno_em') or '').strip() or None

        reuniao_agendada = 1 if request.form.get('reuniao_agendada') else 0
        reuniao_data = (request.form.get('reuniao_data') or '').strip() or None

        # Descobre turno da turma (para registrar)
        turno_turma = None
        try:
            if turma_id:
                conn = conectar_bd()
                cur = conn.cursor()
                cur.execute("SELECT turno FROM turmas WHERE id = ?", (turma_id,))
                row = cur.fetchone()
                cur.close()
                conn.close()
                turno_turma = row['turno'] if row else None
        except Exception:
            turno_turma = None

        # Validações mínimas
        if not turma_id or not aluno_id or not data_atendimento or not relato:
            flash("Preencha os campos obrigatórios: Turma, Aluno(a), Data do atendimento e Relato.")
            return redirect(url_for('atendimentos_novo'))

        if not responsavel_nome:
            flash("Informe o nome do(a) responsável presente no atendimento.")
            return redirect(url_for('atendimentos_novo'))

        if not responsavel_parentesco:
            flash("Informe o parentesco do(a) responsável com o(a) estudante.")
            return redirect(url_for('atendimentos_novo'))

        if not registrador_nome or not registrador_cargo:
            flash("Informe seu nome e cargo/função (quem registrou o atendimento).")
            return redirect(url_for('atendimentos_novo'))

        # Se marcou envolve professor, exige nome
        if envolve_professor and not professor_nome:
            flash("Informe o nome do(a) professor(a) envolvido(a).")
            return redirect(url_for('atendimentos_novo'))

        # Se marcou retorno, exige retorno_em
        if retorno_previsto and not retorno_em:
            flash("Informe a data/hora do retorno.")
            return redirect(url_for('atendimentos_novo'))

        # Se marcou reunião, exige reuniao_data
        if reuniao_agendada and not reuniao_data:
            flash("Informe a data/hora da reunião.")
            return redirect(url_for('atendimentos_novo'))

        # Se protocolo foi preenchido, checa unicidade
        if protocolo:
            conn = conectar_bd()
            cur = conn.cursor()
            cur.execute("SELECT 1 FROM atendimentos_responsaveis WHERE protocolo = ? LIMIT 1", (protocolo,))
            exists = cur.fetchone() is not None
            cur.close()
            conn.close()
            if exists:
                flash("Já existe um atendimento com esse protocolo. Gere outro ou deixe em branco.")
                return redirect(url_for('atendimentos_novo'))

        try:
            conn = conectar_bd()
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO atendimentos_responsaveis (
                    protocolo, turma_id, turno, aluno_id,
                    responsavel_nome, responsavel_parentesco, registrador_nome, registrador_cargo,
                    envolve_professor, professor_nome,
                    data_atendimento, hora_atendimento, assunto,
                    relato, combinados,
                    retorno_previsto, retorno_em,
                    reuniao_agendada, reuniao_data,
                    criado_em
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?, datetime('now'))
            """, (
                protocolo, turma_id, turno_turma, aluno_id,
                responsavel_nome, responsavel_parentesco, registrador_nome, registrador_cargo,
                envolve_professor, professor_nome,
                data_atendimento, hora_atendimento, assunto,
                relato, combinados,
                retorno_previsto, retorno_em,
                reuniao_agendada, reuniao_data
            ))
            conn.commit()
            cur.close()
            conn.close()
            flash("Atendimento salvo com sucesso.")
            return redirect(url_for('atendimentos_historico'))
        except sqlite3.Error as e:
            flash(f"Erro ao salvar atendimento: {e}")
            return redirect(url_for('atendimentos_novo'))

    return render_template('atendimentos_novo.html', turnos=turnos)


@app.route('/atendimentos/historico', methods=['GET'])
def atendimentos_historico():
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    # filtros
    protocolo = (request.args.get('protocolo') or '').strip()
    turno = (request.args.get('turno') or '').strip()
    turma_id = (request.args.get('turma_id') or '').strip()
    aluno_id = (request.args.get('aluno_id') or '').strip()
    data_ini = (request.args.get('data_ini') or '').strip()
    data_fim = (request.args.get('data_fim') or '').strip()

    conn = conectar_bd()
    cursor = conn.cursor()

    # combos
    cursor.execute("SELECT DISTINCT turno FROM turmas ORDER BY turno")
    turnos = [r['turno'] for r in cursor.fetchall() if r['turno']]
    cursor.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cursor.fetchall()

    alunos = []
    if turma_id:
        cursor.execute("SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome", (turma_id,))
        alunos = cursor.fetchall()

    # monta query
    where = []
    params = []

    if protocolo:
        where.append("a.protocolo LIKE ?")
        params.append(f"%{protocolo}%")
    if turno:
        where.append("a.turno = ?")
        params.append(turno)
    if turma_id:
        where.append("a.turma_id = ?")
        params.append(turma_id)
    if aluno_id:
        where.append("a.aluno_id = ?")
        params.append(aluno_id)
    if data_ini:
        where.append("date(a.data_atendimento) >= date(?)")
        params.append(data_ini)
    if data_fim:
        where.append("date(a.data_atendimento) <= date(?)")
        params.append(data_fim)

    sql = '''
        SELECT
            a.*,
            t.nome AS turma_nome,
            t.turno AS turma_turno,
            al.nome AS aluno_nome,
            COALESCE(NULLIF(a.responsavel_nome,''), r.login) AS responsavel_nome_exibicao
        FROM atendimentos_responsaveis a
        JOIN turmas t ON a.turma_id = t.id
        JOIN alunos al ON a.aluno_id = al.id
        LEFT JOIN responsaveis r ON a.responsavel_id = r.id
    '''
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY a.data_atendimento DESC, a.criado_em DESC"

    cursor.execute(sql, tuple(params))
    atendimentos = cursor.fetchall()

    cursor.close()
    conn.close()

    filtros = {
        "protocolo": protocolo,
        "turno": turno,
        "turma_id": turma_id,
        "aluno_id": aluno_id,
        "data_ini": data_ini,
        "data_fim": data_fim,
    }

    return render_template(
        'atendimentos_historico.html',
        atendimentos=atendimentos,
        turnos=turnos,
        turmas=turmas,
        alunos=alunos,
        filtros=filtros
    )


# ----------------- APIs (para o formulário dinâmico) ----------------- #

@app.get('/api/atendimentos/check_protocolo')
def api_atendimentos_check_protocolo():
    """Retorna {ok: True} se o protocolo NÃO existe."""
    protocolo = (request.args.get('protocolo') or '').strip()
    if not protocolo:
        return jsonify({'ok': False, 'error': 'protocolo_vazio'}), 400

    try:
        conn = conectar_bd()
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM atendimentos_responsaveis WHERE protocolo = ? LIMIT 1", (protocolo,))
        exists = cur.fetchone() is not None
        cur.close()
        conn.close()
        return jsonify({'ok': (not exists)})
    except Exception:
        # Em caso de erro, melhor bloquear para evitar duplicidade sem checagem
        return jsonify({'ok': False, 'error': 'erro_consulta'}), 500


# Alias para compatibilidade com template antigo:
@app.get('/api_check_protocolo')
def api_check_protocolo():
    return api_atendimentos_check_protocolo()


@app.get('/api_turmas')
def api_turmas():
    """Lista turmas por turno. Saída: [{id, nome, turno}]."""
    turno = (request.args.get('turno') or '').strip()
    if not turno:
        return jsonify([])

    try:
        conn = conectar_bd()
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        # Ajuste o ORDER BY se quiser
        cur.execute("SELECT id, nome, turno FROM turmas WHERE turno = ? ORDER BY nome", (turno,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return jsonify([{'id': r['id'], 'nome': r['nome'], 'turno': r['turno']} for r in rows])
    except Exception:
        return jsonify([])


@app.get('/api_alunos')
def api_alunos():
    """Lista alunos por turma_id. Saída: [{id, nome}]."""
    turma_id = (request.args.get('turma_id') or '').strip()
    if not turma_id:
        return jsonify([])

    try:
        conn = conectar_bd()
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        cur.execute("SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome", (turma_id,))
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return jsonify([{'id': r['id'], 'nome': r['nome']} for r in rows])
    except Exception:
        return jsonify([])


@app.route('/atendimentos/ver/<int:atendimento_id>')
def atendimentos_ver(atendimento_id):
    # Apenas moderador (gestão) pode ver
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("""
        SELECT
            ar.*,
            t.nome AS turma_nome,
            t.turno AS turma_turno,
            a.nome AS aluno_nome
        FROM atendimentos_responsaveis ar
        JOIN turmas t ON t.id = ar.turma_id
        JOIN alunos a ON a.id = ar.aluno_id
        WHERE ar.id = ?
    """, (atendimento_id,))
    at = cur.fetchone()
    cur.close()
    conn.close()

    if not at:
        flash("Atendimento não encontrado.")
        return redirect(url_for('atendimentos_historico'))

    return render_template('atendimentos_ver.html', at=at)


@app.route('/atendimentos/pdf/<int:atendimento_id>', methods=['GET'])
def atendimentos_pdf(atendimento_id):
    # Apenas moderador pode gerar o PDF do atendimento
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash("Acesso não autorizado.")
        return redirect(url_for('login'))

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('''
        SELECT
            a.*,
            t.nome AS turma_nome,
            t.turno AS turma_turno,
            al.nome AS aluno_nome,
            COALESCE(NULLIF(a.responsavel_nome,''), r.login) AS responsavel_nome_exibicao
        FROM atendimentos_responsaveis a
        JOIN turmas t ON a.turma_id = t.id
        JOIN alunos al ON a.aluno_id = al.id
        LEFT JOIN responsaveis r ON a.responsavel_id = r.id
        WHERE a.id = ?
        LIMIT 1
    ''', (atendimento_id,))
    at = cursor.fetchone()
    cursor.close()
    conn.close()

    if not at:
        flash("Atendimento não encontrado.")
        return redirect(url_for('atendimentos_historico'))

    # =========================
    # CABEÇALHO COM LOGOS
    # =========================
    def _draw_header(pdf, titulo):
        logo_esq = os.path.join(app.root_path, 'static', 'logo.jpg')
        logo_dir = os.path.join(app.root_path, 'static', 'logo1.PNG')

        y_top = 770

        if os.path.exists(logo_esq):
            pdf.drawImage(
                logo_esq,
                40, y_top - 45,
                width=55, height=55,
                preserveAspectRatio=True,
                mask='auto'
            )

        if os.path.exists(logo_dir):
            pdf.drawImage(
                logo_dir,
                520, y_top - 45,
                width=55, height=55,
                preserveAspectRatio=True,
                mask='auto'
            )

        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawCentredString(306, y_top - 10, titulo)

        # Texto institucional (igual ao da imagem)
        pdf.setFont("Helvetica-Bold", 10)
        pdf.drawCentredString(
            306,
            y_top - 24,
            "ESCOLA CLASSE 16"
        )

        pdf.setFont("Helvetica", 9)
        pdf.drawCentredString(
            306,
            y_top - 38,
            "SECRETARIA DE ESTADO DE EDUCAÇÃO DO DISTRITO FEDERAL"
        )
        pdf.drawCentredString(
            306,
            y_top - 50,
            "SUBSECRETARIA DE EDUCAÇÃO BÁSICA"
        )
        pdf.drawCentredString(
            306,
            y_top - 62,
            "Coordenação Regional de Sobradinho"
        )
        pdf.drawCentredString(
            306,
            y_top - 74,
            "ESCOLA CLASSE 16"
        )

        pdf.line(40, y_top - 90, 572, y_top - 90)
        return y_top - 105

    def _draw_block(pdf, x, y, title, text, max_width=510, line_h=13):
        pdf.setFont("Helvetica-Bold", 10)
        pdf.drawString(x, y, title)
        y -= 14

        pdf.setFont("Helvetica", 10)
        text = (text or '').strip()

        if not text:
            pdf.setFillColor(colors.grey)
            pdf.drawString(x, y, "—")
            pdf.setFillColor(colors.black)
            return y - 18

        lines = simpleSplit(text, "Helvetica", 10, max_width)
        for ln in lines:
            if y < 60:
                pdf.showPage()
                y = _draw_header(pdf, "ATA DE ATENDIMENTO")
                pdf.setFont("Helvetica", 10)
            pdf.drawString(x, y, ln)
            y -= line_h

        return y - 6

    buf = BytesIO()
    pdf = canvas.Canvas(buf, pagesize=letter)
    pdf.setTitle("Ata de Atendimento")

    y = _draw_header(pdf, "ATA DE ATENDIMENTO")

    protocolo = at["protocolo"] or "—"
    data_at = at["data_atendimento"] or "—"
    hora_at = at["hora_atendimento"] or ""
    data_hora = f"{data_at} {hora_at}".strip()

    registrador = (at["registrador_nome"] or "").strip()
    cargo = (at["registrador_cargo"] or "").strip()

    pdf.setFont("Helvetica", 10)
    pdf.drawString(40, y, f"Protocolo: {protocolo}")
    pdf.drawRightString(572, y, f"Data(Ano/Mês/Dia)/Hora: {data_hora if data_hora else '—'}")
    y -= 22

    pdf.drawString(40, y, f"Turma: {at['turma_nome']} ({at['turma_turno']})")
    y -= 14
    pdf.drawString(40, y, f"Aluno(a): {at['aluno_nome']}")
    y -= 14
    resp_disp = (at['responsavel_nome_exibicao'] or '—')
    try:
        parentesco = (at['responsavel_parentesco'] or '').strip()
    except Exception:
        parentesco = ''
    if parentesco:
        resp_disp = f"{resp_disp} ({parentesco})"
    pdf.drawString(40, y, f"Responsável atendido: {resp_disp}")
    y -= 14
    pdf.drawString(
        40, y,
        f"Registrado por: {registrador if registrador else '—'} – {cargo if cargo else '—'}"
    )
    y -= 22

    y = _draw_block(pdf, 40, y, "Assunto:", at["assunto"])
    y = _draw_block(
        pdf, 40, y,
        "Envolve professor:",
        f"Sim. Professor(a): {(at['professor_nome'] or '—')}" if at["envolve_professor"] else "Não"
    )
    y = _draw_block(pdf, 40, y, "Relato:", at["relato"])
    y = _draw_block(pdf, 40, y, "Combinados:", at["combinados"])

    retorno_txt = "Sim" if at["retorno_previsto"] else "Não"
    if at["retorno_previsto"] and at["retorno_em"]:
        retorno_txt += f" (previsto para {at['retorno_em']})"

    reuniao_txt = "Sim" if at["reuniao_agendada"] else "Não"
    if at["reuniao_agendada"] and at["reuniao_data"]:
        reuniao_txt += f" (agendada para {at['reuniao_data']})"

    y = _draw_block(pdf, 40, y, "Retorno previsto:", retorno_txt)
    y = _draw_block(pdf, 40, y, "Reunião agendada:", reuniao_txt)

    if y < 120:
        pdf.showPage()
        y = _draw_header(pdf, "ATA DE ATENDIMENTO")

    y -= 8
    pdf.line(40, y, 300, y)
    pdf.line(320, y, 572, y)
    y -= 12

    pdf.setFont("Helvetica", 9)
    pdf.drawString(40, y, "Assinatura do(a) responsável")
    pdf.drawString(320, y, "Assinatura do(a) registrador(a)")

    pdf.showPage()
    pdf.save()
    buf.seek(0)

    filename = f"ata_atendimento_{atendimento_id}.pdf"
    return send_file(
        buf,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=filename
    )


@app.route('/api/calendario/eventos')
def api_calendario_eventos():
    """
    Retorna eventos do calendário para o professor logado:
    - Avaliações agendadas
    - Períodos de planejamento
    - Ocorrências importantes
    """
    if 'usuario' not in session or session.get('tipo') != 'professor':
        return jsonify({'ok': False, 'error': 'Não autorizado'}), 403

    professor_id = obter_professor_id(session['usuario'])
    if not professor_id:
        return jsonify({'ok': False, 'error': 'Professor não encontrado'}), 404

    # Filtros opcionais
    mes = request.args.get('mes')  # 1-12
    ano = request.args.get('ano')  # 2024, 2025...

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    eventos = []

    # ========== AVALIAÇÕES AGENDADAS ==========
    sql_av = """
        SELECT
            a.id,
            a.disciplina,
            a.bimestre,
            a.tipo_avaliacao,
            a.descricao_avaliacao,
            a.data_avaliacao,
            a.pontuacao,
            t.nome AS turma_nome,
            t.turno AS turma_turno
        FROM avaliacoes_bimestrais a
        JOIN turmas t ON a.turma_id = t.id
        WHERE a.professor_id = ?
          AND a.data_avaliacao IS NOT NULL
    """
    params_av = [professor_id]

    if ano and mes:
        sql_av += " AND strftime('%Y', a.data_avaliacao) = ? AND strftime('%m', a.data_avaliacao) = ?"
        params_av.extend([ano, f"{int(mes):02d}"])
    elif ano:
        sql_av += " AND strftime('%Y', a.data_avaliacao) = ?"
        params_av.append(ano)

    sql_av += " ORDER BY a.data_avaliacao"

    cursor.execute(sql_av, params_av)
    avaliacoes = cursor.fetchall()

    for av in avaliacoes:
        eventos.append({
            'id': f"av_{av['id']}",
            'tipo': 'avaliacao',
            'titulo': f"{av['tipo_avaliacao'] or 'Avaliação'} - {av['disciplina']}",
            'descricao': av['descricao_avaliacao'] or '',
            'data': av['data_avaliacao'],
            'turma': f"{av['turma_nome']} ({av['turma_turno']})",
            'pontuacao': av['pontuacao'],
            'cor': '#f59e0b',  # laranja
            'icone': 'clipboard-check'
        })

    # ========== PERÍODOS DE PLANEJAMENTO ==========
    sql_plan = """
        SELECT DISTINCT
            pi.id,
            pi.conteudo,
            pi.data_inicio,
            pi.data_fim,
            p.disciplina,
            p.bimestre,
            GROUP_CONCAT(DISTINCT t.nome || ' (' || t.turno || ')') AS turmas
        FROM planejamento_itens pi
        JOIN planejamentos p ON pi.planejamento_id = p.id
        LEFT JOIN planejamentos_turmas pt ON pt.planejamento_id = p.id
        LEFT JOIN turmas t ON t.id = pt.turma_id
        WHERE p.professor_id = ?
          AND (pi.data_inicio IS NOT NULL OR pi.data_fim IS NOT NULL)
    """
    params_plan = [professor_id]

    if ano and mes:
        sql_plan += """
            AND (
                (strftime('%Y', pi.data_inicio) = ? AND strftime('%m', pi.data_inicio) = ?)
                OR (strftime('%Y', pi.data_fim) = ? AND strftime('%m', pi.data_fim) = ?)
            )
        """
        params_plan.extend([ano, f"{int(mes):02d}", ano, f"{int(mes):02d}"])
    elif ano:
        sql_plan += """
            AND (
                strftime('%Y', pi.data_inicio) = ?
                OR strftime('%Y', pi.data_fim) = ?
            )
        """
        params_plan.extend([ano, ano])

    sql_plan += " GROUP BY pi.id ORDER BY pi.data_inicio, pi.data_fim"

    cursor.execute(sql_plan, params_plan)
    planejamentos = cursor.fetchall()

    for plan in planejamentos:
        # Truncar conteúdo longo
        conteudo = (plan['conteudo'] or '')[:50]
        if len(plan['conteudo'] or '') > 50:
            conteudo += '...'

        eventos.append({
            'id': f"plan_{plan['id']}",
            'tipo': 'planejamento',
            'titulo': f"{plan['disciplina']} - {conteudo}",
            'descricao': plan['conteudo'],
            'data_inicio': plan['data_inicio'],
            'data_fim': plan['data_fim'],
            'turmas': plan['turmas'],
            'bimestre': plan['bimestre'],
            'cor': '#2563eb',  # azul
            'icone': 'book-open'
        })

    cursor.close()
    conn.close()

    return jsonify({
        'ok': True,
        'eventos': eventos,
        'total': len(eventos)
    })


# ============== ROTA 2: NOTIFICAÇÕES DO PROFESSOR ==============

@app.route('/api/notificacoes/professor')
def api_notificacoes_professor():
    """
    Retorna notificações para o professor:
    - Recados visualizados/não visualizados pelos responsáveis
    - Ocorrências com reunião agendada próxima
    - Planejamentos pendentes
    """
    if 'usuario' not in session or session.get('tipo') != 'professor':
        return jsonify({'ok': False, 'error': 'Não autorizado'}), 403

    professor_id = obter_professor_id(session['usuario'])
    if not professor_id:
        return jsonify({'ok': False, 'error': 'Professor não encontrado'}), 404

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    notificacoes = []

    # ========== RECADOS NÃO VISUALIZADOS ==========
    cursor.execute("""
        SELECT
            ra.id,
            ra.data_criacao,
            a.nome AS aluno_nome,
            t.nome AS turma_nome,
            t.turno AS turma_turno
        FROM recados_aluno ra
        JOIN alunos a ON ra.aluno_id = a.id
        JOIN turmas t ON ra.turma_id = t.id
        WHERE ra.professor_id = ?
          AND ra.visualizado = 0
          AND IFNULL(ra.excluido_para_responsavel, 0) = 0
        ORDER BY ra.data_criacao DESC
        LIMIT 10
    """, (professor_id,))

    recados_nao_lidos = cursor.fetchall()

    for rec in recados_nao_lidos:
        # Calcular dias desde o envio
        try:
            data_envio = datetime.strptime(rec['data_criacao'][:10], '%Y-%m-%d')
            dias = (datetime.now() - data_envio).days
        except:
            dias = 0

        notificacoes.append({
            'id': f"rec_{rec['id']}",
            'tipo': 'recado_nao_lido',
            'titulo': f"Recado não visualizado - {rec['aluno_nome']}",
            'descricao': f"{rec['turma_nome']} ({rec['turma_turno']}) • Enviado há {dias} dia(s)",
            'data': rec['data_criacao'],
            'urgencia': 'alta' if dias > 7 else 'media',
            'icone': 'bell',
            'link': url_for('listar_recados_aluno')
        })

    # ========== RECADOS VISUALIZADOS RECENTEMENTE ==========
    cursor.execute("""
        SELECT
            ra.id,
            ra.data_criacao,
            a.nome AS aluno_nome,
            t.nome AS turma_nome
        FROM recados_aluno ra
        JOIN alunos a ON ra.aluno_id = a.id
        JOIN turmas t ON ra.turma_id = t.id
        WHERE ra.professor_id = ?
          AND ra.visualizado = 1
          AND IFNULL(ra.excluido_para_responsavel, 0) = 0
          AND date(ra.data_criacao) >= date('now', '-7 days')
        ORDER BY ra.data_criacao DESC
        LIMIT 5
    """, (professor_id,))

    recados_lidos = cursor.fetchall()

    for rec in recados_lidos:
        notificacoes.append({
            'id': f"rec_lido_{rec['id']}",
            'tipo': 'recado_lido',
            'titulo': f"✓ Recado lido - {rec['aluno_nome']}",
            'descricao': f"{rec['turma_nome']}",
            'data': rec['data_criacao'],
            'urgencia': 'baixa',
            'icone': 'check-circle',
            'link': url_for('listar_recados_aluno')
        })

    # ========== AVALIAÇÕES PRÓXIMAS (7 dias) ==========
    cursor.execute("""
        SELECT
            a.id,
            a.disciplina,
            a.tipo_avaliacao,
            a.data_avaliacao,
            t.nome AS turma_nome
        FROM avaliacoes_bimestrais a
        JOIN turmas t ON a.turma_id = t.id
        WHERE a.professor_id = ?
          AND a.data_avaliacao IS NOT NULL
          AND date(a.data_avaliacao) BETWEEN date('now') AND date('now', '+7 days')
        ORDER BY a.data_avaliacao
        LIMIT 5
    """, (professor_id,))

    avaliacoes_proximas = cursor.fetchall()

    for av in avaliacoes_proximas:
        try:
            data_av = datetime.strptime(av['data_avaliacao'], '%Y-%m-%d')
            dias_restantes = (data_av - datetime.now()).days
        except:
            dias_restantes = 0

        notificacoes.append({
            'id': f"av_prox_{av['id']}",
            'tipo': 'avaliacao_proxima',
            'titulo': f"Avaliação em {dias_restantes} dia(s)",
            'descricao': f"{av['tipo_avaliacao'] or 'Avaliação'} - {av['disciplina']} ({av['turma_nome']})",
            'data': av['data_avaliacao'],
            'urgencia': 'alta' if dias_restantes <= 2 else 'media',
            'icone': 'alert-circle',
            'link': url_for('listar_avaliacoes_professor')
        })

    cursor.close()
    conn.close()

    # Ordenar por urgência e data
    prioridade = {'alta': 0, 'media': 1, 'baixa': 2}
    notificacoes.sort(key=lambda x: (prioridade.get(x['urgencia'], 3), x['data']), reverse=True)

    return jsonify({
        'ok': True,
        'notificacoes': notificacoes,
        'total': len(notificacoes),
        'nao_lidas': len(recados_nao_lidos),
        'urgentes': len([n for n in notificacoes if n['urgencia'] == 'alta'])
    })


# ============== ROTA 3: ESTATÍSTICAS DO PROFESSOR ==============

@app.route('/api/estatisticas/professor')
def api_estatisticas_professor():
    """
    Estatísticas gerais do professor para exibir no dashboard
    """
    if 'usuario' not in session or session.get('tipo') != 'professor':
        return jsonify({'ok': False, 'error': 'Não autorizado'}), 403

    professor_id = obter_professor_id(session['usuario'])
    if not professor_id:
        return jsonify({'ok': False, 'error': 'Professor não encontrado'}), 404

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    stats = {}

    # Total de turmas
    cursor.execute("""
        SELECT COUNT(DISTINCT turma_id) AS total
        FROM professor_turmas
        WHERE professor_id = ?
    """, (professor_id,))
    stats['turmas'] = cursor.fetchone()['total']

    # Total de planejamentos
    cursor.execute("""
        SELECT COUNT(*) AS total
        FROM planejamentos
        WHERE professor_id = ?
    """, (professor_id,))
    stats['planejamentos'] = cursor.fetchone()['total']

    # Total de avaliações
    cursor.execute("""
        SELECT COUNT(*) AS total
        FROM avaliacoes_bimestrais
        WHERE professor_id = ?
    """, (professor_id,))
    stats['avaliacoes'] = cursor.fetchone()['total']

    # Recados não lidos
    cursor.execute("""
        SELECT COUNT(*) AS total
        FROM recados_aluno
        WHERE professor_id = ?
          AND visualizado = 0
          AND IFNULL(excluido_para_responsavel, 0) = 0
    """, (professor_id,))
    stats['recados_nao_lidos'] = cursor.fetchone()['total']

    # Ocorrências registradas
    cursor.execute("""
        SELECT COUNT(*) AS total
        FROM ocorrencias o
        WHERE o.professor = (
            SELECT login FROM professores WHERE id = ?
        )
    """, (professor_id,))
    stats['ocorrencias'] = cursor.fetchone()['total']

    cursor.close()
    conn.close()

    return jsonify({
        'ok': True,
        'stats': stats
    })


# Logout

@app.route('/logout')
def logout():
    session.clear()
    flash("Você saiu com sucesso.")
    return redirect(url_for('index'))


# ============================================================
# LISTA DE PRESENÇA (Moderador) - Word (modelo) e PDF
# ============================================================

def _parse_data_br(data_str: str) -> str:
    data_str = (data_str or "").strip()
    try:
        dt = datetime.strptime(data_str, "%d/%m/%Y")
        return dt.strftime("%d/%m/%Y")
    except Exception:
        raise ValueError("Data inválida. Use o formato dd/mm/aaaa.")


import glob


def _encontrar_logos_lista_presenca() -> tuple[str | None, str | None]:
    """Tenta localizar duas logos no projeto (sem quebrar se não existir)."""
    base_dir = os.path.dirname(os.path.abspath(__file__))

    # ✅ Prioridade absoluta (solicitado):
    #    ESQUERDA = static/logo.jpg
    #    DIREITA  = static/logo1.PNG
    # Nota: em Linux (PythonAnywhere) o nome do arquivo é *case-sensitive*.
    # Então "logo1.PNG" é diferente de "logo1.png".
    static_dir = os.path.join(base_dir, "static")
    logo_esq_prior = os.path.join(static_dir, "logo.jpg")
    logo_dir_prior = os.path.join(static_dir, "logo1.PNG")
    if os.path.isfile(logo_esq_prior) and os.path.isfile(logo_dir_prior):
        return logo_esq_prior, logo_dir_prior

    # fallback 1: aceitar variações comuns de maiúsculas/minúsculas
    for alt_dir in ["logo1.png", "logo1.jpg", "logo1.jpeg", "logo_1.png"]:
        p_alt = os.path.join(static_dir, alt_dir)
        if os.path.isfile(logo_esq_prior) and os.path.isfile(p_alt):
            return logo_esq_prior, p_alt
    candidatos: list[str] = []
    pastas = [
        base_dir,
        os.path.join(base_dir, "static"),
        os.path.join(base_dir, "static", "img"),
        os.path.join(base_dir, "static", "images"),
        os.path.join(base_dir, "static", "assets"),
    ]
    padroes = [
        "logo_left.png", "logo_right.png",
        "logo_esquerda.png", "logo_direita.png",
        "logo1.png", "logo2.png",
        "*logo*.png", "*logo*.jpg", "*logo*.jpeg",
        "*brasao*.png", "*brasao*.jpg", "*emblema*.png",
    ]
    for pasta in pastas:
        if not os.path.isdir(pasta):
            continue
        for pad in padroes:
            for f in sorted(glob.glob(os.path.join(pasta, pad))):
                if os.path.isfile(f) and f not in candidatos:
                    candidatos.append(f)

    # preferir pares explícitos left/right se existirem
    left = None
    right = None
    for nome in ["logo_left.png", "logo_esquerda.png"]:
        for pasta in pastas:
            f = os.path.join(pasta, nome)
            if os.path.isfile(f):
                left = f
                break
        if left:
            break

    for nome in ["logo_right.png", "logo_direita.png"]:
        for pasta in pastas:
            f = os.path.join(pasta, nome)
            if os.path.isfile(f):
                right = f
                break
        if right:
            break

    # fallback: pegar as duas primeiras encontradas
    if left is None and candidatos:
        left = candidatos[0]
    if right is None:
        if len(candidatos) >= 2:
            right = candidatos[1]
        elif left is not None:
            right = left

    return left, right


def _inserir_logos_no_docx(d):
    """Insere uma logo em cada lado do cabeçalho (tabela 0, linha 0)."""
    try:
        from docx.shared import Cm
    except Exception:
        return

    if not getattr(d, "tables", None):
        return
    if len(d.tables) < 1:
        return

    left, right = _encontrar_logos_lista_presenca()
    if not left and not right:
        return

    t0 = d.tables[0]
    if len(t0.rows) < 1 or len(t0.columns) < 3:
        return

    def _clear_cell(cell):
        # limpa parágrafos mantendo um
        for p in list(cell.paragraphs)[1:]:
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass
        if cell.paragraphs:
            cell.paragraphs[0].clear()

    if left and os.path.isfile(left):
        cell = t0.rows[0].cells[0]
        _clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = 1
        run = p.add_run()
        try:
            run.add_picture(left, width=Cm(2.1))
        except Exception:
            pass

    if right and os.path.isfile(right):
        cell = t0.rows[0].cells[2]
        _clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = 1
        run = p.add_run()
        try:
            run.add_picture(right, width=Cm(2.1))
        except Exception:
            pass

    try:
        mid = t0.rows[0].cells[1]
        if not mid.text.strip():
            mid.text = "LISTA DE PRESENÇA"
        mid.paragraphs[0].alignment = 1
    except Exception:
        pass


def _ajustar_tabela_assinaturas_docx(t1):
    """Aumenta a coluna de assinaturas e a altura das linhas para assinatura."""
    try:
        from docx.shared import Cm
        from docx.enum.table import WD_ROW_HEIGHT_RULE
    except Exception:
        return

    try:
        col_w = [Cm(1.0), Cm(12.4), Cm(4.8)]  # nº, nome (mais largo), assinatura
        for c_idx, w in enumerate(col_w):
            try:
                t1.columns[c_idx].width = w
            except Exception:
                pass
            for row in t1.rows:
                try:
                    row.cells[c_idx].width = w
                except Exception:
                    pass

        # aumentar altura das linhas dos estudantes
        for r in range(1, min(35, len(t1.rows))):
            row = t1.rows[r]
            row.height = Cm(0.85)  # mais compacto para caber turmas grandes
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    except Exception:
        pass


def _set_cell_font_arial(cell, size_pt: int = 10):
    """Define fonte Arial (tamanho ajustável) no conteúdo (runs) da célula."""
    try:
        from docx.shared import Pt
    except Exception:
        return
    try:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(size_pt)
    except Exception:
        pass


def _buscar_turmas_para_lista(turno: str, turma_ids: list[int] | None, todas: bool):
    conn = conectar_bd()
    cur = conn.cursor()

    params = []
    where = []
    if turno and turno != "todos":
        where.append("turno = ?")
        params.append(turno)

    if not todas:
        turma_ids = turma_ids or []
        if turma_ids:
            placeholders = ",".join(["?"] * len(turma_ids))
            where.append(f"id IN ({placeholders})")
            params.extend(turma_ids)
        else:
            cur.close()
            conn.close()
            return []

    sql = "SELECT id, nome, turno FROM turmas"
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY turno, nome"

    cur.execute(sql, params)
    turmas = cur.fetchall()
    cur.close()
    conn.close()
    return turmas


def _buscar_alunos_da_turma(turma_id: int):
    conn = conectar_bd()
    cur = conn.cursor()
    cur.execute("SELECT nome FROM alunos WHERE turma_id = ? ORDER BY nome", (turma_id,))
    alunos = [r["nome"] if isinstance(r, sqlite3.Row) else r[0] for r in cur.fetchall()]
    cur.close()
    conn.close()
    return alunos


def _preencher_docx_lista_presenca(modelo_path: str, turma_nome: str, data_br: str, atividade: str,
                                   alunos: list[str]) -> BytesIO:
    from docx import Document
    por_pagina = 34
    total_paginas = max(1, (len(alunos) + por_pagina - 1) // por_pagina)

    doc_final = None

    for p in range(total_paginas):
        inicio = p * por_pagina
        fim = inicio + por_pagina
        bloco = alunos[inicio:fim]

        d = Document(modelo_path)

        _inserir_logos_no_docx(d)

        # Tabela 0: infos
        if d.tables:
            t0 = d.tables[0]
            try:
                t0.rows[1].cells[1].text = atividade
            except Exception:
                pass
            try:
                t0.rows[2].cells[1].text = turma_nome
                t0.rows[2].cells[2].text = f"Data: {data_br}"
            except Exception:
                pass

        # Tabela 1: estudantes
        if len(d.tables) > 1:
            t1 = d.tables[1]
            _ajustar_tabela_assinaturas_docx(t1)
            for i in range(34):
                nome = bloco[i] if i < len(bloco) else ""
                try:
                    cell_nome = t1.rows[i + 1].cells[1]
                    cell_nome.text = nome
                    # Ajuste fino para caber nomes longos (quebra automática dentro da célula)
                    try:
                        for p in cell_nome.paragraphs:
                            p.paragraph_format.space_before = 0
                            p.paragraph_format.space_after = 0
                            p.paragraph_format.line_spacing = 1.0
                    except Exception:
                        pass
                    _set_cell_font_arial(cell_nome, 8)
                except Exception:
                    pass

        if doc_final is None:
            doc_final = d
        else:
            doc_final.add_page_break()
            for element in d.element.body:
                doc_final.element.body.append(element)

    buf = BytesIO()
    doc_final.save(buf)
    buf.seek(0)
    return buf


def _gerar_pdf_lista_presenca(turmas_com_alunos: list[dict], data_br: str, atividade: str) -> BytesIO:
    """
    PDF com uma ou mais páginas por turma.
    🔧 Correção: calcula quantas linhas cabem na página para não "cortar" os últimos nomes.
    """
    buf = BytesIO()

    # --- Fonte: tentar Arial; se não existir, usar LiberationSans/DejaVu (bem próximo) e, por fim, Helvetica ---
    pdf_font = "Helvetica"
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        possiveis = [
            # Windows
            os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts", "arial.ttf"),
            # Linux comuns (PythonAnywhere costuma ter DejaVu)
            "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        ]
        for p in possiveis:
            if p and os.path.isfile(p):
                nome_reg = "Arial" if os.path.basename(p).lower().startswith("arial") else "SansFallback"
                try:
                    pdfmetrics.registerFont(TTFont(nome_reg, p))
                    pdf_font = nome_reg
                    break
                except Exception:
                    continue
    except Exception:
        pass
    # A4 (retraro) costuma casar melhor com a lista
    from reportlab.lib.pagesizes import A4
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    def header():
        y = height - 50
        # logos (se existirem) - um de cada lado
        try:
            from reportlab.lib.utils import ImageReader
            left_logo, right_logo = _encontrar_logos_lista_presenca()
            logo_w = 46
            logo_h = 46
            y_logo = height - 92
            if left_logo and os.path.isfile(left_logo):
                c.drawImage(ImageReader(left_logo), 40, y_logo, width=logo_w, height=logo_h, mask='auto',
                            preserveAspectRatio=True)
            if right_logo and os.path.isfile(right_logo):
                c.drawImage(ImageReader(right_logo), width - 40 - logo_w, y_logo, width=logo_w, height=logo_h,
                            mask='auto', preserveAspectRatio=True)
        except Exception:
            pass
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(width / 2, y, "SECRETARIA DE ESTADO DE EDUCAÇÃO DO DISTRITO FEDERAL")
        y -= 16
        c.drawCentredString(width / 2, y, "ESCOLA CLASSE 16 DE SOBRADINHO II")
        y -= 26
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width / 2, y, "LISTA DE PRESENÇA")
        y -= 22
        c.setFont("Helvetica", 11)
        c.drawString(50, y, f"Atividade: {atividade}")
        y -= 16
        return y

    # Layout da tabela
    margin_top = 40
    margin_bottom = 40
    row_h = 15  # altura da linha
    row_font_size = 8  # fonte menor para caber turmas com ~35 alunos e nomes longos

    # Quantas linhas cabem na página (depois do cabeçalho)
    def _linhas_por_pagina(y_start: float) -> int:
        disponivel = max(0, (y_start - margin_bottom) - row_h)  # -1 linha de segurança
        return max(10, int(disponivel // row_h))

    for item in turmas_com_alunos:
        turma_nome = item["turma"]
        alunos = item["alunos"]

        # primeira página: calcula y de início e quantas linhas cabem
        # (isso evita o problema dos últimos nomes "sumirem" fora da folha)
        # A cada página o cálculo repete porque o cabeçalho é igual.
        #
        # Obs: a tabela começa depois de "Turma/Data" e do cabeçalho da tabela.
        #
        # Determinar linhas por página dinamicamente
        #
        # Faz a contagem e paginação em loop
        idx = 0
        while idx < len(alunos) or idx == 0:
            y = header()
            c.setFont(f"{pdf_font}", 11)
            c.drawString(50, y, f"Turma: {turma_nome}")
            c.drawRightString(width - 50, y, f"Data: {data_br}")
            y -= 22

            x0 = 50
            col_num = 35
            col_ass = 170  # assinatura (ainda confortável) - libera mais largura para o nome
            col_nome = (width - 50) - (x0 + col_num + col_ass)

            # cabeçalho tabela
            def _wrap_nome_pdf(txt: str, max_w: float) -> list[str]:
                """Quebra nome em até 2 linhas para caber na coluna."""
                txt = (txt or "").strip()
                if not txt:
                    return [""]
                try:
                    from reportlab.pdfbase import pdfmetrics
                    def w(s):
                        return pdfmetrics.stringWidth(s, pdf_font, row_font_size)
                except Exception:
                    def w(s):
                        return len(s) * 5
                words = txt.split()
                lines = [""]
                for word in words:
                    cand = (lines[-1] + " " + word).strip()
                    if w(cand) <= (max_w - 12):
                        lines[-1] = cand
                    else:
                        lines.append(word)
                        if len(lines) == 2:
                            # se estourar 2 linhas, corta o restante com reticências
                            rest = " ".join(words[words.index(word):])
                            # tenta encaixar com ...
                            base = lines[-1]
                            while w(base + "...") > (max_w - 12) and len(base) > 3:
                                base = base[:-1]
                            lines[-1] = (base + "...") if base else "..."
                            break
                return lines[:2]

            c.setFont("Helvetica-Bold", 10)
            c.rect(x0, y - 22, col_num, 22)
            c.rect(x0 + col_num, y - 22, col_nome, 22)
            c.rect(x0 + col_num + col_nome, y - 22, col_ass, 22)
            c.drawCentredString(x0 + col_num / 2, y - 13, "Nº")
            c.drawString(x0 + col_num + 6, y - 13, "Estudantes")
            c.drawString(x0 + col_num + col_nome + 6, y - 13, "Assinaturas")

            y -= 22
            # Fonte dos nomes: Arial 12 (ou fallback) como você pediu.
            c.setFont(f"{pdf_font}", 12)

            linhas = _linhas_por_pagina(y)
            bloco = alunos[idx:idx + linhas]
            for i in range(linhas):
                yy = y - (i + 1) * row_h
                c.rect(x0, yy, col_num, row_h)
                c.rect(x0 + col_num, yy, col_nome, row_h)
                c.rect(x0 + col_num + col_nome, yy, col_ass, row_h)
                c.setFont(f"{pdf_font}", row_font_size)
                c.drawCentredString(x0 + col_num / 2, yy + (row_h / 2 - 4), f"{i + 1:02d}")
                nome = bloco[i] if i < len(bloco) else ""
                c.setFont(f"{pdf_font}", row_font_size)
                linhas = _wrap_nome_pdf(str(nome), col_nome)
                # desenha 1 ou 2 linhas dentro da mesma linha (assinatura permanece na mesma linha)
                if len(linhas) == 1:
                    c.drawString(x0 + col_num + 6, yy + (row_h / 2 - 4), linhas[0])
                else:
                    c.drawString(x0 + col_num + 6, yy + (row_h - 7), linhas[0])
                    c.drawString(x0 + col_num + 6, yy + 4, linhas[1])

            idx += len(bloco)
            c.showPage()

    c.save()
    buf.seek(0)
    return buf


@app.route("/gerar_lista_presenca", methods=["GET", "POST"])
def gerar_lista_presenca():
    if "usuario" not in session or session.get("tipo") != "moderador":
        flash("Acesso não autorizado.")
        return redirect(url_for("login"))

    # carregar turmas para exibição
    conn = conectar_bd()
    cur = conn.cursor()
    cur.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
    turmas = cur.fetchall()
    cur.close()
    conn.close()

    if request.method == "GET":
        return render_template("gerar_lista_presenca.html", turmas=turmas)

    try:
        turno = (request.form.get("turno") or "todos").strip()
        todas = (request.form.get("todas_turmas") == "1")

        turma_ids_raw = request.form.getlist("turmas")
        turma_ids = []
        for tid in turma_ids_raw:
            try:
                turma_ids.append(int(tid))
            except Exception:
                pass

        tipo = (request.form.get("tipo") or "").strip()
        outros = (request.form.get("outros") or "").strip()
        atividade = (outros if outros else "Outros") if tipo == "outros" else tipo

        data_br = _parse_data_br(request.form.get("data") or "")
        formato = (request.form.get("formato") or "pdf").strip().lower()

        turmas_selecionadas = _buscar_turmas_para_lista(turno, turma_ids, todas)
        if not turmas_selecionadas:
            flash("Selecione pelo menos uma turma (ou marque 'Todas as turmas').")
            return redirect(url_for("gerar_lista_presenca"))

        turmas_com_alunos = []
        for t in turmas_selecionadas:
            turma_nome = f"{t['nome']} ({t['turno']})" if isinstance(t, sqlite3.Row) else f"{t[1]} ({t[2]})"
            turma_id = t["id"] if isinstance(t, sqlite3.Row) else t[0]
            alunos = _buscar_alunos_da_turma(turma_id)
            turmas_com_alunos.append({"turma": turma_nome, "alunos": alunos})

        modelo_path = os.path.join(os.getcwd(), "modelo lista de presença(provas, reunioes de pais e etc).docx")
        if not os.path.exists(modelo_path):
            flash(
                "Modelo não encontrado na raiz do projeto: 'modelo lista de presença(provas, reunioes de pais e etc).docx'")
            return redirect(url_for("gerar_lista_presenca"))

        if formato == "pdf":
            pdf_buf = _gerar_pdf_lista_presenca(turmas_com_alunos, data_br, atividade)
            nome_arquivo = f"lista_presenca_{data_br.replace('/', '-')}.pdf"
            return send_file(pdf_buf, as_attachment=True, download_name=nome_arquivo, mimetype="application/pdf")

        # Word
        if len(turmas_com_alunos) == 1:
            item = turmas_com_alunos[0]
            docx_buf = _preencher_docx_lista_presenca(modelo_path, item["turma"], data_br, atividade, item["alunos"])
            nome_arquivo = f"lista_presenca_{item['turma'].split('(')[0].strip()}_{data_br.replace('/', '-')}.docx"
            return send_file(
                docx_buf,
                as_attachment=True,
                download_name=nome_arquivo,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        zip_buf = BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for item in turmas_com_alunos:
                docx_buf = _preencher_docx_lista_presenca(modelo_path, item["turma"], data_br, atividade,
                                                          item["alunos"])
                nome_doc = f"lista_presenca_{item['turma'].split('(')[0].strip()}_{data_br.replace('/', '-')}.docx"
                zf.writestr(nome_doc, docx_buf.getvalue())

        zip_buf.seek(0)
        nome_zip = f"listas_presenca_{data_br.replace('/', '-')}.zip"
        return send_file(zip_buf, as_attachment=True, download_name=nome_zip, mimetype="application/zip")

    except ValueError as e:
        flash(str(e))
        return redirect(url_for("gerar_lista_presenca"))
    except Exception as e:
        flash(f"Erro ao gerar lista de presença: {e}")
        return redirect(url_for("gerar_lista_presenca"))


# ============================================================================
# ROTAS PARA MODERADOR ACESSAR CARÔMETRO
# ============================================================================

@app.route('/moderador/carometro')
@app.route('/moderador/carometro')
def moderador_ver_carometro():
    """
    Permite que moderadores acessem a visualização do carômetro.
    """
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash('Acesso não autorizado.')
        return redirect(url_for('login'))

    turma_id = (request.args.get('turma_id') or '').strip()

    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    try:
        # Busca todas as turmas
        cursor.execute('SELECT id, nome, turno FROM turmas ORDER BY turno, nome')
        turmas = cursor.fetchall()

        alunos = []
        if turma_id:
            # Busca alunos da turma com suas fotos
            cursor.execute('''
                SELECT 
                    a.id, 
                    a.nome,
                    cf.arquivo AS arquivo,
                    cf.atualizado_em AS atualizado_em
                FROM alunos a
                LEFT JOIN carometro_fotos cf ON cf.aluno_id = a.id
                WHERE a.turma_id = ?
                ORDER BY a.nome
            ''', (turma_id,))
            alunos = cursor.fetchall()

        cursor.close()
        conn.close()

        return render_template('carometro_ver.html',
                               turmas=turmas,
                               alunos=alunos,
                               turma_id=turma_id,
                               usuario=session['usuario'])

    except Exception as e:
        cursor.close()
        conn.close()
        flash(f'Erro ao carregar carômetro: {str(e)}')
        return redirect(url_for('dashboard_moderador'))


@app.route('/moderador/carometro/registrar')
def moderador_registrar_foto():
    """
    Permite que moderadores registrem fotos no carômetro.
    Usa o MESMO template que o professor, mas com tipo='moderador' para ajustar links.
    """
    if 'usuario' not in session or session.get('tipo') != 'moderador':
        flash('Acesso não autorizado.')
        return redirect(url_for('index'))

    # Busca turmas
    conn = conectar_bd()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    try:
        cursor.execute('SELECT id, nome, turno FROM turmas ORDER BY turno, nome')
        turmas = cursor.fetchall()

        cursor.close()
        conn.close()

        # ✅ IMPORTANTE: Renderiza o template do carômetro com tipo='moderador'
        # para que o JavaScript use as rotas corretas (/professor/carometro/api/...)
        return render_template('carometro_moderador.html',
                               turmas=turmas,
                               usuario=session['usuario'],
                               tipo='moderador')
    except Exception as e:
        cursor.close()
        conn.close()
        flash(f'Erro ao carregar página de registro: {str(e)}')
        return redirect(url_for('dashboard_moderador'))



if __name__ == '__main__':
    app.run(debug=True)