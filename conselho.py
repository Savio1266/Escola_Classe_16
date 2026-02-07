from __future__ import annotations
# conselho.py
# Blueprint: Conselho de Classe Coletivo
# - Professores alimentam por disciplina
# - Moderador acompanha status e gera Word (individual e da turma inteira)
#
# ✅ Compatível com SQLite já existente:
#    - Se a tabela professor_turmas_disciplina já existir com coluna "disciplina" (antiga),
#      o código detecta e usa.
#    - Se existir mas NÃO tiver disciplina_abrev, tenta migrar criando disciplina_abrev.
#    - Evita erro: sqlite3.OperationalError: no such column: ptd.disciplina_abrev

import os
import json
import sqlite3
from datetime import datetime
from copy import deepcopy
from typing import Optional, List, Dict, Any

from flask import Blueprint, render_template, request, redirect, url_for, session, flash, send_file
from docx import Document

bp_conselho = Blueprint("conselho", __name__, template_folder="templates")

# =========================
# CONFIG
# =========================

DB_PATH = os.environ.get("ESCOLA CLASSE 16_DB_PATH", "rfa.db")

# Pasta onde os .docx gerados serão salvos
CONSELHO_OUT_DIR = os.path.join("static", "conselhos_gerados")
os.makedirs(CONSELHO_OUT_DIR, exist_ok=True)

# Caminho padrão do MODELO (mantenha no root do projeto ou configure variável de ambiente)
MODELO_DOCX_PATH = os.environ.get("CONSELHO_MODELO_PATH", "MODELO CONSELHO DE CLASSE.docx")

# Mapeamento Disciplina (cadastro) -> abreviação (modelo)
DISCIPLINA_ABREV = {
    "Artes": "ART",
    "Ciencias": "CN",
    "Ciências": "CN",
    "Educacao Fisica": "E.F",
    "Educação Física": "E.F",
    "Geografia": "GEO",
    "Historia": "HIS",
    "História": "HIS",
    "Ingles": "ING",
    "Inglês": "ING",
    "Portugues": "LP",
    "Português": "LP",
    "Matematica": "MAT",
    "Matemática": "MAT",
}

# Ordem de colunas no modelo
COLUNAS_MODELO = ["ART", "CN", "E.F", "GEO", "HIS", "ING", "LP", "MAT"]

# Aspectos (linhas) do modelo (chave -> rótulo)
ASPECTOS = [
    ("elogios", "Elogios / Avanços percebidos"),
    ("apr_consolidada", "Aprendizagem consolidada"),
    ("apr_desenvolvimento", "Aprendizagem em desenvolvimento"),
    ("dif_persistentes", "Dificuldades persistentes de aprendizagem"),
    ("intervencao_continua", "Necessita intervenção pedagógica contínua"),
    ("cumpriu_ativ", "Cumpriu as atividades propostas"),
    ("cumpriu_parcial", "Cumpriu parcialmente as atividades"),
    ("nao_cumpriu", "Não cumpriu as atividades"),
    ("interesse", "Demonstra interesse"),
    ("desinteresse", "Desinteresse"),
    ("conversa_excessiva", "Conversa excessiva"),
    ("indisciplina", "Indisciplina"),
    ("freq_regular", "Frequência regular"),
    ("freq_irregular", "Frequência irregular"),
    ("infrequente", "Infrequente"),
    ("evolucao_bimestre", "Evolução comportamental no bimestre"),
]



# =========================
# Helpers (BD / Login)
# =========================

def _as_int(v, default=None):
    try:
        if v is None:
            return default
        v = str(v).strip()
        if v == "":
            return default
        return int(v)
    except Exception:
        return default


def conectar_bd():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def _tabela_existe(nome: str) -> bool:
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=? LIMIT 1", (nome,))
        return cur.fetchone() is not None
    finally:
        cur.close()
        conn.close()


def _coluna_existe(tabela: str, coluna: str) -> bool:
    if not _tabela_existe(tabela):
        return False
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute(f"PRAGMA table_info({tabela})")
        cols = [r[1] for r in cur.fetchall()]  # r[1] = nome da coluna
        return coluna in cols
    finally:
        cur.close()
        conn.close()


def _coluna_disciplina_ptd() -> Optional[str]:
    """
    Descobre qual coluna de disciplina existe em professor_turmas_disciplina.
    Prioriza 'disciplina_abrev'. Se não existir, tenta 'disciplina'.
    """
    if not _tabela_existe("professor_turmas_disciplina"):
        return None
    if _coluna_existe("professor_turmas_disciplina", "disciplina_abrev"):
        return "disciplina_abrev"
    if _coluna_existe("professor_turmas_disciplina", "disciplina"):
        return "disciplina"
    return None


def _abrevs_possiveis_para(valor: str) -> List[str]:
    """
    Monta lista de valores possíveis pra comparar no banco:
    - valor pode ser 'LP' (abrev) OU 'Português' (nome)
    - retorna lista com abrev + nomes equivalentes
    """
    valor = (valor or "").strip()
    if not valor:
        return []

    possiveis = [valor]

    # reverse map: "LP" -> ["Portugues", "Português"]
    rev: Dict[str, List[str]] = {}
    for k, v in DISCIPLINA_ABREV.items():
        rev.setdefault(v, []).append(k)

    # Se o valor já é uma abrev (LP/MAT...), agrega nomes.
    if valor in rev:
        possiveis.extend(rev.get(valor, []))
        return possiveis

    # Se o valor é um nome (Português), agrega abrev correspondente.
    ab = DISCIPLINA_ABREV.get(valor)
    if ab and ab not in possiveis:
        possiveis.append(ab)
        possiveis.extend(rev.get(ab, []))

    # remove duplicados preservando ordem
    out = []
    seen = set()
    for x in possiveis:
        x = (x or "").strip()
        if x and x not in seen:
            seen.add(x)
            out.append(x)
    return out


def ensure_conselho_tables():
    """
    Cria tabelas necessárias sem quebrar o restante do sistema.
    Também tenta MIGRAR professor_turmas_disciplina se existir com coluna antiga.
    """
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("PRAGMA foreign_keys = ON")

        # Registro do professor por aluno + disciplina + bimestre/ano
        cur.execute("""
            CREATE TABLE IF NOT EXISTS conselhos_registros (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                professor_id INTEGER NOT NULL,
                turma_id INTEGER NOT NULL,
                aluno_id INTEGER NOT NULL,
                disciplina_abrev TEXT NOT NULL,
                bimestre INTEGER NOT NULL,
                ano INTEGER NOT NULL,
                aspectos_json TEXT NOT NULL,
                atualizado_em TEXT NOT NULL DEFAULT (datetime('now','localtime')),
                UNIQUE (turma_id, aluno_id, disciplina_abrev, bimestre, ano),
                FOREIGN KEY (professor_id) REFERENCES professores(id),
                FOREIGN KEY (turma_id) REFERENCES turmas(id),
                FOREIGN KEY (aluno_id) REFERENCES alunos(id)
            )
        """)

        # Vínculo do professor com a turma + disciplina (separar quando ele tem 2 disciplinas)
        # (Se já existir com outra estrutura, não quebra — apenas cria se não existir)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS professor_turmas_disciplina (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                professor_id INTEGER NOT NULL,
                turma_id INTEGER NOT NULL,
                disciplina_abrev TEXT,
                disciplina TEXT,
                UNIQUE (professor_id, turma_id, COALESCE(disciplina_abrev, disciplina)),
                FOREIGN KEY (professor_id) REFERENCES professores(id),
                FOREIGN KEY (turma_id) REFERENCES turmas(id)
            )
        """)

        # Observações do moderador por aluno + bimestre/ano
        cur.execute("""
            CREATE TABLE IF NOT EXISTS conselhos_observacoes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                turma_id INTEGER NOT NULL,
                aluno_id INTEGER NOT NULL,
                bimestre INTEGER NOT NULL,
                ano INTEGER NOT NULL,
                observacoes TEXT,
                atualizado_em TEXT NOT NULL DEFAULT (datetime('now','localtime')),
                UNIQUE (turma_id, aluno_id, bimestre, ano),
                FOREIGN KEY (turma_id) REFERENCES turmas(id),
                FOREIGN KEY (aluno_id) REFERENCES alunos(id)
            )
        """)

        # =========================
        # MIGRAÇÃO segura
        # =========================
        # Se a tabela já existe no seu banco, mas NÃO tem disciplina_abrev (tem disciplina),
        # adiciona disciplina_abrev e tenta preencher.
        if _tabela_existe("professor_turmas_disciplina"):
            cur.execute("PRAGMA table_info(professor_turmas_disciplina)")
            cols = [r[1] for r in cur.fetchall()]

            if "disciplina_abrev" not in cols:
                try:
                    cur.execute("ALTER TABLE professor_turmas_disciplina ADD COLUMN disciplina_abrev TEXT")
                except Exception:
                    pass

            # tenta preencher disciplina_abrev a partir de disciplina (se existir)
            if "disciplina" in cols:
                try:
                    cur.execute("SELECT id, disciplina, disciplina_abrev FROM professor_turmas_disciplina")
                    rows = cur.fetchall()
                    for r in rows:
                        if (r["disciplina_abrev"] or "").strip():
                            continue
                        disc_txt = (r["disciplina"] or "").strip()
                        ab = DISCIPLINA_ABREV.get(disc_txt)
                        if ab:
                            cur.execute(
                                "UPDATE professor_turmas_disciplina SET disciplina_abrev=? WHERE id=?",
                                (ab, r["id"])
                            )
                except Exception:
                    pass

        conn.commit()
    finally:
        cur.close()
        conn.close()


def _login_required(tipo: str) -> bool:
    if "usuario" not in session or session.get("tipo") != tipo:
        flash("Acesso nao autorizado.")
        return False
    return True


def _obter_professor_id(login: str) -> Optional[int]:
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("SELECT id FROM professores WHERE login = ? LIMIT 1", (login,))
        row = cur.fetchone()
        return row["id"] if row else None
    finally:
        cur.close()
        conn.close()


def _obter_professor_nome(login: str) -> str:
    """
    Tenta pegar um nome amigável do professor.
    Se não existir coluna 'nome', cai no login.
    """
    if not login:
        return "Professor(a)"

    conn = conectar_bd()
    cur = conn.cursor()
    try:
        # tenta com 'nome'
        try:
            cur.execute("SELECT nome, login FROM professores WHERE login = ? LIMIT 1", (login,))
            row = cur.fetchone()
            if row:
                nome = (row["nome"] or "").strip() if "nome" in row.keys() else ""
                if nome:
                    return nome
                return (row["login"] or login).strip()
        except Exception:
            # fallback seguro
            cur.execute("SELECT login FROM professores WHERE login = ? LIMIT 1", (login,))
            row = cur.fetchone()
            if row:
                return (row["login"] or login).strip()
    finally:
        cur.close()
        conn.close()

    return (login or "").strip() or "Professor(a)"


# =========================
# Helpers (Professor / Turmas / Disciplinas)
# =========================

def _disciplinas_professor_nomes(professor_id: int) -> List[str]:
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute(
            "SELECT disciplina FROM professor_disciplinas WHERE professor_id = ? ORDER BY disciplina",
            (professor_id,),
        )
        disciplinas = [(r["disciplina"] or "").strip() for r in cur.fetchall()]
        return [d for d in disciplinas if d]
    finally:
        cur.close()
        conn.close()


def _disciplinas_professor_abrev(professor_id: int) -> List[str]:
    nomes = _disciplinas_professor_nomes(professor_id)
    abrev: List[str] = []
    for d in nomes:
        a = DISCIPLINA_ABREV.get((d or "").strip())
        if a and a not in abrev:
            abrev.append(a)
    return abrev


def _turmas_professor(professor_id: int, disciplina_abrev: Optional[str] = None):
    """
    ✅ Robustez:
    - Se existir professor_turmas_disciplina e disciplina selecionada, tenta filtrar por disciplina.
    - Se não retornar nada (ou se não existir tabela/coluna), faz fallback para professor_turmas.
    """
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        ptd_col = _coluna_disciplina_ptd()

        # 1) tentativa filtrada por disciplina
        if disciplina_abrev and ptd_col:
            possiveis = _abrevs_possiveis_para(disciplina_abrev)
            if not possiveis:
                possiveis = [disciplina_abrev]

            placeholders = ",".join(["?"] * len(possiveis))

            cur.execute(f"""
                SELECT DISTINCT t.id, t.nome, t.turno
                FROM turmas t
                JOIN professor_turmas_disciplina ptd ON ptd.turma_id = t.id
                WHERE ptd.professor_id = ?
                  AND ptd.{ptd_col} IN ({placeholders})
                ORDER BY t.turno, t.nome
            """, (professor_id, *possiveis))

            rows = cur.fetchall()
            if rows:
                return rows  # ✅ achou turmas com disciplina

            # ✅ se não achou nada, cai no fallback (não trava o professor)

        # 2) fallback antigo (sempre deve funcionar se o vínculo estiver em professor_turmas)
        cur.execute("""
            SELECT DISTINCT t.id, t.nome, t.turno
            FROM turmas t
            JOIN professor_turmas pt ON pt.turma_id = t.id
            WHERE pt.professor_id = ?
            ORDER BY t.turno, t.nome
        """, (professor_id,))
        return cur.fetchall()

    finally:
        cur.close()
        conn.close()



# =========================
# Helpers (Turma / Aluno / Registros)
# =========================

def _alunos_da_turma(turma_id: int):
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("SELECT id, nome FROM alunos WHERE turma_id = ? ORDER BY nome", (turma_id,))
        return cur.fetchall()
    finally:
        cur.close()
        conn.close()


def _alunos_pendentes_professor_disciplina(
    turma_id: int,
    professor_id: int,
    disciplina_abrev: str,
    bimestre: int,
    ano: int
):
    """
    Retorna SOMENTE alunos que AINDA NÃO foram preenchidos pelo professor
    naquela turma + disciplina + bimestre + ano.
    """
    todos = _alunos_da_turma(turma_id)
    if not turma_id or not professor_id or not disciplina_abrev:
        return todos

    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT aluno_id
            FROM conselhos_registros
            WHERE turma_id=? AND professor_id=? AND disciplina_abrev=? AND bimestre=? AND ano=?
        """, (turma_id, professor_id, disciplina_abrev, bimestre, ano))
        feitos = {r["aluno_id"] for r in cur.fetchall()}
        return [a for a in todos if a["id"] not in feitos]
    finally:
        cur.close()
        conn.close()


def _turma_nome(turma_id: int) -> str:
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("SELECT nome, turno FROM turmas WHERE id = ? LIMIT 1", (turma_id,))
        row = cur.fetchone()
        if not row:
            return ""
        return f"{row['nome']} ({row['turno']})"
    finally:
        cur.close()
        conn.close()


def _aluno_nome(aluno_id: int) -> str:
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("SELECT nome FROM alunos WHERE id = ? LIMIT 1", (aluno_id,))
        row = cur.fetchone()
        return row["nome"] if row else ""
    finally:
        cur.close()
        conn.close()


def _registros_por_aluno(turma_id: int, aluno_id: int, bimestre: int, ano: int) -> Dict[str, Dict[str, Any]]:
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT disciplina_abrev, aspectos_json, atualizado_em, professor_id
            FROM conselhos_registros
            WHERE turma_id=? AND aluno_id=? AND bimestre=? AND ano=?
        """, (turma_id, aluno_id, bimestre, ano))
        rows = cur.fetchall()

        out: Dict[str, Dict[str, Any]] = {}
        for r in rows:
            try:
                asp = json.loads(r["aspectos_json"] or "{}")
            except Exception:
                asp = {}
            out[r["disciplina_abrev"]] = {
                "aspectos": asp,
                "atualizado_em": r["atualizado_em"],
                "professor_id": r["professor_id"],
            }
        return out
    finally:
        cur.close()
        conn.close()


def _observacao_moderador(turma_id: int, aluno_id: int, bimestre: int, ano: int) -> Dict[str, Any]:
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT observacoes, atualizado_em
            FROM conselhos_observacoes
            WHERE turma_id=? AND aluno_id=? AND bimestre=? AND ano=?
            LIMIT 1
        """, (turma_id, aluno_id, bimestre, ano))
        row = cur.fetchone()
        if not row:
            return {"observacoes": "", "atualizado_em": None}
        return {"observacoes": row["observacoes"] or "", "atualizado_em": row["atualizado_em"]}
    finally:
        cur.close()
        conn.close()


def _disciplinas_esperadas_turma_abrev(turma_id: int) -> List[str]:
    """
    Para o status do moderador ficar correto:
    - Se existir professor_turmas_disciplina, usa ela (preciso).
      ✅ Aceita coluna disciplina_abrev OU disciplina (banco antigo).
    - Senão, fallback para professor_turmas + professor_disciplinas (modo antigo).
    """
    conn = conectar_bd()
    cur = conn.cursor()
    try:
        abrev: List[str] = []

        ptd_col = _coluna_disciplina_ptd()
        if ptd_col:
            cur.execute(f"""
                SELECT DISTINCT ptd.{ptd_col} AS disc
                FROM professor_turmas_disciplina ptd
                JOIN professores p ON p.id = ptd.professor_id
                WHERE ptd.turma_id = ?
                  AND p.status = 'aprovado'
            """, (turma_id,))
            vals = [(r["disc"] or "").strip() for r in cur.fetchall() if (r["disc"] or "").strip()]

            for v in vals:
                # se já for abrev
                if v in COLUNAS_MODELO and v not in abrev:
                    abrev.append(v)
                    continue
                # se for nome, converte
                conv = DISCIPLINA_ABREV.get(v)
                if conv and conv not in abrev:
                    abrev.append(conv)

        else:
            # fallback: mistura disciplina do professor mesmo que ele não lecione nela na turma
            cur.execute("""
                SELECT DISTINCT pd.disciplina
                FROM professor_turmas pt
                JOIN professores p ON p.id = pt.professor_id
                JOIN professor_disciplinas pd ON pd.professor_id = p.id
                WHERE pt.turma_id = ?
                  AND p.status = 'aprovado'
            """, (turma_id,))
            disciplinas = [r["disciplina"] for r in cur.fetchall()]
            for d in disciplinas:
                a = DISCIPLINA_ABREV.get((d or "").strip())
                if a and a not in abrev:
                    abrev.append(a)

        # ordena conforme o modelo
        return [c for c in COLUNAS_MODELO if c in abrev]
    finally:
        cur.close()
        conn.close()


def _progressos_turma(turma_id: int, bimestre: int, ano: int) -> Dict[str, Any]:
    alunos = _alunos_da_turma(turma_id)
    esperadas = _disciplinas_esperadas_turma_abrev(turma_id)
    total_disc = len(esperadas)

    progresso_alunos = []
    feitos = 0

    conn = conectar_bd()
    cur = conn.cursor()
    try:
        for a in alunos:
            cur.execute("""
                SELECT COUNT(DISTINCT disciplina_abrev) AS total
                FROM conselhos_registros
                WHERE turma_id=? AND aluno_id=? AND bimestre=? AND ano=?
            """, (turma_id, a["id"], bimestre, ano))
            qtd = cur.fetchone()["total"] or 0

            ok = (total_disc > 0 and qtd >= total_disc)
            if ok:
                feitos += 1

            progresso_alunos.append({
                "aluno_id": a["id"],
                "aluno_nome": a["nome"],
                "qtd": qtd,
                "total": total_disc,
                "ok": ok,
            })

        return {
            "alunos": progresso_alunos,
            "feitos": feitos,
            "total_alunos": len(alunos),
            "disciplinas_esperadas": esperadas,
            "total_disciplinas": total_disc
        }
    finally:
        cur.close()
        conn.close()


# =========================
# DOCX helpers
# =========================

def _set_line_value(paragraph, label: str, value: str):
    txt = paragraph.text or ""
    if label in txt:
        prefix = txt.split(label, 1)[0] + label
        new_txt = (prefix + " " + value).strip()
        for r in paragraph.runs:
            r.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_txt
        else:
            paragraph.add_run(new_txt)


def _preencher_doc_modelo(doc: Document, turma_id: int, aluno_id: int, bimestre: int, ano: int):
    """Preenche um Document(modelo) para UM aluno."""
    if not doc.tables:
        raise RuntimeError("Modelo nao possui tabelas. Verifique o arquivo MODELO CONSELHO DE CLASSE.docx")

    turma_nome = _turma_nome(turma_id)
    aluno_nome = _aluno_nome(aluno_id)

    registros = _registros_por_aluno(turma_id, aluno_id, bimestre, ano)
    obs = _observacao_moderador(turma_id, aluno_id, bimestre, ano)

    # Atualiza cabeçalho
    for p in doc.paragraphs:
        if "CONSELHO DE CLASSE" in (p.text or "") and "BIMESTRE" in (p.text or ""):
            new_txt = f"CONSELHO DE CLASSE {bimestre}o BIMESTRE /{ano}"
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new_txt
            else:
                p.add_run(new_txt)

    # TURMA / ESTUDANTE
    for p in doc.paragraphs:
        _set_line_value(p, "TURMA:", turma_nome)
        _set_line_value(p, "ESTUDANTE:", aluno_nome)

    tabela = doc.tables[0]

    header = [tabela.cell(0, c).text.strip().replace("\n", " ") for c in range(len(tabela.columns))]
    col_index = {h: i for i, h in enumerate(header) if h in COLUNAS_MODELO}

    row_index: Dict[str, int] = {}
    for r in range(1, len(tabela.rows)):
        label = tabela.cell(r, 0).text.strip().replace("\n", " ")
        for key, nome in ASPECTOS:
            if label == nome:
                row_index[key] = r

    # Limpa tudo (evita “herdar X” do modelo)
    for asp_key, _asp_name in ASPECTOS:
        rr = row_index.get(asp_key)
        if rr is None:
            continue
        for col in col_index.values():
            tabela.cell(rr, col).text = ""

    # Aplica X onde foi marcado
    for disc_abrev, payload in registros.items():
        if disc_abrev not in col_index:
            continue
        c = col_index[disc_abrev]
        aspectos = payload.get("aspectos", {}) or {}
        for asp_key, _asp_name in ASPECTOS:
            r = row_index.get(asp_key)
            if r is None:
                continue
            marcar = bool(aspectos.get(asp_key))
            tabela.cell(r, c).text = "X" if marcar else ""

    # Observações em segunda tabela (se existir)
    if len(doc.tables) > 1:
        obs_tbl = doc.tables[1]
        cell = obs_tbl.cell(0, 0)
        cell.text = "Observacoes:\n" + (obs.get("observacoes") or "").strip()


def _render_docx_conselho(turma_id: int, aluno_id: int, bimestre: int, ano: int) -> str:
    if not os.path.exists(MODELO_DOCX_PATH):
        raise FileNotFoundError(f"Modelo nao encontrado: {MODELO_DOCX_PATH}")

    doc = Document(MODELO_DOCX_PATH)
    _preencher_doc_modelo(doc, turma_id, aluno_id, bimestre, ano)

    out_name = f"conselho_t{turma_id}_a{aluno_id}_b{bimestre}_{ano}.docx"
    out_path = os.path.join(CONSELHO_OUT_DIR, out_name)
    doc.save(out_path)
    return out_path


def _render_docx_turma_unico(turma_id: int, bimestre: int, ano: int) -> str:
    """Gera um ÚNICO DOCX com TODOS os alunos da turma (1 aluno por página)."""
    if not os.path.exists(MODELO_DOCX_PATH):
        raise FileNotFoundError(f"Modelo nao encontrado: {MODELO_DOCX_PATH}")

    alunos = _alunos_da_turma(turma_id)
    if not alunos:
        raise RuntimeError("Turma sem alunos para gerar.")

    # 1) Primeiro aluno vira o documento mestre
    master = Document(MODELO_DOCX_PATH)
    _preencher_doc_modelo(master, turma_id, alunos[0]["id"], bimestre, ano)

    # 2) Demais alunos: cria doc temp preenchido e “anexa” no master com quebra de página
    for a in alunos[1:]:
        temp = Document(MODELO_DOCX_PATH)
        _preencher_doc_modelo(temp, turma_id, a["id"], bimestre, ano)

        master.add_page_break()

        # Copia elementos do body (exceto sectPr final)
        for el in list(temp.element.body):
            if el.tag.endswith("}sectPr"):
                continue
            master.element.body.append(deepcopy(el))

    out_name = f"conselho_TURMA_{turma_id}_b{bimestre}_{ano}.docx"
    out_path = os.path.join(CONSELHO_OUT_DIR, out_name)
    master.save(out_path)
    return out_path


# =========================
# PROFESSOR
# =========================

@bp_conselho.route("/conselho/professor", methods=["GET", "POST"])
def conselho_professor():
    if not _login_required("professor"):
        return redirect(url_for("login"))

    professor_login = session.get("usuario")
    professor_id = _obter_professor_id(professor_login)
    if not professor_id:
        flash("Professor nao encontrado.")
        return redirect(url_for("dashboard_professor"))

    professor_nome = _obter_professor_nome(professor_login)
    professor_disciplinas_nomes = _disciplinas_professor_nomes(professor_id)

    # DISCIPLINAS do professor (abreviações do modelo)
    disciplinas = _disciplinas_professor_abrev(professor_id)

    turma_id = (request.values.get("turma_id", "") or "").strip()
    aluno_id = (request.values.get("aluno_id", "") or "").strip()
    disc = (request.values.get("disciplina_abrev", "") or "").strip()
    bimestre = _as_int(request.values.get("bimestre", "1"), 1) or 1
    ano = _as_int(request.values.get("ano", str(datetime.now().year)), datetime.now().year) or datetime.now().year

    turma_id_i = _as_int(turma_id, 0) or 0
    aluno_id_i = _as_int(aluno_id, 0) or 0

    # Turmas filtradas por disciplina (quando selecionada)
    turmas = _turmas_professor(professor_id, disc if disc else None)

    modo = (request.values.get("modo", "") or "").strip().lower()  # "" ou "editar"
    is_editar = (modo == "editar")

    # Lista de alunos
    if turma_id_i:
        if disc:
            if is_editar:
                # ✅ edição: mostra TODOS os alunos da turma
                alunos = _alunos_da_turma(turma_id_i)
            else:
                # padrão: só pendentes
                alunos = _alunos_pendentes_professor_disciplina(turma_id_i, professor_id, disc, bimestre, ano)

                # ✅ se veio um aluno_id específico na URL (ex: botão "Editar"),
                # garante que ele apareça na lista, mesmo já concluído
                if aluno_id_i and all(a["id"] != aluno_id_i for a in alunos):
                    try:
                        conn2 = conectar_bd()
                        cur2 = conn2.cursor()
                        cur2.execute(
                            "SELECT id, nome FROM alunos WHERE id=? AND turma_id=? LIMIT 1",
                            (aluno_id_i, turma_id_i)
                        )
                        row_al = cur2.fetchone()
                        if row_al:
                            alunos = alunos + [{"id": row_al["id"], "nome": row_al["nome"]}]
                    finally:
                        try: cur2.close()
                        except Exception: pass
                        try: conn2.close()
                        except Exception: pass
        else:
            alunos = _alunos_da_turma(turma_id_i)
    else:
        alunos = []

    # ❌ REMOVE este comportamento em edição:
    # "Se o aluno selecionado sumiu, limpa seleção"
    # Em modo editar a seleção NÃO pode sumir.
    if (not is_editar) and aluno_id_i and alunos and all(a["id"] != aluno_id_i for a in alunos):
        aluno_id = ""
        aluno_id_i = 0

    # Se o aluno selecionado "sumiu" (porque já foi feito), limpa seleção
    if aluno_id_i and alunos and all(a["id"] != aluno_id_i for a in alunos):
        aluno_id = ""
        aluno_id_i = 0

    registro_existente: Dict[str, Any] = {}
    if turma_id_i and aluno_id_i and disc:
        conn = conectar_bd()
        cur = conn.cursor()
        try:
            cur.execute("""
                SELECT aspectos_json
                FROM conselhos_registros
                WHERE turma_id=? AND aluno_id=? AND disciplina_abrev=? AND bimestre=? AND ano=?
                  AND professor_id=?
                LIMIT 1
            """, (turma_id_i, aluno_id_i, disc, bimestre, ano, professor_id))
            row = cur.fetchone()
            if row:
                try:
                    registro_existente = json.loads(row["aspectos_json"] or "{}")
                except Exception:
                    registro_existente = {}
        finally:
            cur.close()
            conn.close()

    if request.method == "POST":
        acao = (request.form.get("acao") or "").strip()

        # quando for "só mudar select", NÃO salva
        if acao != "salvar":
            return redirect(url_for(
                "conselho.conselho_professor",
                turma_id=request.form.get("turma_id", ""),
                aluno_id=request.form.get("aluno_id", ""),
                disciplina_abrev=request.form.get("disciplina_abrev", ""),
                bimestre=request.form.get("bimestre", "1"),
                ano=request.form.get("ano", str(datetime.now().year))
            ))

        turma_id_i = _as_int(request.form.get("turma_id"))
        aluno_id_i = _as_int(request.form.get("aluno_id"))
        disc = (request.form.get("disciplina_abrev") or "").strip()
        bimestre = _as_int(request.form.get("bimestre"), 1) or 1
        ano = _as_int(request.form.get("ano"), datetime.now().year) or datetime.now().year

        if not turma_id_i or not aluno_id_i or not disc:
            flash("Selecione Turma, Aluno e Disciplina antes de salvar.")
            return redirect(url_for(
                "conselho.conselho_professor",
                turma_id=request.form.get("turma_id", ""),
                aluno_id=request.form.get("aluno_id", ""),
                disciplina_abrev=disc,
                bimestre=bimestre,
                ano=ano
            ))

        if disc not in disciplinas:
            flash("Disciplina invalida para este professor.")
            return redirect(url_for(
                "conselho.conselho_professor",
                turma_id=turma_id_i,
                aluno_id=aluno_id_i,
                disciplina_abrev=disc,
                bimestre=bimestre,
                ano=ano
            ))

        aspectos = {key: (1 if request.form.get(f"asp_{key}") == "1" else 0) for key, _ in ASPECTOS}

        conn = conectar_bd()
        cur = conn.cursor()
        try:
            cur.execute("""
                INSERT INTO conselhos_registros
                    (professor_id, turma_id, aluno_id, disciplina_abrev, bimestre, ano, aspectos_json, atualizado_em)
                VALUES (?, ?, ?, ?, ?, ?, ?, datetime('now','localtime'))
                ON CONFLICT(turma_id, aluno_id, disciplina_abrev, bimestre, ano)
                DO UPDATE SET
                    professor_id=excluded.professor_id,
                    aspectos_json=excluded.aspectos_json,
                    atualizado_em=datetime('now','localtime')
            """, (professor_id, turma_id_i, aluno_id_i, disc, bimestre, ano, json.dumps(aspectos, ensure_ascii=False)))
            conn.commit()
            flash("Conselho salvo com sucesso!")
        except Exception as e:
            conn.rollback()
            flash(f"Erro ao salvar: {e}")
        finally:
            cur.close()
            conn.close()

        # Depois de salvar: mantém turma/discipiplina/bimestre/ano e limpa aluno (vai pro próximo pendente)
        return redirect(url_for(
            "conselho.conselho_professor",
            turma_id=turma_id_i,
            aluno_id="",
            disciplina_abrev=disc,
            bimestre=bimestre,
            ano=ano
        ))

    return render_template(
        "conselho_professor.html",
        turmas=turmas,
        disciplinas=disciplinas,
        alunos=alunos,
        turma_id=turma_id,
        aluno_id=aluno_id,
        disciplina_abrev=disc,
        bimestre=bimestre,
        ano=ano,
        aspectos=ASPECTOS,
        registro=registro_existente,
        professor_nome=professor_nome,
        professor_disciplinas_nomes=professor_disciplinas_nomes
    )


@bp_conselho.route("/conselho/professor/visualizar", methods=["GET"])
def conselho_professor_visualizar():
    if not _login_required("professor"):
        return redirect(url_for("login"))

    professor_id = _obter_professor_id(session.get("usuario"))
    if not professor_id:
        flash("Professor nao encontrado.")
        return redirect(url_for("dashboard_professor"))

    bimestre = _as_int(request.args.get("bimestre", "1"), 1) or 1
    ano = _as_int(request.args.get("ano", str(datetime.now().year)), datetime.now().year) or datetime.now().year

    disc_prof = _disciplinas_professor_abrev(professor_id)
    turmas = _turmas_professor(professor_id)

    progresso_turmas = []

    conn = conectar_bd()
    cur = conn.cursor()
    try:
        for t in turmas:
            alunos = _alunos_da_turma(t["id"])
            total_alunos = len(alunos)
            total_disc = len(disc_prof) if disc_prof else 0

            feitos = 0
            for a in alunos:
                cur.execute("""
                    SELECT COUNT(DISTINCT disciplina_abrev) AS total
                    FROM conselhos_registros
                    WHERE turma_id=? AND aluno_id=? AND bimestre=? AND ano=?
                      AND professor_id=?
                """, (t["id"], a["id"], bimestre, ano, professor_id))
                qtd = cur.fetchone()["total"] or 0
                if total_disc > 0 and qtd >= total_disc:
                    feitos += 1

            progresso_turmas.append({
                "turma_id": t["id"],
                "turma_nome": f"{t['nome']} ({t['turno']})",
                "feitos": feitos,
                "total": total_alunos
            })
    finally:
        cur.close()
        conn.close()

    return render_template(
        "conselho_professor_turmas.html",
        progresso_turmas=progresso_turmas,
        bimestre=bimestre,
        ano=ano
    )


@bp_conselho.route("/conselho/professor/turma/<int:turma_id>", methods=["GET"])
def conselho_professor_turma(turma_id: int):
    """
    Lista alunos da turma com status POR DISCIPLINA do professor (quando ele tem 2 disciplinas).
    Template sugerido: conselho_professor_turma.html
    """
    if not _login_required("professor"):
        return redirect(url_for("login"))

    professor_id = _obter_professor_id(session.get("usuario"))
    if not professor_id:
        flash("Professor nao encontrado.")
        return redirect(url_for("dashboard_professor"))

    bimestre = _as_int(request.args.get("bimestre", "1"), 1) or 1
    ano = _as_int(request.args.get("ano", str(datetime.now().year)), datetime.now().year) or datetime.now().year

    alunos = _alunos_da_turma(turma_id)
    disc_prof = _disciplinas_professor_abrev(professor_id)

    conn = conectar_bd()
    cur = conn.cursor()
    try:
        itens = []
        feitos = 0

        for a in alunos:
            cur.execute("""
                SELECT disciplina_abrev
                FROM conselhos_registros
                WHERE turma_id=? AND aluno_id=? AND bimestre=? AND ano=? AND professor_id=?
            """, (turma_id, a["id"], bimestre, ano, professor_id))
            feitas_set = {r["disciplina_abrev"] for r in cur.fetchall()}

            status_por_disc = {}
            qtd_ok = 0
            for d in disc_prof:
                ok_d = (d in feitas_set)
                status_por_disc[d] = ok_d
                if ok_d:
                    qtd_ok += 1

            ok_total = (len(disc_prof) > 0 and qtd_ok >= len(disc_prof))
            if ok_total:
                feitos += 1

            itens.append({
                "aluno_id": a["id"],
                "aluno_nome": a["nome"],
                "status_por_disc": status_por_disc,
                "qtd": qtd_ok,
                "total": len(disc_prof),
                "ok": ok_total
            })

        return render_template(
            "conselho_professor_turma.html",
            turma_id=turma_id,
            turma_nome=_turma_nome(turma_id),
            itens=itens,
            disciplinas=disc_prof,
            feitos=feitos,
            total=len(itens),
            bimestre=bimestre,
            ano=ano
        )
    finally:
        cur.close()
        conn.close()


@bp_conselho.route("/conselho/professor/ver", methods=["GET"])
def conselho_professor_ver():
    """
    Consolidado para o professor:
    - mostra apenas as colunas das disciplinas do professor (LP/MAT etc.)
    - permite filtrar por uma disciplina (disciplina_abrev)
    Template: conselho_consolidado.html
    """
    if not _login_required("professor"):
        return redirect(url_for("login"))

    professor_id = _obter_professor_id(session.get("usuario"))
    if not professor_id:
        flash("Professor nao encontrado.")
        return redirect(url_for("dashboard_professor"))

    turma_id = _as_int(request.args.get("turma_id"))
    aluno_id = _as_int(request.args.get("aluno_id"))
    bimestre = _as_int(request.args.get("bimestre", "1"), 1) or 1
    ano = _as_int(request.args.get("ano", str(datetime.now().year)), datetime.now().year) or datetime.now().year

    if not turma_id or not aluno_id:
        flash("Turma e aluno sao obrigatorios para visualizar.")
        return redirect(url_for("conselho.conselho_professor_visualizar"))

    disc_prof = _disciplinas_professor_abrev(professor_id)
    disciplina_filtro = (request.args.get("disciplina_abrev") or "").strip()

    if disciplina_filtro and disciplina_filtro in disc_prof:
        colunas = [disciplina_filtro]
    else:
        colunas = disc_prof[:]  # mostra todas do professor

    registros = _registros_por_aluno(turma_id, aluno_id, bimestre, ano)
    obs = _observacao_moderador(turma_id, aluno_id, bimestre, ano)

    return render_template(
        "conselho_consolidado.html",
        aluno_nome=_aluno_nome(aluno_id),
        turma_nome=_turma_nome(turma_id),
        bimestre=bimestre,
        ano=ano,
        colunas=colunas,
        aspectos=ASPECTOS,
        registros=registros,
        observacao=obs,
        disciplina_filtro=disciplina_filtro,
        disciplinas_prof=disc_prof
    )


# =========================
# MODERADOR
# =========================

@bp_conselho.route("/conselho/moderador", methods=["GET"])
def conselho_moderador():
    if not _login_required("moderador"):
        return redirect(url_for("login"))

    conn = conectar_bd()
    cur = conn.cursor()
    try:
        cur.execute("SELECT id, nome, turno FROM turmas ORDER BY turno, nome")
        turmas = cur.fetchall()
    finally:
        cur.close()
        conn.close()

    bimestre = _as_int(request.args.get("bimestre", "1"), 1) or 1
    ano = _as_int(request.args.get("ano", str(datetime.now().year)), datetime.now().year) or datetime.now().year

    return render_template("conselho_moderador.html", turmas=turmas, bimestre=bimestre, ano=ano)


@bp_conselho.route("/conselho/moderador/turma/<int:turma_id>", methods=["GET", "POST"])
def conselho_moderador_turma(turma_id):
    if not _login_required("moderador"):
        return redirect(url_for("login"))

    bimestre = _as_int(request.values.get("bimestre", "1"), 1) or 1
    ano = _as_int(request.values.get("ano", str(datetime.now().year)), datetime.now().year) or datetime.now().year

    progresso = _progressos_turma(turma_id, bimestre, ano)

    if request.method == "POST":
        aluno_id = _as_int(request.form.get("aluno_id"))
        observacoes = (request.form.get("observacoes") or "").strip()

        if not aluno_id:
            flash("Selecione um aluno para salvar observacoes.")
            return redirect(url_for("conselho.conselho_moderador_turma", turma_id=turma_id, bimestre=bimestre, ano=ano))

        conn = conectar_bd()
        cur = conn.cursor()
        try:
            cur.execute("""
                INSERT INTO conselhos_observacoes (turma_id, aluno_id, bimestre, ano, observacoes, atualizado_em)
                VALUES (?, ?, ?, ?, ?, datetime('now','localtime'))
                ON CONFLICT(turma_id, aluno_id, bimestre, ano)
                DO UPDATE SET observacoes=excluded.observacoes, atualizado_em=datetime('now','localtime')
            """, (turma_id, aluno_id, bimestre, ano, observacoes))
            conn.commit()
            flash("Observacoes salvas.")
        except Exception as e:
            conn.rollback()
            flash(f"Erro ao salvar observacoes: {e}")
        finally:
            cur.close()
            conn.close()

        return redirect(url_for("conselho.conselho_moderador_turma", turma_id=turma_id, bimestre=bimestre, ano=ano))

    return render_template(
        "conselho_moderador_turma.html",
        turma_id=turma_id,
        turma_nome=_turma_nome(turma_id),
        bimestre=bimestre,
        ano=ano,
        progresso=progresso
    )


@bp_conselho.route("/conselho/moderador/ver/<int:turma_id>/<int:aluno_id>", methods=["GET"])
def conselho_moderador_ver(turma_id, aluno_id):
    if not _login_required("moderador"):
        return redirect(url_for("login"))

    bimestre = _as_int(request.args.get("bimestre", "1"), 1) or 1
    ano = _as_int(request.args.get("ano", str(datetime.now().year)), datetime.now().year) or datetime.now().year

    registros = _registros_por_aluno(turma_id, aluno_id, bimestre, ano)
    obs = _observacao_moderador(turma_id, aluno_id, bimestre, ano)
    esperadas = _disciplinas_esperadas_turma_abrev(turma_id)

    return render_template(
        "conselho_ver.html",
        modo="moderador",
        turma_nome=_turma_nome(turma_id),
        aluno_nome=_aluno_nome(aluno_id),
        turma_id=turma_id,
        aluno_id=aluno_id,
        bimestre=bimestre,
        ano=ano,
        colunas=COLUNAS_MODELO,
        disciplinas_esperadas=esperadas,
        aspectos=ASPECTOS,
        registros=registros,
        observacao=obs
    )


# ✅ INDIVIDUAL: gera e baixa direto
@bp_conselho.route("/conselho/moderador/gerar/<int:turma_id>/<int:aluno_id>", methods=["POST"])
def conselho_moderador_gerar(turma_id, aluno_id):
    if not _login_required("moderador"):
        return redirect(url_for("login"))

    bimestre = _as_int(request.form.get("bimestre"), 1) or 1
    ano = _as_int(request.form.get("ano"), datetime.now().year) or datetime.now().year

    try:
        out_path = _render_docx_conselho(turma_id, aluno_id, bimestre, ano)
        filename = os.path.basename(out_path)
        return send_file(out_path, as_attachment=True, download_name=filename)
    except Exception as e:
        flash(f"Erro ao gerar Word individual: {e}")
        return redirect(url_for("conselho.conselho_moderador_turma", turma_id=turma_id, bimestre=bimestre, ano=ano))


# ✅ TURMA: gera um único Word e baixa direto
@bp_conselho.route("/conselho/moderador/gerar_turma/<int:turma_id>", methods=["POST"])
def conselho_moderador_gerar_turma(turma_id):
    if not _login_required("moderador"):
        return redirect(url_for("login"))

    bimestre = _as_int(request.form.get("bimestre"), 1) or 1
    ano = _as_int(request.form.get("ano"), datetime.now().year) or datetime.now().year

    try:
        out_path = _render_docx_turma_unico(turma_id, bimestre, ano)
        filename = os.path.basename(out_path)
        return send_file(out_path, as_attachment=True, download_name=filename)
    except Exception as e:
        flash(f"Erro ao gerar Word da turma: {e}")
        return redirect(url_for("conselho.conselho_moderador_turma", turma_id=turma_id, bimestre=bimestre, ano=ano))


# =========================
# Inicialização do módulo
# =========================
ensure_conselho_tables()